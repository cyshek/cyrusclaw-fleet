from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil, os, zipfile

BUNDLE_DIR = '/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/lance-pm'
os.makedirs(BUNDLE_DIR, exist_ok=True)

JD_SRC = '/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/lance-294e2602-725c-4d84-a7a0-a1fc602acec8/JD.md'
RESUME_SRC = '/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/lance-294e2602-725c-4d84-a7a0-a1fc602acec8/Cyrus_Shekari_Resume_ashby-lance_294e2602_v2.pdf'
MASTER_SRC = '/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Master_Interview_Prep_Guide.docx'
SHORTER_SRC = '/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Shorter_Master_Interview_Prep_Guide.docx'

FULL_DEST = os.path.join(BUNDLE_DIR, 'Lance_PM_Interview_Prep_Guide.docx')
SHORTER_DEST = os.path.join(BUNDLE_DIR, 'Lance_PM_Shorter_Interview_Prep_Guide.docx')
JD_DEST = os.path.join(BUNDLE_DIR, 'Lance_PM_JD.md')
RESUME_DEST = os.path.join(BUNDLE_DIR, 'Cyrus_Shekari_Resume_Lance_PM.pdf')

shutil.copy2(MASTER_SRC, FULL_DEST)
shutil.copy2(SHORTER_SRC, SHORTER_DEST)
shutil.copy2(JD_SRC, JD_DEST)
shutil.copy2(RESUME_SRC, RESUME_DEST)

# Q2 content
Q2_LABEL1 = 'Why do I want to work at Lance:'
Q2_TEXT1  = 'Lance is building autonomous AI agents that handle real hotel operations end-to-end — not just a chatbot layer on top of existing software, but agents that actually take actions inside live systems. That combination of Computer Use Agent technology and a focused vertical (hospitality) is a genuinely interesting product problem, and the traction with 50+ hotels across Marriott, Hilton, and Hyatt shows the approach is working.'
Q2_LABEL2 = 'Why I want this role:'
Q2_TEXT2  = 'The PM role here maps directly to what I\'ve been doing — owning a roadmap, coordinating cross-functional teams, and driving a product from ambiguous early-stage through scaled execution. I\'m drawn to the pace of a YC company where decisions move fast and the PM has real influence over what gets built.'

# Q3 content
Q3_LABEL1 = 'What Lance does:'
Q3_TEXT1  = 'Lance builds autonomous AI agents for the hospitality industry — agents that handle calls, close sales, and manage operations for hotels, integrating directly into existing hotel software to make real-time decisions. Backed by YC W26, they work with 50+ hotels across major brands.'
Q3_LABEL2 = 'What this role is:'
Q3_TEXT2  = 'The Product Manager owns the roadmap and lifecycle for assigned products, working across engineering, design, and business teams to define requirements, prioritize features, and ship. It\'s an in-person role in San Francisco with a strong focus on AI and SaaS product delivery.'

PLACEHOLDER_Q2 = '[Fill in here]'
PLACEHOLDER_Q3 = '[Fill in here]'

def make_bold_plain_run(para, label, text):
    """Add a bold label run + plain text run to a paragraph, clearing existing runs first."""
    # Clear existing runs
    for r in para.runs:
        r.text = ''
    # Bold label
    r1 = para.add_run(label + ' ')
    r1.bold = True
    # Plain text
    r2 = para.add_run(text)
    r2.bold = False

def fill_q2_q3(dest_path, q2_label1, q2_text1, q2_label2, q2_text2,
               q3_label1, q3_text1, q3_label2, q3_text2):
    doc = Document(dest_path)
    in_q2 = False
    in_q3 = False
    q2_fills = 0
    q3_fills = 0

    for p in doc.paragraphs:
        t = p.text.strip()

        if 'Q2:' in t and ('Why' in t or 'here' in t or 'company' in t):
            in_q2 = True; in_q3 = False
            continue
        if 'Q3:' in t and ('do we do' in t or 'What' in t or 'know' in t):
            in_q3 = True; in_q2 = False
            continue
        if ('Q4:' in t or 'Q1:' in t) and (in_q2 or in_q3):
            in_q2 = False; in_q3 = False
            continue

        if (in_q2 or in_q3) and PLACEHOLDER_Q2 in t:\n            if in_q2:
                if q2_fills == 0:
                    make_bold_plain_run(p, q2_label1, q2_text1)
                    q2_fills += 1
                elif q2_fills == 1:
                    make_bold_plain_run(p, q2_label2, q2_text2)
                    q2_fills += 1
            else:
                if q3_fills == 0:
                    make_bold_plain_run(p, q3_label1, q3_text1)
                    q3_fills += 1
                elif q3_fills == 1:
                    make_bold_plain_run(p, q3_label2, q3_text2)
                    q3_fills += 1

    doc.save(dest_path)
    return q2_fills, q3_fills

r = fill_q2_q3(FULL_DEST,
    Q2_LABEL1, Q2_TEXT1, Q2_LABEL2, Q2_TEXT2,
    Q3_LABEL1, Q3_TEXT1, Q3_LABEL2, Q3_TEXT2)
print(f'Full guide: Q2 fills={r[0]}, Q3 fills={r[1]}')

r = fill_q2_q3(SHORTER_DEST,
    Q2_LABEL1, Q2_TEXT1, Q2_LABEL2, Q2_TEXT2,
    Q3_LABEL1, Q3_TEXT1, Q3_LABEL2, Q3_TEXT2)
print(f'Shorter guide: Q2 fills={r[0]}, Q3 fills={r[1]}')

# Verify no placeholders remain + bold labels intact
for path, label in [(FULL_DEST, 'Full'), (SHORTER_DEST, 'Shorter')]:
    doc = Document(path)
    placeholders = [p.text for p in doc.paragraphs if PLACEHOLDER_Q2 in p.text]
    q2q3_bullets = []
    in_q = False
    for p in doc.paragraphs:
        if 'Q2:' in p.text or 'Q3:' in p.text:
            in_q = True
        if in_q and p.runs and p.runs[0].bold and ':' in (p.runs[0].text or ''):
            if p.text.strip() not in ('Script:',):
                q2q3_bullets.append((p.runs[0].text, p.runs[0].bold))
        if 'Q4:' in p.text:
            in_q = False
    print(f'\n{label} — placeholders remaining: {len(placeholders)}')
    print(f'{label} — Q2/Q3 bold labels: {q2q3_bullets}')

# Zip
zip_path = '/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/Lance_PM_Bundle.zip'
with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
    for f in [FULL_DEST, SHORTER_DEST, JD_DEST, RESUME_DEST]:
        zf.write(f, os.path.basename(f))
print(f'\nBundle zipped: {zip_path}')
