#!/usr/bin/env python3
"""
Rebuild the Datadog PTSE guide so it inherits the MASTER's exact styling
(Arial body / Georgia headings, 11pt, master Heading1/2/3 + 'normal').

Faithful approach:
  - Start from a LITERAL copy of master.docx (keeps theme/styles/fonts/sizes/numbering defs).
  - Replace Q2/Q3 `[Fill in here]` bullets with SHORTENED answers (reusing the master bullet paragraph,
    so they look exactly like every other bullet).
  - Append Section 5 by DEEP-COPYING each paragraph's XML from the pandoc doc into the master copy,
    then remapping each pasted paragraph's style to a style that exists in the master
    (Heading 2/3 or 'normal'), and importing any numbering so bullets survive.

Only writes the per-company copy. templates/ master is never touched.
"""
import copy, shutil
from docx import Document
from docx.oxml.ns import qn

MASTER = "templates/Master_Interview_Prep_Guide.docx"
PANDOC = "guides/datadog-partner-technology-solutions-engineer/Datadog_PTSE_Interview_Prep_Guide.docx"
OUT    = "guides/datadog-partner-technology-solutions-engineer/Datadog_PTSE_Interview_Prep_Guide.docx"

Q2_BULLETS = [
    "What excites me about Datadog is that you've become the observability standard for the AI era, and the Partner TSE role sits at the exact intersection of what I love: deep distributed-systems knowledge, hands-on technical consulting, and turning hard engineering into something other developers can adopt.",
    "My whole career has been taking complex infrastructure and making it accessible through platforms and automation. At Microsoft I turned a manual recovery process into a self-service platform, and spent that time as the technical bridge between platform engineering, availability teams, and service owners. A Partner TSE is that same translator role, pointed at an external developer community.",
    "I'm also an observability person at heart: in my recovery work, telemetry was the product. Metrics, logs, and traces were how I proved systems survived failure and surfaced latent defects before production. Datadog isn't abstract to me, it's the category I already live in from the customer side.",
    "And the role rewards what I'm best at: building 0-to-1 and driving adoption. The JD calls out spotting friction in the Integration Developer Platform and partnering with Product and Eng, and I've run that exact loop, taking operator feedback, influencing a platform roadmap, and building onboarding to drive adoption.",
]
Q3_BULLETS = [
    "Absolutely. Datadog is the leading observability and security platform for the AI era. You give engineering teams unified, end-to-end visibility across their whole stack in one place, built on the three pillars, metrics, logs, and traces, through products like Infrastructure Monitoring, APM, and Log Management.",
    "A big part of why Datadog wins is the integration ecosystem, 1,000-plus integrations that pull every tool in a customer's stack into a single pane of glass. The more Datadog can see, the more valuable it gets, and that ecosystem is exactly where this role lives.",
    "The Partner TSE is the technical bridge between Datadog and your third-party developer community. I'd be the primary technical contact guiding partners through the whole integration lifecycle: from architecture, through building on the IDP (OAuth, log pipelines, OpenTelemetry, agent vs API-based), to publication on the Datadog Marketplace after passing your Quality Rubric.",
    "So it's part deep technical troubleshooting, architectural and code reviews plus solving partner issues, and part consulting and platform advocacy, feeding the friction I see back to Product and Eng. It's high-leverage technical work where my contribution shows up directly in the ecosystem, which is exactly what I want to be doing.",
]

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def find_idx(doc, contains, start=0):
    for i in range(start, len(doc.paragraphs)):
        if contains in doc.paragraphs[i].text:
            return i
    return None

def set_bullet_text(p, text):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(text)

def replace_placeholder_block(doc, heading_contains, next_contains, bullets):
    qi = find_idx(doc, heading_contains)
    assert qi is not None, f"heading not found: {heading_contains}"
    ph = None
    for j in range(qi+1, len(doc.paragraphs)):
        if "Fill in here" in doc.paragraphs[j].text:
            ph = j; break
        if next_contains in doc.paragraphs[j].text:
            break
    assert ph is not None, f"placeholder not found after {heading_contains}"
    ph_para = doc.paragraphs[ph]
    set_bullet_text(ph_para, bullets[0])
    anchor = ph_para._p
    clones = []
    for _ in bullets[1:]:
        newp = copy.deepcopy(ph_para._p)
        anchor.addnext(newp)
        anchor = newp
        clones.append(newp)
    # set text on clones
    from docx.text.paragraph import Paragraph
    for cp_el, btxt in zip(clones, bullets[1:]):
        cp = Paragraph(cp_el, ph_para._parent)
        set_bullet_text(cp, btxt)
    print(f"  filled {heading_contains[:25]}: {len(bullets)} bullets")

def main():
    shutil.copyfile(MASTER, "/tmp/_mcopy.docx")
    doc = Document("/tmp/_mcopy.docx")

    # 1) Q2 / Q3 shortened answers
    replace_placeholder_block(doc, "Why do you want to work at our company", "Do you know what this role is", Q2_BULLETS)
    replace_placeholder_block(doc, "Do you know what this role is", "Why are you looking to leave", Q3_BULLETS)

    # 2) Build a style-name map present in master
    master_styles = set(s.name for s in doc.styles)
    def remap(nm):
        if nm and nm.startswith("Heading 2"): return "Heading 2"
        if nm and nm.startswith("Heading 3"): return "Heading 3"
        if nm and nm.startswith("Heading 4"): return "Heading 3"   # master has no H4
        if nm and nm.startswith("Heading 1"): return "Heading 1"
        return "normal"

    # 3) Grab a reference MASTER bullet paragraph (style='normal' + inline numPr/numId).
    #    Clone THIS for every Section 5 bullet so numbering+font match the rest of the doc exactly.
    ref_bullet_el = None
    for p in doc.paragraphs:
        ppr = p._p.find(qn('w:pPr'))
        if ppr is not None and ppr.find(qn('w:numPr')) is not None and (p.style and p.style.name == 'normal'):
            ref_bullet_el = p._p
            break
    print("  reference master bullet found:", ref_bullet_el is not None)

    def add_master_bullet(text):
        from docx.text.paragraph import Paragraph
        newp = copy.deepcopy(ref_bullet_el)
        for r in newp.findall(qn('w:r')):
            newp.remove(r)
        doc.element.body.append(newp)
        para = Paragraph(newp, doc.paragraphs[-1]._parent)
        para.add_run(text)
        return para

    src = Document(PANDOC)
    doc.add_page_break()

    porting = False
    n = 0
    for p in src.paragraphs:
        t = p.text
        s = t.strip()
        if not porting:
            if s.startswith("Section 5:"):
                porting = True
            else:
                continue
        if not s:
            continue
        srcsty = p.style.name if p.style else None
        ppr = p._p.find(qn('w:pPr'))
        has_num = ppr is not None and ppr.find(qn('w:numPr')) is not None

        if s.startswith("Section 5:"):
            np = doc.add_paragraph(); np.style = doc.styles["Heading 2"]; np.add_run(t)
        elif s.startswith("Bucket "):
            np = doc.add_paragraph(); np.style = doc.styles["Heading 3"]; np.add_run(t)
        elif srcsty and srcsty.startswith("Heading 4"):
            np = doc.add_paragraph(); np.style = doc.styles["Heading 3"]; np.add_run(t)
        elif srcsty and srcsty.startswith("Heading"):
            np = doc.add_paragraph(); np.style = doc.styles[remap(srcsty)]; np.add_run(t)
        elif has_num and ref_bullet_el is not None:
            add_master_bullet(t)
        else:
            np = doc.add_paragraph(); np.style = doc.styles["normal"]
            run = np.add_run(t)
            if s == "Script:":
                run.bold = True
        n += 1
    print(f"  ported {n} Section 5 paragraphs")

    doc.save(OUT)
    print("SAVED", OUT)

if __name__ == "__main__":
    main()
