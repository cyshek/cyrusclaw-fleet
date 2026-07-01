#!/usr/bin/env python3
"""
Build the Inworld AI Founding SE guide on the NEW corrected template:
  - Copy templates/Master_Interview_Prep_Guide.docx -> Inworld guide (backup current first)
  - Fill Company block (Inworld AI) and Role block (Founding AI Solutions Engineer)
  - Append "Round 1 Game Plan — Florin Radu, COO" section tuned to the Jul 1 HM screen,
    in bold-topic + exact-spoken-line format (LOCKED style), then cheat-sheet / stories /
    questions / mindset sections reformatted from build_src/inworld_section6.txt.
  - Rebuild guides/inworld-ai-founding-se.zip (guide + JD + resume).
Mirrors build_src/build_google.py structure exactly. Edits the real template copy in place
so headings stay BLACK and the theme is preserved.
"""
import shutil, zipfile, datetime
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATE = WS / "templates/Master_Interview_Prep_Guide.docx"
GDIR = WS / "guides/inworld-ai-founding-se"
GUIDE = GDIR / "Inworld_AI_Founding_SE_Interview_Prep_Guide.docx"
JD = GDIR / "Inworld_Founding_SE_JD.md"
RESUME = GDIR / "Cyrus_Shekari_Resume_ashby-inworld-ai_9aef36c8_v2.pdf"
ZIP = WS / "guides/inworld-ai-founding-se.zip"

# backup current guide
bk = Path("/tmp") / f"inworld_guide_backup_{datetime.datetime.now():%Y%m%d_%H%M%S}.docx"
shutil.copy2(GUIDE, bk)
print("backup ->", bk)

shutil.copy2(TEMPLATE, GUIDE)
doc = Document(str(GUIDE))


def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for r in p._p.findall(qn('w:r')):
        p._p.remove(r)


def fill_placeholder(marker, label, rest):
    for p in doc.paragraphs:
        if marker in p.text:
            clear_runs(p)
            r1 = p.add_run(label)
            r1.bold = True
            p.add_run(rest)
            return True
    return False


COMPANY_REST = ("Inworld AI is a research lab building the world's #1-ranked realtime voice models — "
  "TTS, STT, an LLM Router, and a single-session Realtime API — and the compute to run them, powering "
  "hundreds of millions of users across partners like NVIDIA, Xbox, and Niantic, on $125M-plus raised. "
  "Why I want in: it's a frontier voice-AI company at the exact 0-to-1 build phase where I add the most "
  "value, and I want to be the technical backbone of a revenue team at that seam.")
ROLE_REST = ("The Founding AI Solutions Engineer is the technical backbone of the revenue team at the "
  "intersection of sales, product, and engineering — running technical discovery, POCs and prototypes, "
  "owning the technical side of deals before and after signature, and getting customers all the way to "
  "production. Why it fits me: I live at the engineering / product / stakeholder seam, I've run "
  "discovery, onboarding, and adoption with internal 'customers,' and I build working tools, like the AI "
  "agent I shipped on GitHub Copilot. I'll be honest that it's a stretch on pure SE years — I'm leading "
  "with the intersection, the building, and the ambiguity fit.")

assert fill_placeholder("[COMPANY", "Inworld AI — what + why: ", COMPANY_REST), "company ph missing"
assert fill_placeholder("[ROLE", "Founding AI Solutions Engineer — what + why: ", ROLE_REST), "role ph missing"

# ── Build the Round 1 Game Plan section at the end (before sectPr) ──────────────
new_body = doc.element.body
sectPr = new_body.find(qn('w:sectPr'))

# get a List Bullet style name if any (mirrors build_google.py: falls back to Normal)
bullet_style = None
try:
    _ = doc.styles["List Bullet"]
    bullet_style = "List Bullet"
except KeyError:
    bullet_style = None


def add_par(style=None):
    p_el = OxmlElement('w:p')
    new_body.insert(list(new_body).index(sectPr), p_el)
    p = Paragraph(p_el, doc)
    if style:
        p.style = doc.styles[style]
    else:
        p.style = doc.styles["Normal"]
    return p


def heading(text, level):
    p = add_par(style=f"Heading {level}")
    p.add_run(text)
    return p


def plain(text):
    p = add_par(style="Normal")
    p.add_run(text)
    return p


def topic_line(topic, said):
    """bold 'Topic: ' + plain spoken line"""
    p = add_par(style="Normal")
    r = p.add_run(topic + ": ")
    r.bold = True
    p.add_run(said)
    return p


def bullet(text):
    p = add_par(style=(bullet_style or "Normal"))
    p.add_run(text)
    return p


def note(text):
    """plain bracketed guidance paragraph (for Cyrus, not a spoken line)"""
    p = add_par(style="Normal")
    p.add_run(text)
    return p


# ---- Section content ----
heading("Round 1 Game Plan — Florin Radu, COO (Wed Jul 1, 2:00 PM PDT, Google Meet)", 2)
plain("30-minute Round 1 hiring-manager screen with Florin Radu, COO of Inworld AI, on Google Meet (meet.google.com/vsa-ztai-fvi); recruiter is Sadia Fatima. This is run by a founder-level operator: expect a fast, signal-dense conversation about fit, technical credibility, and whether you think like an owner. This is the DEEP one — the seat asks for 5-plus years customer-facing and strong Python, and you're coming from a TPM title. Be truthful about the stretch and sell the intersection: the eng/product/stakeholder seam, the hands-on building, and the 0-to-1 ambiguity fit. Keep answers crisp (30-45s) and let Florin drive.")

heading("Say-it-ready answers (bold topic, then the exact line)", 3)
topic_line("Walk me through your background",
    "\"I'm a Technical Program Manager at Microsoft, but really I live at the seam of engineering, product, "
    "and stakeholders — which is what a Solutions Engineer does. I ran a 0-to-1 cloud-resilience program "
    "to a 94% autonomous recovery rate, did discovery and onboarding with 20-plus internal service teams "
    "as my customers, and I build the tooling myself — including an AI agent on GitHub Copilot that runs "
    "our whole drill lifecycle. I translate deep technical tradeoffs for engineers and executives every "
    "day, and I want to do exactly that in a revenue seat at Inworld.\"")
topic_line("Why Inworld / why this role",
    "\"Inworld is a frontier voice-AI company with the number-one-ranked realtime models, and the Founding "
    "Solutions Engineer sits at the intersection of sales, product, and engineering — which is the exact "
    "seam I already work at. You're at the 0-to-1 build phase where the SE playbook is still being "
    "written, and that high-ambiguity build phase is where I add the most value. I want to be the "
    "technical backbone of the revenue team while that's still being built.\"")
topic_line("The honest SE-stretch reframe",
    "\"I'll be straight with you — my title is TPM, not SE, so on paper it's a stretch on customer-facing "
    "years. But the SE motion is what I actually do: discovery, POC, onboarding, production, expansion — "
    "I just run it with internal Azure service teams as my customers instead of external accounts. I'd "
    "rather be honest about the gap and show you the intersection, the building, and the ramp speed than "
    "oversell years I don't have. Truthful and hungry is how I want to win this.\"")
topic_line("Hands-on coding / Python",
    "\"Let me be honest about my level: I'm not a full-time production engineer, but I build working tools "
    "that solve real problems. My strongest proof is the AI agent I built with GitHub Copilot to run our "
    "drill lifecycle end-to-end — I wired together the APIs and telemetry and shipped something the team "
    "actually uses. I'm comfortable in Python, I ramp fast on a new stack, and I'd rather show you "
    "something I built than claim a depth I don't have.\"")
topic_line("Location / relocation",
    "\"I'm based in the Seattle area right now, and I know this role is Bay Area onsite a few days a week. "
    "I'm genuinely excited enough about Inworld that I'm open to relocating for the right role — I'd just "
    "want to talk through timing and the specifics with you.\"")
note("[Cyrus: confirm your true relocation stance before the call. You are in Kirkland WA; this role is "
    "SF Bay Area / South Bay onsite a few days a week, so relocation is a real question here. If you are "
    "genuinely open to relocating, the line above works as-is. If you would only do hybrid/remote or need "
    "a specific timeline or relocation support, swap in your real stance — do NOT imply you are already "
    "local to the Bay Area, because you are not.]")
topic_line("Work authorization",
    "\"I'm a US citizen — no sponsorship needed, now or in the future.\"")
topic_line("Why leaving Microsoft",
    "\"I've had a great run, but the resilience program matured from a 0-to-1 build into steady-state, and "
    "I add the most value in the high-ambiguity building phase — which is exactly where Inworld is right "
    "now. Nothing negative; I'm running toward this, not away.\"")

heading("Inworld product & company cheat sheet (know cold)", 3)
bullet("Six products: Realtime TTS (text-to-speech), Realtime STT (speech-to-text with voice profiling), a Realtime Router (one OpenAI-compatible endpoint routing across 200+ LLMs), Realtime Inference (their own optimized open-source models), a Realtime API (STT + LLM + TTS in one streaming WebSocket session), and managed Compute.")
bullet("Why they win: their TTS/voice models are #1-ranked on public realtime arenas — the pitch is quality + latency + the full pipeline + developer experience. Voice that feels human enough that users stay on the call and come back.")
bullet("Who buys: consumer-facing AI apps — companions, character chat, customer-support voice agents, sales/SDR and phone agents, language learning, interactive media. Top verticals: gaming, CCaaS (contact-center-as-a-service), and media/entertainment.")
bullet("Traction & backing: hundreds of millions of end users on apps powered by their tech; customers/partners have included NVIDIA, Microsoft Xbox, Niantic, Logitech Streamlabs, Wishroll, Bible Chat. Raised $125M-plus (Lightspeed, Section 32, Kleiner Perkins, M12, Founders Fund, Meta, Stanford). CB Insights AI 100; LinkedIn Top 10 US startups.")
bullet("Technical vocabulary to speak credibly: latency tradeoffs, streaming architecture (WebSocket/WebRTC), on-prem vs. cloud vs. edge deployment, model selection/routing, quantization and model serving, OpenAI-compatible APIs. You don't need to be an expert in all of it — speak it credibly and show you learn fast.")

heading("Your stories -> what Florin screens for", 3)
bullet("\"Can you run technical discovery + POCs?\" -> The GDOT / maintenance-data integration and the 20-plus service-team discovery: you ran discovery, found the real blocker (adoption, not tech), and drove it to a working integration.")
bullet("\"Can you go deep technically without backup?\" -> The Service Healing data-loss pivot: you stress-tested an architectural assumption with availability engineers, caught a real customer-data-loss risk, and re-architected mid-flight. You can hold your own in a deep technical room.")
bullet("\"Can you bridge technical and business / talk to executives?\" -> The sovereign-cloud network-isolation test tied to a $1.5B-plus contract — you were the bridge lead translating tradeoffs for engineers AND senior stakeholders under exec visibility.")
bullet("\"Customer empathy / handling a tough customer\" -> The escalated tier-one team-leader story: de-escalated, unblocked them in 24 hours, then fixed the root cause with self-service. Critic to advocate — a clean SE customer-success narrative.")
bullet("\"Can you build, not just coordinate?\" -> The AI agent (GitHub Copilot, full drill lifecycle) and the self-service automation platform. This is your proof of hands-on building — lean on it whenever coding or prototyping comes up.")
bullet("\"Why leave Microsoft?\" -> Program matured from 0-to-1 into steady-state; you add the most value in high-ambiguity build phases. Tie it directly to Inworld being at exactly that phase.")

heading("Questions to ask Florin (COO lens)", 3)
bullet("\"As the founding Solutions Engineer, what does 'great' look like in the first 6 months — closing technical wins, building the demo/integration library, or standing up the SE process itself?\"")
bullet("\"Where is most of the technical friction in deals today — latency/performance proof, integration complexity, deployment model (cloud vs. on-prem), or something else?\"")
bullet("\"Which verticals are you leaning into hardest right now — gaming, CCaaS, media — and where do you see the SE function mattering most as you scale?\"")
bullet("\"As COO, how do you see the GTM and SE motion evolving over the next 12-18 months, and how much of that playbook is still being written?\"")
bullet("\"What's the split between pre-sales POC work and post-sale production onboarding for this role today, and do you expect that to shift?\"")

heading("Mindset", 3)
bullet("Match the founder energy: Florin is a COO at a fast startup — be crisp, high-signal, and bias toward 'here's what I'd do.' No corporate meandering.")
bullet("Show genuine product curiosity: you actually looked at the Realtime API — say so. 'The single-WebSocket STT+LLM+TTS session is a clean DX story' reads as someone who'd sell it well.")
bullet("Be honest about the stretch, confident about the trajectory: don't oversell SE years you don't have. Sell the intersection, the building, the ambiguity fit, and how fast you ramp. Truthful + hungry beats inflated.")
bullet("Close with intent: it's a 30-minute screen — end by signaling real interest and asking about next steps in the process.")

doc.save(str(GUIDE))
print("Inworld guide built. paras:", len(Document(str(GUIDE)).paragraphs))

# Rebuild zip (guide + JD + resume)
# Rebuild zip (guide + JD + resume)
with zipfile.ZipFile(str(ZIP), "w", zipfile.ZIP_DEFLATED) as z:
    z.write(GUIDE, GUIDE.name)
    if JD.exists():
        z.write(JD, JD.name)
    if RESUME.exists():
        z.write(RESUME, RESUME.name)
with zipfile.ZipFile(str(ZIP)) as z:
    print("ZIP:", z.namelist())
