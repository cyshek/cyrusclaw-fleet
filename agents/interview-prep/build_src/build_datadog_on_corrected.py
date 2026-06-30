#!/usr/bin/env python3
"""
Build Datadog bundle on top of the CORRECTED (in-place-styled) template.
 1. Copy corrected template -> Datadog guide.
 2. Fill Company [COMPANY...] and Role [ROLE...] placeholders with Datadog spoken blocks
    (bold label run + plain rest), removing the bracket markers.
 3. Append the Round-1 Pulkit section (Heading1 + everything after it) grafted from the
    PREVIOUS datadog build's XML so the cheat sheet formatting/bold runs are preserved.
 4. Rebuild the prep zip (docx + JD + resume).
"""
import shutil, zipfile, copy, glob
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATE = WS / "templates/Master_Interview_Prep_Guide.docx"
OUTDIR = WS / "bundles/datadog-partner-tse"
GUIDE = OUTDIR / "Datadog_Partner_TSE_Interview_Prep_Guide.docx"
PREV = OUTDIR / "_prev_datadog_blue.docx"   # backup of the old build we graft Round-1 from
ZIP = OUTDIR / "Datadog_Partner_TSE_PrepBundle.zip"

# backup current (blue) datadog build to graft from, then build fresh from corrected template
shutil.copy2(GUIDE, PREV)
shutil.copy2(TEMPLATE, GUIDE)

doc = Document(str(GUIDE))

COMPANY_LABEL = "Datadog — what + why: "
COMPANY_REST = ("Datadog is the leading cloud observability and monitoring platform — metrics, logs, "
  "traces, and security all in one place, used by engineering and DevOps teams worldwide to understand "
  "what's happening across their infrastructure and applications. Why I want in: Datadog won the "
  "observability consolidation war, and the integration ecosystem is the growth engine that keeps "
  "extending that platform — I want to be at that edge, where new partners plug the rest of the world into Datadog.")

ROLE_LABEL = "Partner TSE — what + why: "
ROLE_REST = ("The Partner TSE is the primary technical contact for third-party developers building "
  "integrations on the Integration Developer Platform (IDP) — I'd guide them on architecture, review "
  "their code against the Quality Rubric, troubleshoot the hard issues, and feed the friction they hit "
  "back to product and engineering. Why it's right for me: it's the exact intersection of technical depth "
  "and external impact — consultant, engineer, and product voice at once — which is where I do my best work.")

def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for r in p._p.findall(qn('w:r')):
        p._p.remove(r)

def fill_placeholder(marker, label, rest):
    for p in doc.paragraphs:
        if marker in p.text:
            clear_runs(p)
            r1 = p.add_run(label); r1.bold = True
            p.add_run(rest)
            return True
    return False

assert fill_placeholder("[COMPANY", COMPANY_LABEL, COMPANY_REST), "company placeholder not found"
assert fill_placeholder("[ROLE", ROLE_LABEL, ROLE_REST), "role placeholder not found"

# ── Graft Round-1 section from PREV build ─────────────────────────────────────
prev = Document(str(PREV))
prev_body = prev.element.body
# find the <w:p> that is the 'Round 1 Prep — Pulkit' Heading 1 in prev; take it + all following block elements
round1_start = None
for el in list(prev_body):
    if el.tag == qn('w:p'):
        # get text
        texts = el.findall('.//' + qn('w:t'))
        txt = "".join(t.text or "" for t in texts)
        if txt.strip().startswith("Round 1 Prep"):
            round1_start = el
            break

assert round1_start is not None, "Round 1 section not found in previous build"

# collect round1_start and all following siblings (stop before sectPr)
graft = []
collecting = False
for el in list(prev_body):
    if el is round1_start:
        collecting = True
    if collecting:
        if el.tag == qn('w:sectPr'):
            break
        graft.append(el)

# append deep-copies to the new doc body, BEFORE its sectPr
new_body = doc.element.body
sectPr = new_body.find(qn('w:sectPr'))
for el in graft:
    new_body.insert(list(new_body).index(sectPr), copy.deepcopy(el))

doc.save(str(GUIDE))
print("Datadog guide built on corrected template. paras:", len(Document(str(GUIDE)).paragraphs))

# ── Rebuild zip (preserve JD + resume already referenced) ─────────────────────
# Find JD + resume: reuse whatever the current zip had.
jd = None; resume = None
if ZIP.exists():
    with zipfile.ZipFile(ZIP) as z:
        names = z.namelist()
        for n in names:
            if n.lower().endswith(".md"):
                (OUTDIR / "_jd_tmp.md").write_bytes(z.read(n)); jd = (n, OUTDIR / "_jd_tmp.md")
            if n.lower().endswith(".pdf"):
                (OUTDIR / "_resume_tmp.pdf").write_bytes(z.read(n)); resume = (n, OUTDIR / "_resume_tmp.pdf")

with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as z:
    z.write(GUIDE, GUIDE.name)
    if jd: z.write(jd[1], jd[0])
    if resume: z.write(resume[1], resume[0])

for t in ["_jd_tmp.md", "_resume_tmp.pdf"]:
    q = OUTDIR / t
    if q.exists(): q.unlink()

with zipfile.ZipFile(ZIP) as z:
    print("ZIP namelist:", z.namelist())
