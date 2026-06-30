#!/usr/bin/env python3
"""
Rewrite the Datadog 'Theme-by-theme prep' block into: bold topic label + exact spoken line.
In place on the delivered guide; preserves styles. Then rebuild the zip.
"""
import shutil, zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
OUTDIR = WS / "bundles/datadog-partner-tse"
GUIDE = OUTDIR / "Datadog_Partner_TSE_Interview_Prep_Guide.docx"
ZIP = OUTDIR / "Datadog_Partner_TSE_PrepBundle.zip"

doc = Document(str(GUIDE))

# topic -> exact spoken line (plain). Bold label = topic + ": "
themes = [
    ("Why Datadog",
     "\"I've been driving large-scale technical programs at Microsoft, but I'm most energized at the "
     "intersection of technical depth and external impact. The Partner TSE role is exactly that — and "
     "Datadog is the platform that won the observability consolidation war, where the integration "
     "ecosystem is how you keep extending it. That's the edge I want to be on.\""),
    ("Owning a partner relationship end-to-end",
     "\"At Microsoft I owned the GDOT platform adoption. The local data center teams were my external "
     "developers and the central platform team was my internal product partner. They weren't adopting "
     "the tool, so I diagnosed why, packaged their feedback into a product brief, drove the fix, and "
     "got them to 100% adoption — that's the exact TSE motion.\""),
    ("Observability stack knowledge",
     "\"I know how the pieces connect: metrics are numbers over time, logs are timestamped events, and "
     "traces follow one request across services. OpenTelemetry is the open standard Datadog supports "
     "natively, and integrations are either agent-based — running on the host in near real-time — or "
     "API-based, polling a remote endpoint. I don't need to write the code; I need to advise on the "
     "architecture.\""),
    ("Code quality / integration advising",
     "\"If a partner's integration works but doesn't meet the Quality Rubric, I start with curiosity — "
     "what was their constraint? — then explain the standard and why it protects the ecosystem, and "
     "collaborate on a path to compliance instead of just blocking them.\""),
    ("Pushing back / influencing product",
     "\"On the GDOT project I had no authority over the product team or the local operators, but I "
     "packaged external feedback into a product brief, got the feature shipped, and drove adoption from "
     "the ground up. That's the identify-friction, feed-it-back-to-engineering loop the TSE role runs on.\""),
]

def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for r in p._p.findall(qn('w:r')):
        p._p.remove(r)

# Locate the 5 theme paragraphs between 'Theme-by-theme prep' and 'Datadog observability cheat sheet'.
paras = doc.paragraphs
start = end = None
for i, p in enumerate(paras):
    if p.text.strip() == "Theme-by-theme prep":
        start = i
    if start is not None and p.text.strip() == "Datadog observability cheat sheet":
        end = i
        break

assert start is not None and end is not None, "theme block bounds not found"

# The theme content paragraphs are the non-empty normal paras between start+1 and end.
content_idxs = [i for i in range(start+1, end) if paras[i].text.strip()]
assert len(content_idxs) == len(themes), f"expected {len(themes)} theme paras, found {len(content_idxs)}"

for idx, (topic, line) in zip(content_idxs, themes):
    p = paras[idx]
    clear_runs(p)
    r1 = p.add_run(topic + ": ")
    r1.bold = True
    p.add_run(line)

doc.save(str(GUIDE))
print("Theme prep rewritten. Verifying...")

# verify each theme para: runs[0] bold, ends with colon
d2 = Document(str(GUIDE))
ok = 0
inblock = False
for p in d2.paragraphs:
    t = p.text.strip()
    if t == "Theme-by-theme prep":
        inblock = True; continue
    if t == "Datadog observability cheat sheet":
        break
    if inblock and t:
        b = p.runs[0].bold if p.runs else None
        lead = p.runs[0].text if p.runs else ""
        good = bool(b) and lead.strip().endswith(":")
        print(f"  {(chr(79)+chr(75)+chr(32)) if good else (chr(66)+chr(65)+chr(68))} bold={b} | {lead.strip()[:42]}")
        if good: ok += 1
print(f"themes OK: {ok}/5")

jd = resume = None
with zipfile.ZipFile(ZIP) as z:
    names = z.namelist()
    for n in names:
        if n.lower().endswith(".md"):
            (OUTDIR / "_jd.md").write_bytes(z.read(n)); jd = (n, OUTDIR / "_jd.md")
        if n.lower().endswith(".pdf"):
            (OUTDIR / "_res.pdf").write_bytes(z.read(n)); resume = (n, OUTDIR / "_res.pdf")
with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as z:
    z.write(GUIDE, GUIDE.name)
    if jd: z.write(jd[1], jd[0])
    if resume: z.write(resume[1], resume[0])
for t in ["_jd.md", "_res.pdf"]:
    q = OUTDIR / t
    if q.exists(): q.unlink()
with zipfile.ZipFile(ZIP) as z:
    print("ZIP:", z.namelist())
