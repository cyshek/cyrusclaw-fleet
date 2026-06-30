#!/usr/bin/env python3
"""
Remove the three italic helper lines from BOTH the canonical template and the Datadog guide:
  - "Say it as ONE spoken answer — what the company is, ..."
  - "Say it as ONE spoken answer — what the role is, ..."
  - "Rehearse from the bullets — each card is a fact skeleton, ..."
Then rebuild the Datadog zip. Preserves everything else / styles.
"""
import zipfile
from pathlib import Path
from docx import Document

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATE = WS / "templates/Master_Interview_Prep_Guide.docx"
OUTDIR = WS / "bundles/datadog-partner-tse"
GUIDE = OUTDIR / "Datadog_Partner_TSE_Interview_Prep_Guide.docx"
ZIP = OUTDIR / "Datadog_Partner_TSE_PrepBundle.zip"

PREFIXES = (
    "Say it as ONE spoken answer",
    "Rehearse from the bullets",
)

def strip_helpers(path):
    doc = Document(str(path))
    removed = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if any(t.startswith(pref) for pref in PREFIXES):
            p._p.getparent().remove(p._p)
            removed += 1
    doc.save(str(path))
    return removed

r_tpl = strip_helpers(TEMPLATE)
r_dd = strip_helpers(GUIDE)
print(f"template removed: {r_tpl}")
print(f"datadog removed:  {r_dd}")

# verify none remain
for label, path in (("template", TEMPLATE), ("datadog", GUIDE)):
    d = Document(str(path))
    txt = "\n".join(p.text for p in d.paragraphs)
    still = [pref for pref in PREFIXES if pref in txt]
    print(f"  {label}: remaining helper prefixes -> {still if still else 'NONE'}")

# rebuild zip
jd = resume = None
with zipfile.ZipFile(ZIP) as z:
    names = z.namelist()
    for n in names:
        if n.lower().endswith(".md"):
            (OUTDIR / "_jd.md").write_bytes(z.read(n)); jd = (n, OUTDIR / "_jd.md")
        if n.lower().endswith(".pdf"):
            (OUTDIR / "_res.pdf").write_bytes(z.read(n)); resume = (n, OUTDIR / "_res.pdf")
with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as z:
    z.write(GUIDE, GUIDE.name)
    if jd: z.write(jd[1], jd[0])
    if resume: z.write(resume[1], resume[0])
for t in ("_jd.md", "_res.pdf"):
    q = OUTDIR / t
    if q.exists(): q.unlink()
with zipfile.ZipFile(ZIP) as z:
    print("ZIP:", z.namelist())
