"""
TASK 1 — Build the new canonical Master Interview Prep Guide template.

Overwrites templates/Master_Interview_Prep_Guide.docx with the new structure:
  - Title + how-to-use line
  - Section 1: Core Questions (Q1 kept verbatim, Q2 Company placeholder,
    Q3 Role placeholder, Q4 kept verbatim)
  - Section 2: Behavioral Story Bank (3 study cards from structural_bullets.txt)
  - Section 3: Managing Senior/Executive Requests (verbatim, bold step labels)
  - Section 4: Product Thinking Questions (verbatim, bold component labels)
  - Section 5: Questions for the Interviewer (verbatim)

Template KEEPS the [Fill ...] placeholders. Base font Calibri 11.
Every label is its OWN bold run so p.runs[0].bold == True.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
OUT = WORKSPACE / "templates" / "Master_Interview_Prep_Guide.docx"


# ---------------------------------------------------------------------------
# Low-level paragraph / run helpers
# ---------------------------------------------------------------------------

def set_base_font(document):
    """Set the document default font to Calibri 11."""
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_heading(document, text, level):
    p = document.add_heading(text, level=level)
    return p


def add_plain(document, text, style="Normal"):
    """Add a paragraph with a single plain (non-bold) run."""
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = False
    return p


def add_italic(document, text, style="Normal"):
    """Add a paragraph with a single italic run (used for helper lines)."""
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    run.italic = True
    return p


def add_bold_only(document, text, style="Normal"):
    """Add a paragraph whose only run is bold (e.g. a lone 'Script:' line)."""
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = True
    return p


def add_label_then_plain(document, label, rest, style="Normal"):
    """
    Paragraph with TWO runs:
      run[0] bold = label (ends with ':')
      run[1] plain = ' ' + rest
    """
    p = document.add_paragraph(style=style)
    r0 = p.add_run(label)
    r0.bold = True
    r1 = p.add_run(" " + rest)
    r1.bold = False
    return p


def add_bullet_plain(document, text):
    """A plain List Bullet paragraph (a single fact)."""
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.bold = False
    return p


def add_bullet_label_only(document, label):
    """A List Bullet paragraph whose ONLY run is a bold phase label."""
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(label)
    run.bold = True
    return p


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

doc = Document()
set_base_font(doc)

# ===== Title + how-to-use =====
add_heading(doc, "Master Interview Prep Guide", level=1)
add_italic(
    doc,
    "How to use this guide: Section 1 = spoken openers — speak these, don't "
    "recite them. Section 2 = your story bank as study cards — rehearse from "
    "the bullets, hit the 4 beats in order, then stop. The role cheat sheet at "
    "the end (added per role) is for drilling the morning of — say each term "
    "out loud.",
)

# ===== SECTION 1: Core Questions =====
add_heading(doc, "Section 1: Core Questions", level=2)

# --- Q1 (kept verbatim) ---
add_heading(doc, "Q1: Tell me about yourself / Walk me through your background", level=3)
add_bold_only(doc, "Script:")
add_plain(
    doc,
    "I\u2019m currently a Technical Program Manager at Microsoft, where I focus "
    "on cloud reliability, recovery validation, and platform automation. My core "
    "responsibility is ensuring our large-scale infrastructure can gracefully "
    "survive worst-case system failures.",
)
add_plain(
    doc,
    "Over the past couple of years, I\u2019ve scaled our recovery program into a "
    "highly automated, platformized system, leading over 14 high-stakes, "
    "cross-functional resilience executions under executive visibility. This "
    "includes leading critical incident bridges for key enterprise contracts and "
    "driving architectural optimizations that cut system failover latencies from "
    "over 20 minutes down to just two.",
)
add_plain(
    doc,
    "One of my biggest milestones was pioneering a 0-to-1 proactive testing "
    "capability that achieved a 94% autonomous recovery rate and successfully "
    "surfaced latent hardware defects before they hit production. To support this "
    "scale, I also led the requirements for an internal automation platform that "
    "cut manual planning cycles by 30%, and built custom AI agents to turn "
    "complex recovery playbooks into self-service engineering tools.",
)
add_plain(
    doc,
    "Before this, I earned my Computer Science degree and completed internships "
    "at Amazon Robotics and twice here at Microsoft, always focusing on "
    "data-driven optimization\u2014like mapping migration dependencies across "
    "2,000 robotics units or designing frameworks to deploy infrastructure 28% "
    "faster.",
)
add_plain(
    doc,
    "I\u2019ve loved solving these massive-scale challenges, but I\u2019m ready "
    "for my next step where I can bring this exact intersection of distributed "
    "systems reliability, cross-functional leadership, and automation to your "
    "team.",
)

# --- Q2: The Company (consolidated placeholder) ---
add_heading(doc, "Q2: The Company \u2014 what it is and why you want to be there", level=3)
add_italic(
    doc,
    "Say it as ONE spoken answer (what the company is, then why you want in). "
    "Filled per role.",
)
add_label_then_plain(
    doc,
    "[COMPANY \u2014 what + why]:",
    "[Fill per role: one-breath answer \u2014 what the company is/does, then the "
    "specific reason you want to be there. Show basic homework; keep it natural, "
    "not a deep-dive recitation.]",
)

# --- Q3: The Role (consolidated placeholder) ---
add_heading(doc, "Q3: The Role \u2014 what it is and why you want it", level=3)
add_italic(
    doc,
    "Say it as ONE spoken answer (what the role is, then why it's the right next "
    "step). Filled per role.",
)
add_label_then_plain(
    doc,
    "[ROLE \u2014 what + why]:",
    "[Fill per role: one-breath answer \u2014 what the role actually does "
    "day-to-day, then why it's the intersection where you do your best work.]",
)

# --- Q4 (kept verbatim) ---
add_heading(doc, "Q4: Why are you looking to leave Microsoft?", level=3)
add_bold_only(doc, "Script:")
add_plain(
    doc,
    "I\u2019ve had a fantastic experience at Microsoft, and I\u2019m incredibly "
    "proud of the impact I\u2019ve been able to deliver\u2014scaling our recovery "
    "program into an automated platform, launching our first proactive testing "
    "capabilities, and driving significant architecture and automation wins.",
)
add_plain(
    doc,
    "The main reason I\u2019m looking for my next step is that my current program "
    "has reached a very high level of operational maturity. It has naturally "
    "transitioned from that ambiguous, 0-to-1 engineering phase where we were "
    "designing systems from scratch into a steady-state, operational phase "
    "focused on maintaining what we built.",
)
add_plain(
    doc,
    "Through this, I\u2019ve realized that I add the absolute most value\u2014and "
    "am most energized\u2014when I\u2019m operating in that sweet spot of high "
    "ambiguity. I love being handed a complex, undefined platform problem, "
    "building the cross-functional engineering alignment required to solve it, "
    "and driving it from zero to execution. I\u2019m looking to bring that exact "
    "experience in systems reliability, platform automation, and execution "
    "leadership to a team here that is tackling its next major building phase.",
)

# ===== SECTION 2: Behavioral Story Bank (Study Cards) =====
add_heading(doc, "Section 2: Behavioral Story Bank (Study Cards)", level=2)
add_italic(
    doc,
    "Rehearse from the bullets \u2014 each card is a fact skeleton, not a script. "
    "Hit the labeled beats in order and stop. Map any behavioral question to the "
    "closest card.",
)

# --- Card 1 ---
add_heading(
    doc,
    "Card 1 \u2014 Proactive Resilience Testing Program (0\u21921, disagreement & pivot)",
    level=3,
)
add_italic(
    doc,
    "Maps to: \u201CTell me about an ambiguous problem\u201D (Q1), \u201Ca "
    "disagreement with a stakeholder\u201D (Q2), and troubleshooting a complex "
    "technical issue.",
)
add_bullet_label_only(doc, "Situation:")
add_bullet_plain(doc, "Ensured cloud infrastructure resilience within the team at Microsoft.")
add_bullet_plain(
    doc,
    "Historically, the recovery program was entirely reactive, validating "
    "capabilities only when Service Teams requested a test.",
)
add_bullet_plain(
    doc,
    "Drills were historically larger blast radius, increasing the likelihood of "
    "impact to other services within the cluster that did not want to participate.",
)
add_bullet_plain(
    doc,
    "No existing framework was in place for proactively validating resilience at "
    "scale without impacting service teams.",
)
add_bullet_plain(
    doc,
    "Owned the 0-to-1 incubation of Azure's first-ever proactive resilience "
    "testing program with a broad mandate to convert unused drill slots and scale "
    "the program by increasing the frequency of smaller blast radius drills.",
)
add_bullet_label_only(doc, "The Disagreement & Pivot:")
add_bullet_plain(
    doc,
    "The initial execution model relied on an existing script used for power "
    "cycling nodes, tweaked to target a smaller blast radius (rack instead of "
    "cluster).",
)
add_bullet_plain(
    doc,
    "The script was provided by engineering counterparts who develop programming "
    "for failure modes invoked for hardware power operations.",
)
add_bullet_plain(
    doc,
    "The Availability Platform team (owning the Azure Node and VM automated "
    "recovery process) and infrastructure engineering leads raised a key concern.",
)
add_bullet_plain(
    doc,
    "The concern was that Azure's safety mechanism marks nodes as Human "
    "Investigate when less than 20% of nodes in a cluster are powered off, as "
    "Fabric (the control plane triggering recovery) interprets this as unexpected "
    "behavior.",
)
add_bullet_plain(
    doc,
    "After data center team inspection, Service Healing (an infrastructure "
    "fallback mechanism where VMs are migrated to healthy hardware) takes place to "
    "keep services available.",
)
add_bullet_plain(
    doc,
    "The main issue was that ephemeral data on the nodes is wiped out during this "
    "process, making it impossible to track node recovery times and detect "
    "restart issues.",
)
add_bullet_label_only(doc, "Action Taken:")
add_bullet_plain(
    doc,
    "Pushed back the execution date for the first Proactive drill to allow for "
    "cross-functional alignment between drill engineering and infrastructure teams.",
)
add_bullet_plain(
    doc,
    "Decided to still use the script to power cycle nodes, but established a need "
    "to block Service Healing before the drill and unblock it afterward.",
)
add_bullet_plain(
    doc,
    "Partnered with Availability Platform team engineers to develop scripts to "
    "trigger this blocker.",
)
add_bullet_plain(
    doc,
    "To ensure full future autonomy, gathered engineering documents, scripts, "
    "APIs, and telemetry queries.",
)
add_bullet_plain(
    doc,
    "Leveraged GitHub Copilot to architect an AI agent capable of programmatically "
    "executing the entire drill lifecycle.",
)
add_bullet_plain(
    doc,
    "Successfully ran subsequent proactive drills end-to-end utilizing this AI "
    "solution.",
)
add_bullet_label_only(doc, "Result:")
add_bullet_plain(
    doc,
    "The very first end-to-end execution achieved a 94% autonomous recovery rate, "
    "safely recovering 15 out of 16 infrastructure nodes with zero data loss.",
)
add_bullet_plain(
    doc,
    "The drill successfully surfaced a latent hardware defect that could have "
    "impacted customers if left undetected.",
)
add_bullet_plain(
    doc,
    "The program ultimately gained high-level executive visibility, receiving "
    "strong praise from multiple CVP leadership members and setting a new standard "
    "for resilience automation across Microsoft.",
)

# --- Card 2 ---
add_heading(
    doc,
    "Card 2 \u2014 GDOT Platform Integration & User Adoption (influence w/o authority)",
    level=3,
)
add_italic(
    doc,
    "Maps to: \u201Ca project that didn't go as planned / how you pivoted\u201D "
    "(Q3), \u201Cinfluencing without authority\u201D (Q6), \u201Ca frustrated "
    "customer\u201D (Q7), and being the primary technical contact for users.",
)
add_bullet_label_only(doc, "Situation:")
add_bullet_plain(
    doc,
    "Faced a major bottleneck while scaling the cloud recovery program due to "
    "manual confirmation requirements with the Data Center Operations team "
    "regarding planned hardware maintenance.",
)
add_bullet_plain(
    doc,
    "Service teams lacked visibility into which slots were unavailable due to rack "
    "maintenance.",
)
add_bullet_plain(
    doc,
    "Goal was to automate this conflict-checking process by establishing an API "
    "integration with the Global Datacenter Operations Team (GDOT) platform, the "
    "central source of truth for maintenance data.",
)
add_bullet_plain(
    doc,
    "Integration aimed to enable the platform to automatically block unavailable "
    "dates on the internal calendar.",
)
add_bullet_label_only(doc, "The Roadblock:")
add_bullet_plain(
    doc,
    "The project hit a major wall upon discovering that local data center "
    "operators for the Canary region were not using the central GDOT platform at "
    "all.",
)
add_bullet_plain(
    doc,
    "Operators found the system unintuitive and difficult to navigate, leading "
    "them to track all campus maintenance data on localized, manual spreadsheets.",
)
add_bullet_plain(
    doc,
    "Because their data was completely siloed, the automated API strategy was "
    "initially stalled due to the lack of a reliable data source to query, "
    "combined with a lack of direct authority over local operations teams.",
)
add_bullet_label_only(doc, "Action Taken:")
add_bullet_plain(
    doc,
    "Recognized the issue as a human adoption challenge rather than purely an "
    "engineering problem, requiring a collaborative approach.",
)
add_bullet_plain(
    doc,
    "Gathered granular usability feedback from operators and packaged it into a "
    "data-driven product gap analysis for the central GDOT product team, "
    "successfully influencing their long-term UX roadmap.",
)
add_bullet_plain(
    doc,
    "Fed AI videos and documentation covering how to migrate data to the new "
    "platform to solve the immediate data silo.",
)
add_bullet_plain(
    doc,
    "Had AI develop automations to assist local campus teams in migrating their "
    "spreadsheet data, including maintenance schedules, scope of work, and "
    "impacted hardware.",
)
add_bullet_plain(
    doc,
    "Built dedicated onboarding resources, structured wiki guides, and configured "
    "a custom help copilot to support teams through the learning curve.",
)
add_bullet_plain(
    doc,
    "Demonstrated immediate value to the 11 operators by showing how centralized "
    "data would automatically prevent recovery drills from disrupting physical "
    "data center maintenance windows.",
)
add_bullet_label_only(doc, "Result:")
add_bullet_plain(
    doc,
    "Local teams successfully transitioned 100% of their maintenance data into the "
    "central platform, completely eliminating siloed spreadsheets in the canary "
    "region.",
)
add_bullet_plain(
    doc,
    "Unlocked the technical roadmap, allowing the establishment of a clean API "
    "connection that eliminated manual scheduling friction and back-and-forth "
    "coordination.",
)
add_bullet_plain(
    doc,
    "Demonstrated that true ownership means solving human adoption challenges to "
    "make technology effective.",
)

# --- Card 3 ---
add_heading(
    doc,
    "Card 3 \u2014 Resilience Automation Platform / Self-Service Intake (trade-offs & scale)",
    level=3,
)
add_italic(
    doc,
    "Maps to: \u201Cautomating a manual process to scale a team\u201D (Q4), \u201Ca "
    "difficult trade-off between competing priorities\u201D (Q5), and product "
    "strategy / prioritization questions.",
)
add_bullet_label_only(doc, "Situation:")
add_bullet_plain(
    doc,
    "The team was responsible for running Azure\u2019s high-impact resilience and "
    "recovery drills, but hit a massive scaling wall due to operational toil.",
)
add_bullet_plain(
    doc,
    "Every internal service team scheduling request required hours of manual "
    "intake, timeline coordination, scope definition, and failure mode mapping.",
)
add_bullet_plain(doc, "Manual friction consumed nearly 6 hours per drill for planning and scoping.")
add_bullet_plain(
    doc,
    "Faced a team KPI to scale up and sustain over 45 major resilience drills a "
    "year, which was unsustainable using manual processes.",
)
add_bullet_plain(doc, "Trying to go from running 1 drill a month to 1 drill a week.")
add_bullet_label_only(doc, "The Tension & Trade-off:")
add_bullet_plain(
    doc,
    "As the owner of the automation initiative to transition to a self-service "
    "platform model, I received numerous competing feature requests from internal "
    "service teams and engineering leads.",
)
add_bullet_plain(doc, "Operated under a strict deadline to launch the platform MVP.")
add_bullet_plain(
    doc,
    "Established a strict prioritization framework based on operational impact vs. "
    "execution velocity to protect project momentum.",
)
add_bullet_plain(
    doc,
    "Focused 100% of initial engineering resources on the self-service scheduling "
    "and intake flow to offload manual work.",
)
add_bullet_plain(
    doc,
    "Explicitly deprioritized secondary features, such as requested post-drill "
    "telemetry dashboards containing recovery metrics and RCAs, because that data "
    "could temporarily be pulled manually.",
)
add_bullet_plain(
    doc,
    "Partnered with the Observability team to identify the source of truth for "
    "Node-to-Service team mapping, allowing Service Teams to see what hardware "
    "resources their subscriptions were tied to.",
)
add_bullet_plain(
    doc,
    "Eliminated a tedious and inaccurate process where the team previously looked "
    "through various databases to construct ad-hoc queries for this information.",
)
add_bullet_label_only(doc, "Action Taken:")
add_bullet_plain(doc, "Partnered with Dev and PM leads to prioritize the MVP feature set effectively.")
add_bullet_plain(
    doc,
    "Leveraged AI agents to rapidly prototype key features needed, including a "
    "calendar view, drill request form, Service ID lookup resource mapping, "
    "compute infra failure modes, executive communications, RBAC, and remote "
    "execution.",
)
add_bullet_plain(
    doc,
    "Set up an A/B testing framework to evaluate different layout configurations "
    "and visual hierarchies.",
)
add_bullet_plain(
    doc,
    "Gathered real-world telemetry on user drop-off rates and completion times to "
    "ensure the finalized layout was frictionless for both technical and "
    "non-technical operational teams.",
)
add_bullet_label_only(doc, "Result:")
add_bullet_plain(doc, "Launched the self-service platform exactly on time by protecting the MVP scope.")
add_bullet_plain(
    doc,
    "The A/B-tested UI directly resulted in a 35% reduction in intake form "
    "completion times with a near-zero error rate during self-service scheduling.",
)
add_bullet_plain(
    doc,
    "Transformed the operating model by crushing a process that previously "
    "required six hours of manual back-and-forth coordination down to just six "
    "minutes.",
)
add_bullet_plain(
    doc,
    "Immediately reclaimed critical engineering hours, allowing the team to "
    "successfully scale and sustain the 45-drills-per-year milestone without "
    "adding headcount.",
)

# ===== SECTION 3: Managing Senior/Executive Requests (verbatim) =====
add_heading(doc, "Section 3: Managing Senior / Executive Requests", level=2)
add_plain(doc, "Use this structure for questions like:")
add_bullet_plain(
    doc,
    "How do you react when someone senior says 'this is high priority, get it "
    "done\u201D?",
)
add_bullet_plain(
    doc,
    "How would you react if you are alone at 6:30 p.m. and your VP walks in with "
    "work?",
)
add_bullet_plain(doc, "What if a VP asks for something urgently with no clear business impact?")
add_bullet_plain(
    doc,
    "How would you deal with a VP asking for something totally out of your area?",
)
add_bold_only(doc, "Script:")
add_plain(
    doc,
    "1. My approach to handling urgent requests from senior leadership is anchored "
    "in three principles: maintaining executive empathy, validating business "
    "impact, and protecting my core commitments. I never want to blindly say "
    "'yes' and risk delivering low-quality work, but I also never say 'no' without "
    "offering a constructive alternative.",
)
add_plain(
    doc,
    "2. If a [senior/executive/etc] approached me with an urgent, high-priority "
    "request\u2014whether it was after hours, entirely out of my domain, or "
    "lacking immediate context\u2014I would handle it using a deliberate "
    "three-step framework:",
)
add_label_then_plain(
    doc,
    "De-escalate and Diagnose:",
    "First, I would pause to establish context. I would ask targeted, clarifying "
    "questions to understand the underlying business driver, the hard deadline, "
    "and the intended audience. If it\u2019s out of my domain, I\u2019d ask what "
    "specific expertise they need from me. If it\u2019s 6:30 p.m., I\u2019d ask, "
    "'Is this blocking an executive review tomorrow morning, or can we establish a "
    "game plan first thing at 8:00 a.m.?'",
)
add_label_then_plain(
    doc,
    "Evaluate the Trade-offs:",
    "Next, I look at the request through the lens of program trade-offs. I look at "
    "my current high-priority commitments\u2014such as our active resilience drill "
    "schedules\u2014and determine what would have to drop or slide to accommodate "
    "this new request.",
)
add_label_then_plain(
    doc,
    "Propose a Scoped Solution:",
    "Finally, I present options rather than roadblocks. I will say, 'I can "
    "absolutely deliver X for you by tomorrow morning. To do that, I will pause "
    "work on Y, which will delay our milestone by 48 hours. Does that trade-off "
    "align with your expectations?' If I am completely out of my depth "
    "technically, I will immediately leverage my network to loop in the exact "
    "domain expert who can unblock them faster than I can.",
)
add_plain(
    doc,
    "3. I actually apply this framework regularly at Microsoft. For example, when "
    "multiple CVPs took notice of our proactive resilience testing capability, we "
    "frequently received high-priority, urgent inquiries regarding infrastructure "
    "risk profiles. By immediately diagnosing the exact data they needed, "
    "evaluating our team's engineering bandwidth, and setting clear boundaries on "
    "delivery timelines, I ensured we satisfied executive inquiries without "
    "derailing our core roadmap to hit our 45-drill annual milestone.",
)

# ===== SECTION 4: Product Thinking Questions (verbatim) =====
add_heading(doc, "Section 4: Product Thinking Questions", level=2)

add_heading(doc, "Q1: What is a product vision? Can you give an example?", level=3)
add_bold_only(doc, "Script:")
add_plain(
    doc,
    "A product vision is the long-term North Star for a product \u2014 a clear "
    "statement of what it ultimately exists to accomplish and why it matters. It "
    "defines the problem being solved, who it serves, and what success looks like "
    "at full scale. It's not a feature list or a roadmap; it's the destination "
    "everything else is navigating toward.",
)
add_plain(
    doc,
    "A good product vision does three things: it answers why the product needs to "
    "exist, it sets a direction ambitious enough to inspire the team, and it's "
    "specific enough to filter decisions. When I'm evaluating whether to build a "
    "feature or cut scope, I ask whether it moves us closer to the vision. If it "
    "doesn't, it waits.",
)
add_plain(
    doc,
    "A concrete example from my work at Microsoft: I owned the vision for an "
    "AI-driven autonomous drill agent. The vision was to make cloud infrastructure "
    "resilience fully self-sustaining \u2014 where unused capacity is automatically "
    "identified and converted into proactive recovery validation, without any "
    "human scheduling it. The problem it addressed was that customers would book "
    "resilience drill slots and then cancel last-minute, wasting planning cycles "
    "and leaving gaps in our validation coverage. The vision wasn't \"build an "
    "agent\" \u2014 that's implementation. The vision was that Azure "
    "infrastructure should continuously validate its own recovery health, "
    "proactively, at scale, without engineers manually orchestrating it. "
    "Everything we built was in service of that.",
)

add_heading(doc, "Q2: How do you build a product strategy?", level=3)
add_bold_only(doc, "Script:")
add_plain(
    doc,
    "Product strategy is the bridge between the vision and the roadmap \u2014 it's "
    "how you decide where to focus and in what order. When I build a product "
    "strategy, I anchor it in five components:",
)
add_label_then_plain(
    doc,
    "First, I start with customer and stakeholder input.",
    "Before anything else, I talk to the people the product serves. At Microsoft, "
    "that meant sitting with internal Azure service teams \u2014 understanding "
    "where the friction was, what was blocking them, and what outcomes they "
    "actually cared about. That pain became the foundation of the strategy.",
)
add_label_then_plain(
    doc,
    "Second, I map the gap between current state and the vision.",
    "I'm explicit about where we are today versus where we need to be. That gap "
    "analysis surfaces the highest-leverage problems to solve \u2014 the ones "
    "that, if removed, unlock everything downstream.",
)
add_label_then_plain(
    doc,
    "Third, I build a prioritized roadmap.",
    "Not everything can be first. I rank work by impact vs. effort and protect the "
    "critical path. When we built the Resilience Automation Platform, the first "
    "milestone was self-service intake \u2014 not dashboards or advanced "
    "analytics. Intake was the single unlock that cut ~6 hours of manual planning "
    "per drill. Everything else was downstream of that.",
)
add_label_then_plain(
    doc,
    "Fourth, I align stakeholders before writing a single spec.",
    "A strategy nobody believes in doesn't get resourced. I socialize early, take "
    "objections seriously, and adjust. When I needed the GDOT platform team to "
    "change their roadmap to support our API integration, I brought them operator "
    "feedback packaged as a product gap analysis \u2014 a mutual win, not a "
    "demand.",
)
add_label_then_plain(
    doc,
    "Fifth, I define the metrics upfront.",
    "Strategy without measurement is just a plan that never gets revisited. I set "
    "KPIs before we build: for the automation platform it was time-per-drill-"
    "intake, drill cadence per month, and engineering hours recovered. Those "
    "metrics became the scoreboard we used to justify every roadmap decision.",
)

add_heading(doc, "Q3: What metrics do you use to measure success? Can you give an example?", level=3)
add_bold_only(doc, "Script:")
add_plain(
    doc,
    "I anchor metrics to the problem the product was solving, not to activity. "
    "Outputs matter less than outcomes.",
)
add_plain(
    doc,
    "The clearest example from my work is what I call the \"6 hours to 6 "
    "minutes\" metric. Before we built the self-service intake platform, "
    "coordinating a single resilience drill \u2014 scoping the blast radius, "
    "confirming resources, scheduling around maintenance windows \u2014 took an "
    "average of six hours of back-and-forth. After the platform launched, that "
    "same process took six minutes. That single number told the whole story: the "
    "manual coordination bottleneck was gone.",
)
add_plain(doc, "Beyond intake, the metrics I tracked for the full program were:")
add_label_then_plain(
    doc,
    "94% autonomous recovery rate",
    "on the first proactive drill execution \u2014 15 of 16 nodes recovered "
    "without human intervention.",
)
add_label_then_plain(
    doc,
    "Drill execution time cut",
    "from a full business day down to approximately 2 hours through remote "
    "API-driven power-down, replacing physical data center walkthroughs.",
)
add_label_then_plain(
    doc,
    "45+ resilience drills per year",
    "sustained without adding headcount \u2014 the platform absorbed the scale.",
)
add_label_then_plain(
    doc,
    "Zero customer data loss",
    "across all proactive drills, validating the architectural safety decisions we "
    "made early.",
)
add_plain(
    doc,
    "The rule I follow: pick metrics that would change a decision if they moved. "
    "If the number going up or down doesn't affect what you build next, it's the "
    "wrong metric.",
)

# ===== SECTION 5: Questions for the Interviewer (verbatim) =====
add_heading(doc, "Section 5: Questions for the Interviewer", level=2)
add_bold_only(doc, "Script:")
add_bullet_plain(doc, "What does 'great' look like 6 months in?")
add_bullet_plain(
    doc,
    "How does the team balance long-term strategic bets with short-term customer "
    "requests?",
)
add_bullet_plain(
    doc,
    "How do you see the company's product priorities shifting over the next 12 to "
    "18 months?",
)
add_bullet_plain(
    doc,
    "What is the biggest operational bottleneck or friction point the team faces "
    "today that you are hoping the person in this role can help solve?",
)
add_bullet_plain(
    doc,
    "If you could hit the reset button and start over in your current role from "
    "scratch, knowing what you know now, what would you do differently?",
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT)
print("SAVED:", OUT)
print("Paragraph count:", len(doc.paragraphs))
