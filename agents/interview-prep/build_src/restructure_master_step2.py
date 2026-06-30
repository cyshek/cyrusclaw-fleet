#!/usr/bin/env python3
"""
STEP 2: Replace the 7 prose STAR Q-scripts in Section 2 with 3 bulleted STUDY CARDS
built from Cyrus's Structural Study Guide (build_src/structural_bullets.txt).
Preserves his styles. Operates on templates/Master_Interview_Prep_Guide.docx (already
has Step 1 applied).
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import copy

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
OUT = WS / "templates/Master_Interview_Prep_Guide.docx"
doc = Document(str(OUT))

def idx_of(pred, start=0):
    for i in range(start, len(doc.paragraphs)):
        if pred(doc.paragraphs[i]):
            return i
    return -1

def delete_para(p):
    p._p.getparent().remove(p._p)

# Locate Section 2 heading and Section 3 heading.
i_s2 = idx_of(lambda p: p.text.strip().startswith("Section 2"))
i_s3 = idx_of(lambda p: p.text.strip().startswith("Section 3"), i_s2+1)
sec2_head = doc.paragraphs[i_s2]
sec3_head = doc.paragraphs[i_s3]

# Capture a List Bullet style name if present, else fall back.
bullet_style = "List Bullet"
try:
    _ = doc.styles[bullet_style]
except KeyError:
    bullet_style = None

# Delete everything strictly between Section 2 heading and Section 3 heading.
for k in range(i_s3-1, i_s2, -1):
    delete_para(doc.paragraphs[k])

# Retitle Section 2 (keep its Heading 2 style/run formatting -> rewrite text only).
for r in list(sec2_head.runs):
    r._r.getparent().remove(r._r)
for r in sec2_head._p.findall(qn('w:r')):
    sec2_head._p.remove(r)
sec2_head.add_run("Section 2: Behavioral Story Bank (Study Cards)")

# Insertion helper: build paragraphs in order right BEFORE Section 3 heading.
def new_para_before(ref_p, style=None):
    new_el = copy.deepcopy(ref_p._p)
    for r in new_el.findall(qn('w:r')):
        new_el.remove(r)
    # also drop any numbering/style override leftovers by resetting pPr style below
    ref_p._p.addprevious(new_el)
    np = Paragraph(new_el, ref_p._parent)
    if style is not None:
        np.style = doc.styles[style]
    else:
        np.style = doc.styles["Normal"]
    return np

def add_heading3(text):
    p = new_para_before(sec3_head, style="Heading 3")
    p.add_run(text)
    return p

def add_italic(text):
    p = new_para_before(sec3_head, style="Normal")
    r = p.add_run(text); r.italic = True
    return p

def add_phase_label(text):
    p = new_para_before(sec3_head, style=(bullet_style or "Normal"))
    r = p.add_run(text); r.bold = True
    return p

def add_bullet(text):
    p = new_para_before(sec3_head, style=(bullet_style or "Normal"))
    p.add_run(text)
    return p

def add_blank():
    new_para_before(sec3_head, style="Normal")

# ── CARD DATA (verbatim bullets from structural_bullets.txt) ──────────────────
cards = [
  {
    "title": "Card 1 — Proactive Resilience Testing Program (0→1, disagreement & pivot)",
    "maps": "Maps to: ambiguous 0-to-1 problem · disagreement with a stakeholder · solving a complex technical issue.",
    "phases": [
      ("Situation:", [
        "Ensured cloud infrastructure resilience within the team at Microsoft; the recovery program was historically reactive — validating only when a service team requested a test.",
        "Drills had a large blast radius, risking impact to other services in the cluster that didn't want to participate.",
        "No framework existed for proactively validating resilience at scale without impacting service teams.",
        "Owned the 0-to-1 incubation of Azure's first-ever proactive resilience testing program: convert unused drill slots and scale by running smaller-blast-radius drills more often.",
      ]),
      ("The Disagreement & Pivot:", [
        "My initial model reused an existing node power-cycling script, retargeted to a rack (smaller blast radius) instead of a cluster.",
        "The Availability Platform team (owns Azure Node/VM automated recovery) and infra engineering leads raised a key concern.",
        "Azure marks nodes as Human Investigate when <20% of nodes in a cluster are powered off — Fabric (the control plane) reads it as unexpected behavior.",
        "That triggers Service Healing (VMs migrate to healthy hardware to stay available) — which wipes ephemeral data on the nodes, making it impossible to track recovery times or detect restart issues.",
      ]),
      ("Action Taken:", [
        "Pushed the first drill date back to get cross-functional alignment between drill engineering and infra teams.",
        "Kept the power-cycling script, but established the need to BLOCK Service Healing before the drill and UNBLOCK it after.",
        "Partnered with Availability Platform engineers to build the scripts that trigger that blocker.",
        "For full future autonomy: gathered engineering docs, scripts, APIs, telemetry queries and used GitHub Copilot to architect an AI agent that runs the entire drill lifecycle end-to-end.",
      ]),
      ("Result:", [
        "First end-to-end execution hit a 94% autonomous recovery rate — safely recovered 15 of 16 nodes with zero data loss.",
        "Surfaced a latent hardware defect that could have hit customers if left undetected.",
        "Gained executive visibility; strong praise from multiple CVPs; set a new standard for resilience automation across Microsoft.",
      ]),
    ],
  },
  {
    "title": "Card 2 — GDOT Platform Integration & User Adoption (influence without authority, frustrated stakeholders)",
    "maps": "Maps to: project that didn't go as planned / pivot · influencing without authority · frustrated customer · being the primary technical contact.",
    "phases": [
      ("Situation:", [
        "Hit a major bottleneck scaling the recovery program: scheduling required manual confirmation with Data Center Operations about planned hardware maintenance.",
        "Service teams had no visibility into which slots were unavailable due to rack maintenance.",
        "Goal: automate conflict-checking via an API integration with the Global Datacenter Operations Team (GDOT) platform — the central source of truth — to auto-block unavailable dates on the internal calendar.",
      ]),
      ("The Roadblock:", [
        "Discovered local data center operators for the Canary region weren't using the central GDOT platform at all.",
        "They found it unintuitive, so they tracked all campus maintenance in localized manual spreadsheets.",
        "Their data was siloed → the API strategy stalled (no reliable source to query), and I had no direct authority over local ops teams.",
      ]),
      ("Action Taken:", [
        "Reframed it as a human-adoption challenge, not a pure engineering problem.",
        "Packaged operators' usability feedback into a data-driven product-gap analysis for the central GDOT product team — influenced their long-term UX roadmap.",
        "Fed AI-generated videos/docs on how to migrate data to the new platform to solve the immediate silo.",
        "Had AI build automations to help local teams migrate spreadsheet data (maintenance schedules, scope of work, impacted hardware).",
        "Built onboarding resources, structured wiki guides, and a custom help copilot for the learning curve.",
        "Showed the 11 operators immediate value: centralized data automatically prevents drills from disrupting physical maintenance windows.",
      ]),
      ("Result:", [
        "Local teams moved 100% of their maintenance data into the central platform — eliminated siloed spreadsheets in the Canary region.",
        "Unlocked the roadmap: a clean API connection that killed manual scheduling friction and back-and-forth.",
        "Proved that true ownership means solving the human-adoption problem to make the technology effective.",
      ]),
    ],
  },
  {
    "title": "Card 3 — Resilience Automation Platform / Self-Service Intake (trade-offs at scale)",
    "maps": "Maps to: automating a manual/inefficient process · difficult trade-off (completeness vs. speed) · product strategy / prioritization.",
    "phases": [
      ("Situation:", [
        "Team ran Azure's high-impact resilience drills but hit a scaling wall from operational toil.",
        "Every service-team request needed hours of manual intake, timeline coordination, scope definition, and failure-mode mapping.",
        "Manual friction consumed nearly 6 hours per drill for planning and scoping.",
        "Team KPI: scale up and sustain 45+ major resilience drills a year — unsustainable manually.",
        "Trying to go from 1 drill a month to 1 drill a week.",
      ]),
      ("The Tension & Trade-off:", [
        "As owner of the self-service automation initiative, I got numerous competing feature requests from service teams and engineering leads — under a strict MVP launch deadline.",
        "Set a strict prioritization framework: operational impact vs. execution velocity, to protect momentum.",
        "Focused 100% of initial engineering on the self-service scheduling/intake flow to offload manual work.",
        "Explicitly deprioritized secondary features (e.g. post-drill telemetry dashboards with recovery metrics/RCAs) because that data could be pulled manually for now.",
        "Partnered with Observability to find the source of truth for Node-to-Service mapping, so teams could see which hardware their subscriptions were tied to — replacing a tedious, inaccurate ad-hoc query process.",
      ]),
      ("Action Taken:", [
        "Partnered with Dev and PM leads to lock the MVP feature set.",
        "Used AI agents to rapidly prototype: calendar view, drill request form, Service ID lookup / resource mapping, compute-infra failure modes, exec comms, RBAC, and remote execution.",
        "Set up an A/B testing framework to evaluate layout configurations and visual hierarchy.",
        "Gathered telemetry on drop-off and completion times to make the final layout frictionless for technical and non-technical teams.",
      ]),
      ("Result:", [
        "Launched the self-service platform exactly on time by protecting MVP scope.",
        "The A/B-tested UI drove a 35% reduction in intake-form completion time with a near-zero error rate.",
        "Crushed a 6-hour manual coordination process down to 6 minutes.",
        "Reclaimed critical engineering hours — sustained the 45-drills-per-year milestone without adding headcount.",
      ]),
    ],
  },
]

# Intro italic line under the section heading (insert first, before Section 3).
add_italic("Rehearse from the bullets — each card is a fact skeleton, not a script. Hit the labeled beats in order and stop. Map any behavioral question to the closest card.")
add_blank()

for c in cards:
    add_heading3(c["title"])
    add_italic(c["maps"])
    for label, bullets in c["phases"]:
        add_phase_label(label)
        for b in bullets:
            add_bullet(b)
    add_blank()

doc.save(str(OUT))
print("STEP 2 done. Total paras:", len(doc.paragraphs))
