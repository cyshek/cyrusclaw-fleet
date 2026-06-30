#!/usr/bin/env python3
"""
Build the Google TPM AI-Enablement guide on the NEW corrected template:
  - Copy templates/Master_Interview_Prep_Guide.docx -> Google guide
  - Fill Company block (Google AI+Infra) and Role block (TPM AI Enablement) spoken answers
  - Append "Recruiter Screen Game Plan" section tuned to Sirisha's 15-min call,
    in bold-topic + exact-spoken-line format, with LOCKED answers:
      Kirkland local / US citizen no sponsorship / comp / why-Google / why-leaving.
  - Rebuild guides/google-tpm-ai-enablement/ files + zip (guide + JD + resume).
Preserves template styling (we edit the real template copy in place).
"""
import shutil, zipfile, datetime
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATE = WS / "templates/Master_Interview_Prep_Guide.docx"
GDIR = WS / "guides/google-tpm-ai-enablement"
GUIDE = GDIR / "Google_TPM_AI_Enablement_Interview_Prep_Guide.docx"
JD = GDIR / "Google_TPM_AI_Enablement_JD.md"
RESUME = GDIR / "Cyrus_Shekari_Resume.pdf"
ZIP = WS / "guides/google-tpm-ai-enablement.zip"

# backup current guide
bk = Path("/tmp") / f"google_guide_backup_{datetime.datetime.now():%Y%m%d_%H%M%S}.docx"
shutil.copy2(GUIDE, bk)
print("backup ->", bk)

shutil.copy2(TEMPLATE, GUIDE)
doc = Document(str(GUIDE))

def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for r in p._p.findall(qn('w:r')):
        p._p.remove(r)

def fill_placeholder(marker, label, rest):
    for p in doc.paragraphs:
        if marker in p.text:
            clear_runs(p)
            r1 = p.add_run(label); r1.bold = True
            p.add_run(rest)
            return True
    return False

COMPANY_REST = ("Google's AI and Infrastructure org is the team that delivers AI and compute at "
  "hyperscale — TPUs, Vertex AI, Global Networking, data center operations — for Googlers, Cloud "
  "customers, and billions of users. Why I want in: it's the highest-leverage place in the world to do "
  "exactly what I already do — make massive infrastructure reliable and get AI into the operational "
  "loop — and Google opens doors to places like DeepMind and Waymo over time.")
ROLE_REST = ("The TPM, AI Enablement role leads complex, cross-functional infrastructure programs end to "
  "end — scoping, milestones, risk, and driving consensus across many interdependent stakeholders to "
  "scale and operationalize AI solutions. Why it's right for me: 'scale and operationalize AI' is "
  "essentially my job at Microsoft today — I ran a 0-to-1 resilience program and built AI agents to "
  "automate it — and I want to do that at Google's scale.")

assert fill_placeholder("[COMPANY", "Google (AI & Infrastructure) — what + why: ", COMPANY_REST), "company ph missing"
assert fill_placeholder("[ROLE", "TPM, AI Enablement — what + why: ", ROLE_REST), "role ph missing"

# ── Build the Recruiter Screen section at the end (before sectPr) ──────────────
new_body = doc.element.body
sectPr = new_body.find(qn('w:sectPr'))

# get a List Bullet style name if any
bullet_style = None
try:
    _ = doc.styles["List Bullet"]; bullet_style = "List Bullet"
except KeyError:
    bullet_style = None

def add_par(style=None):
    p_el = OxmlElement('w:p')
    new_body.insert(list(new_body).index(sectPr), p_el)
    p = Paragraph(p_el, doc)
    if style:
        p.style = doc.styles[style]
    else:
        p.style = doc.styles["Normal"]
    return p

def heading(text, level):
    p = add_par(style=f"Heading {level}")
    p.add_run(text)
    return p

def plain(text):
    p = add_par(style="Normal")
    p.add_run(text)
    return p

def topic_line(topic, said):
    """bold 'Topic: ' + plain spoken line"""
    p = add_par(style="Normal")
    r = p.add_run(topic + ": "); r.bold = True
    p.add_run(said)
    return p

def bullet(text):
    p = add_par(style=(bullet_style or "Normal"))
    p.add_run(text)
    return p

# ---- Section content ----
heading("Recruiter Screen Game Plan — Sirisha Susarla (Wed Jul 1, 11:00 AM PDT, phone)", 2)
plain("15-minute non-technical recruiter call. Her stated goal: discuss the opportunity, next steps, and how to prepare. This is a FIT + LOGISTICS gate, not a technical or hiring-manager interview. Be warm, crisp, obviously qualified, clearly interested, and easy to schedule. ~6-8 exchanges total — keep answers 30-45 seconds and let her drive.")

heading("Say-it-ready answers (bold topic, then the exact line)", 3)
topic_line("Walk me through your background",
    "\"I'm a Technical Program Manager at Microsoft on cloud reliability — I make sure our infrastructure "
    "survives worst-case failures. I scaled a 0-to-1 proactive resilience program that hit a 94% "
    "autonomous recovery rate, and I built AI agents on GitHub Copilot to automate the whole drill "
    "lifecycle. Now I'm looking to bring that infrastructure and AI-operationalization work to a bigger "
    "platform like Google.\"")
topic_line("Why Google / why this role",
    "\"Google's AI and Infrastructure org is the highest-leverage place in the world to do exactly what I "
    "already do — make hyperscale infrastructure reliable and get AI into the operational loop. The "
    "role's mandate to scale and operationalize AI solutions is essentially my current job, at Google "
    "scale.\"")
topic_line("Do the must-have quals line up",
    "\"Yes — I've got the 2-plus years of program management as a TPM, and the AI/ML side is the AI agents "
    "I built to run our drill lifecycle plus the AI-prototyped automation platform. I also hit the "
    "preferred data-center infrastructure experience through my GDOT data-center maintenance integration "
    "and node-to-service infra mapping.\"")
topic_line("Work authorization / sponsorship",
    "\"I'm a US citizen — no sponsorship needed, now or in the future.\"")
topic_line("Location — are you open to one of the sites",
    "\"I'm actually based in Kirkland, so I'm local to that campus — no relocation needed, I can be onsite there.\"")
topic_line("Compensation expectations",
    "\"I'm looking at total compensation and I'm flexible for the right fit. The posted range works for me "
    "— I care most about the role and the platform.\"")
topic_line("Why are you leaving Microsoft",
    "\"I've had a great run at Microsoft, but the resilience program matured from a 0-to-1 build into "
    "steady-state, and I add the most value in the high-ambiguity building phase — which is exactly where "
    "this team is. Nothing negative; I'm running toward this, not away.\"")

heading("Light Google context (only if it comes up — it's a recruiter call)", 3)
bullet("The org — AI & Infrastructure: delivers AI + compute at hyperscale. Key teams: TPUs, Vertex AI, Google Global Networking, Data Center operations. Reliability + scale is the through-line — your home turf.")
bullet("\"AI Enablement\" = scaling and operationalizing AI across the enterprise: find gaps, standardize the approach, drive consensus across interdependent workstreams. You've literally done 0-to-1 AI-operationalization.")
bullet("Heads-up for the LATER loop (not this call): Google TPMs are expected to be technical — system design + engineering tradeoffs. Ask Sirisha what the loop looks like so you can prep for it.")

heading("Questions to ask Sirisha", 3)
bullet("\"What does the full process look like after this — how many rounds, and what should I focus my prep on?\"")
bullet("\"Who's the hiring manager, and what's the team's top priority for this role in the first 6-12 months?\"")
bullet("\"Is the role tied to a specific site or team, or is that decided later?\"")
bullet("\"How technical is the loop — should I expect system design or a live exercise?\"")
bullet("\"What's the timeline, and is there anything on my background you'd want me to elaborate on for the hiring manager?\"")

heading("Mindset & logistics", 3)
bullet("You've cleared Google's candidate bar before (Jan TPM-2: 'strong candidate, not a fit for that team') and a recruiter just sourced you again — you're a known quantity they reached for twice. Walk in from strength.")
bullet("Lead with the easy wins early: you're local (Kirkland) and a US citizen — recruiters love a clean, no-friction candidate. Get those on the table.")
bullet("It's a phone call: quiet room, good signal, resume + JD open, water, notepad. Pick up promptly and energetic.")
bullet("Same day: Inworld HM screen is at 2:00 PM PDT — totally different conversation. Don't let them bleed; this one is warm + concise + fit-focused.")
bullet("Close by confirming next steps, reiterating genuine interest, and thanking her by name. Ask when you'll hear back.")

doc.save(str(GUIDE))
print("Google guide built. paras:", len(Document(str(GUIDE)).paragraphs))

# Rebuild zip (guide + JD + resume)
with zipfile.ZipFile(str(ZIP), "w", zipfile.ZIP_DEFLATED) as z:
    z.write(GUIDE, GUIDE.name)
    if JD.exists(): z.write(JD, JD.name)
    if RESUME.exists(): z.write(RESUME, RESUME.name)
with zipfile.ZipFile(str(ZIP)) as z:
    print("ZIP:", z.namelist())
