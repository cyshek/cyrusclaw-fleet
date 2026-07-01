#!/usr/bin/env python3
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
GUIDE = WS / "guides/inworld-ai-founding-se/Inworld_AI_Founding_SE_Interview_Prep_Guide.docx"
d = Document(str(GUIDE))

# 1) placeholders
bad_markers = ["[COMPANY", "[ROLE", "[Fill"]
ph_hits = []
for i, p in enumerate(d.paragraphs):
    for m in bad_markers:
        if m in p.text:
            ph_hits.append((i, m, p.text[:60]))
print("PLACEHOLDERS REMAINING:", len(ph_hits))
for h in ph_hits:
    print("   ", h)

# 2) say-it topic lines: in the appended "Say-it-ready" block, every topic_line
#    is Normal style whose first run is bold and text ends with ':'.
#    Identify them: Normal paragraphs where runs[0].bold is True and runs[0].text ends ': '
say_lines = []
for p in d.paragraphs:
    if p.style.name == "normal" and p.runs:
        r0 = p.runs[0]
        if r0.bold and r0.text.rstrip().endswith(":"):
            say_lines.append(p)
print("SAY-IT TOPIC LINES FOUND:", len(say_lines))
all_ok = True
for p in say_lines:
    r0 = p.runs[0]
    ok_bold = (r0.bold is True)
    ok_colon = r0.text.rstrip().endswith(":")
    if not (ok_bold and ok_colon):
        all_ok = False
        print("   FAIL:", repr(r0.text[:40]), "bold=", r0.bold)
    else:
        print("   ok  :", repr(r0.text.strip()))
print("ALL SAY-IT LINES bold==True AND end ':' ->", all_ok)

# 3) key facts present anywhere
full = "\n".join(p.text for p in d.paragraphs)
facts = ["Florin", "Realtime API", "GitHub Copilot", "US citizen", "Service Healing"]
print("FACTS:")
for f in facts:
    print("   ", f, "->", ("PRESENT" if f in full else "*** MISSING ***"))

# 4) heading colors: confirm no explicit blue; check Heading runs color
def run_color(r):
    rPr = r._r.find(qn('w:rPr'))
    if rPr is None:
        return None
    c = rPr.find(qn('w:color'))
    if c is None:
        return None
    return c.get(qn('w:val'))

blue_like = []
heads = 0
for p in d.paragraphs:
    if p.style.name.startswith("Heading"):
        heads += 1
        for r in p.runs:
            col = run_color(r)
            if col and col.lower() not in ("000000", "auto"):
                # flag anything that looks blue-ish
                blue_like.append((p.style.name, col, p.text[:40]))
print("HEADINGS:", heads, "| explicit non-black/auto run colors:", len(blue_like))
for b in blue_like:
    print("   ", b)

# Also check the Heading styles' defined color in styles.xml (theme)
from docx.oxml.ns import nsmap
styles_part = d.part.styles_part if hasattr(d.part, 'styles_part') else None
print("--- Heading style defined colors ---")
for sname in ["Heading 1", "Heading 2", "Heading 3"]:
    st = d.styles[sname]
    el = st.element
    col = None
    rpr = el.find(qn('w:rPr'))
    if rpr is not None:
        c = rpr.find(qn('w:color'))
        if c is not None:
            col = (c.get(qn('w:val')), c.get(qn('w:themeColor')))
    print("   ", sname, "->", col)
