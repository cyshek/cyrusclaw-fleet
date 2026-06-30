"""
Verification checker for the interview-prep build.

Checks:
  T1  Template opens cleanly via Document(); print paragraph count.
  T2  Template STILL has the [Fill placeholders (and [COMPANY/[ROLE).
  D1  Datadog guide opens cleanly via Document(); print paragraph count.
  D2  Datadog guide has NO [COMPANY / [ROLE / [Fill placeholders.
  D3  Company + Role blocks: first run bold and ends with ':'.
  D4  Every cheat-sheet entry: first run bold and ends with ':'.
  D5  Critical facts survive (string presence) in the Datadog guide.
  Z1  Zip namelist contains the docx + a JD + a resume.
  B1  Master: random spot-check that key labels are own bold run.
"""

import zipfile
from pathlib import Path
from docx import Document

WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATE = WORKSPACE / "templates" / "Master_Interview_Prep_Guide.docx"
GUIDE = WORKSPACE / "bundles" / "datadog-partner-tse" / "Datadog_Partner_TSE_Interview_Prep_Guide.docx"
ZIP_OUT = WORKSPACE / "bundles" / "datadog-partner-tse" / "Datadog_Partner_TSE_PrepBundle.zip"

results = []


def check(name, ok, detail=""):
    status = "PASS" if ok else "FAIL"
    results.append(ok)
    line = "[" + status + "] " + name
    if detail:
        line += " :: " + detail
    print(line)


# ---------------------------------------------------------------------------
# T1 / T2 — template
# ---------------------------------------------------------------------------
try:
    tdoc = Document(TEMPLATE)
    t_paras = len(tdoc.paragraphs)
    check("T1 template opens via Document()", True, "paragraphs=" + str(t_paras))
except Exception as exc:
    check("T1 template opens via Document()", False, repr(exc))
    tdoc = None
    t_paras = -1

if tdoc is not None:
    t_text = "\n".join(p.text for p in tdoc.paragraphs)
    has_fill = "[Fill" in t_text
    has_company = "[COMPANY" in t_text
    has_role = "[ROLE" in t_text
    check(
        "T2 template KEEPS placeholders ([Fill, [COMPANY, [ROLE)",
        has_fill and has_company and has_role,
        "Fill=" + str(has_fill) + " COMPANY=" + str(has_company) + " ROLE=" + str(has_role),
    )

# ---------------------------------------------------------------------------
# D1 / D2 — datadog guide opens + placeholders gone
# ---------------------------------------------------------------------------
try:
    gdoc = Document(GUIDE)
    g_paras = len(gdoc.paragraphs)
    check("D1 datadog guide opens via Document()", True, "paragraphs=" + str(g_paras))
except Exception as exc:
    check("D1 datadog guide opens via Document()", False, repr(exc))
    gdoc = None
    g_paras = -1

if gdoc is not None:
    g_text = "\n".join(p.text for p in gdoc.paragraphs)
    no_company = "[COMPANY" not in g_text
    no_role = "[ROLE" not in g_text
    no_fill = "[Fill" not in g_text
    check(
        "D2 datadog guide has NO [COMPANY / [ROLE / [Fill",
        no_company and no_role and no_fill,
        "no_COMPANY=" + str(no_company) + " no_ROLE=" + str(no_role) + " no_Fill=" + str(no_fill),
    )

    # -----------------------------------------------------------------------
    # D3 — Company + Role blocks: first run bold + ends ':'
    # -----------------------------------------------------------------------
    company_ok = False
    role_ok = False
    for p in gdoc.paragraphs:
        if p.runs and p.runs[0].text.strip().startswith("Datadog \u2014 what + why"):
            r0 = p.runs[0]
            company_ok = (r0.bold is True) and r0.text.strip().endswith(":")
        if p.runs and p.runs[0].text.strip().startswith("Partner TSE \u2014 what + why"):
            r0 = p.runs[0]
            role_ok = (r0.bold is True) and r0.text.strip().endswith(":")
    check("D3 Company block first run bold + ':'", company_ok)
    check("D3 Role block first run bold + ':'", role_ok)

    # -----------------------------------------------------------------------
    # D4 — every cheat-sheet entry: first run bold + ends ':'
    # -----------------------------------------------------------------------
    cheat_terms = [
        "Metrics:",
        "Logs:",
        "Traces / APM:",
        "OpenTelemetry (OTEL):",
        "Datadog Agent (agent-based vs API-based):",
        "Integration Developer Platform (IDP) / Marketplace:",
        "OAuth:",
    ]
    found = {}
    for p in gdoc.paragraphs:
        if not p.runs:
            continue
        first = p.runs[0].text.strip()
        for term in cheat_terms:
            if first == term:
                found[term] = (p.runs[0].bold is True) and first.endswith(":")
    all_cheat_ok = True
    for term in cheat_terms:
        ok = found.get(term, False)
        if not ok:
            all_cheat_ok = False
        check("D4 cheat term bold+':' -> " + term, ok)
    check("D4 ALL cheat-sheet entries bold+':'", all_cheat_ok, str(len(found)) + "/" + str(len(cheat_terms)) + " found")

    # -----------------------------------------------------------------------
    # D5 — critical facts survive
    # -----------------------------------------------------------------------
    critical = [
        "Human Investigate",
        "less than 20% of nodes",
        "Service Healing",
        "94% autonomous recovery rate",
        "15 out of 16",
        "zero data loss",
        "11 operators",
        "6 hours",
        "six minutes",
        "35%",
        "45",
        "RBAC",
        "Node-to-Service",
    ]
    missing = [c for c in critical if c not in g_text]
    check("D5 critical facts survive in Datadog guide", len(missing) == 0, "missing=" + str(missing))

# ---------------------------------------------------------------------------
# B1 — master spot-check key labels are own bold run
# ---------------------------------------------------------------------------
if tdoc is not None:
    want_bold_first = ["Situation:", "The Disagreement & Pivot:", "Action Taken:", "Result:",
                       "De-escalate and Diagnose:", "Evaluate the Trade-offs:",
                       "Propose a Scoped Solution:"]
    seen = {}
    for p in tdoc.paragraphs:
        if not p.runs:
            continue
        first = p.runs[0].text.strip()
        for lbl in want_bold_first:
            if first == lbl:
                seen[lbl] = (p.runs[0].bold is True)
    all_master_ok = all(seen.get(lbl, False) for lbl in want_bold_first)
    check("B1 master phase/step labels are own bold run", all_master_ok,
          str(sum(1 for lbl in want_bold_first if seen.get(lbl))) + "/" + str(len(want_bold_first)))

    # Also confirm master critical facts survive
    t_crit = ["Human Investigate", "Service Healing", "94% autonomous recovery rate",
              "15 out of 16", "zero data loss", "11 operators", "6 hours per drill",
              "six minutes", "35%", "45", "RBAC", "Node-to-Service"]
    t_missing = [c for c in t_crit if c not in t_text]
    check("B1b master critical facts survive", len(t_missing) == 0, "missing=" + str(t_missing))

# ---------------------------------------------------------------------------
# Z1 — zip namelist
# ---------------------------------------------------------------------------
try:
    with zipfile.ZipFile(ZIP_OUT) as zf:
        names = zf.namelist()
    has_docx = any(n.endswith(".docx") for n in names)
    has_jd = any("JD" in n for n in names)
    has_resume = any("Resume" in n and n.endswith(".pdf") for n in names)
    check("Z1 zip has docx + JD + resume PDF", has_docx and has_jd and has_resume, str(names))
except Exception as exc:
    check("Z1 zip opens / namelist", False, repr(exc))

# ---------------------------------------------------------------------------
print("=" * 60)
print("TEMPLATE paragraphs:", t_paras)
print("DATADOG paragraphs :", g_paras)
total = len(results)
passed = sum(1 for r in results if r)
print("OVERALL:", passed, "/", total, "checks PASS")
print("RESULT:", "ALL PASS" if passed == total else "SOME FAIL")
