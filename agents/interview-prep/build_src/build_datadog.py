"""
TASK 2 — Rebuild the Datadog Partner TSE bundle on top of the NEW master template.

Steps:
  1. Copy the new templates/Master_Interview_Prep_Guide.docx -> guide.
  2. Replace the Q2 [COMPANY ...] placeholder with the spoken Company answer.
  3. Replace the Q3 [ROLE ...] placeholder with the spoken Role answer.
  4. Append "Round 1 Prep — Pulkit Chandra" (theme-by-theme prep, kept) and a
     REFORMATTED observability cheat sheet (bold term + plain spoken one-liner),
     plus "Questions to ask Pulkit" and "Mindset".
  5. Rebuild the zip (docx + JD.md + resume PDF) via python zipfile.

Every label is its OWN bold run (p.runs[0].bold == True, ends with ':').
No heredocs / inline \\n. Real multi-line bodies.
"""

import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt

WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
JOB_SEARCH = Path(
    "/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/"
    "applications/submitted/datadog-7961297"
)
OUT_DIR = WORKSPACE / "bundles" / "datadog-partner-tse"
TEMPLATE = WORKSPACE / "templates" / "Master_Interview_Prep_Guide.docx"
GUIDE_NAME = "Datadog_Partner_TSE_Interview_Prep_Guide.docx"
GUIDE_OUT = OUT_DIR / GUIDE_NAME
RESUME_PDF = JOB_SEARCH / "Cyrus_Shekari_Resume_datadog_7961297_v2.pdf"
JD_MD = JOB_SEARCH / "JD.md"
ZIP_OUT = OUT_DIR / "Datadog_Partner_TSE_PrepBundle.zip"

OUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def clear_paragraph(para):
    """Remove every run element from a paragraph."""
    p_elem = para._p
    for r_elem in list(p_elem):
        if r_elem.tag.endswith("}r"):
            p_elem.remove(r_elem)


def fill_label_then_plain(para, label, rest):
    """Rewrite a paragraph to two runs: bold label + plain rest."""
    clear_paragraph(para)
    r0 = para.add_run(label)
    r0.bold = True
    r1 = para.add_run(" " + rest)
    r1.bold = False


def add_heading(document, text, level):
    return document.add_heading(text, level=level)


def add_plain(document, text, style="Normal"):
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = False
    return p


def add_italic(document, text, style="Normal"):
    p = document.add_paragraph(style=style)
    run = p.add_run(text)
    run.italic = True
    return p


def add_bullet_plain(document, text):
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.bold = False
    return p


def add_term_then_plain(document, term, rest, style="List Bullet"):
    """
    Cheat-sheet / spoken entry: bold term label (ends in ':') + plain one-liner.
    run[0].bold == True, run[0].text endswith ':'.
    """
    p = document.add_paragraph(style=style)
    r0 = p.add_run(term)
    r0.bold = True
    r1 = p.add_run(" " + rest)
    r1.bold = False
    return p


# ---------------------------------------------------------------------------
# Step 1 — fresh copy of the new template
# ---------------------------------------------------------------------------
shutil.copy2(TEMPLATE, GUIDE_OUT)
print("[1] Copied new template ->", GUIDE_OUT.name)

doc = Document(GUIDE_OUT)

# ---------------------------------------------------------------------------
# Step 2 + 3 — fill Company (Q2) and Role (Q3) placeholders
# ---------------------------------------------------------------------------
company_para = None
role_para = None
for p in doc.paragraphs:
    if "[COMPANY" in p.text:
        company_para = p
    elif "[ROLE" in p.text:
        role_para = p

assert company_para is not None, "Could not find [COMPANY placeholder"
assert role_para is not None, "Could not find [ROLE placeholder"

fill_label_then_plain(
    company_para,
    "Datadog \u2014 what + why:",
    "Datadog is the leading cloud observability and monitoring platform \u2014 "
    "metrics, logs, traces, and security all in one place, used by engineering "
    "and DevOps teams worldwide to understand what's happening across their "
    "infrastructure and applications. Why I want in: Datadog won the "
    "observability consolidation war, and the integration ecosystem is the growth "
    "engine that keeps extending that platform \u2014 I want to be at that edge, "
    "where new partners plug the rest of the world into Datadog.",
)
print("[2] Filled Q2 Company block")

fill_label_then_plain(
    role_para,
    "Partner TSE \u2014 what + why:",
    "The Partner TSE is the primary technical contact for third-party developers "
    "building integrations on the Integration Developer Platform (IDP) \u2014 I'd "
    "guide them on architecture, review their code against the Quality Rubric, "
    "troubleshoot the hard issues, and feed the friction they hit back to product "
    "and engineering. Why it's right for me: it's the exact intersection of "
    "technical depth and external impact \u2014 consultant, engineer, and product "
    "voice at once \u2014 which is where I do my best work.",
)
print("[3] Filled Q3 Role block")

doc.save(GUIDE_OUT)

# ---------------------------------------------------------------------------
# Step 4 — append Round 1 Pulkit section + reformatted cheat sheet
# ---------------------------------------------------------------------------
doc = Document(GUIDE_OUT)

add_heading(doc, "Round 1 Prep \u2014 Pulkit Chandra (Hiring Manager)", level=1)

# --- What Pulkit is evaluating ---
add_heading(doc, "What Pulkit is evaluating", level=2)
add_plain(
    doc,
    "Per Taylor's email, this is a conversational interview \u2014 not a case "
    "study. Pulkit wants to understand how you think and how you've worked. Five "
    "themes from the prep note: (1) why you're here / why Datadog specifically, "
    "(2) real partner or developer relationships you've owned end-to-end, "
    "(3) observability stack knowledge and architecture decisions, (4) code "
    "quality and integration design advising, (5) times you pushed back or "
    "influenced product/eng based on external signal.",
)

# --- Theme-by-theme prep (kept) ---
add_heading(doc, "Theme-by-theme prep", level=2)
add_bullet_plain(
    doc,
    "Why you're here / why Datadog: Lead with the transition story \u2014 you've "
    "been driving large-scale technical programs at Microsoft, you're most "
    "energized at the intersection of technical depth and external impact, and "
    "Datadog's Partner TSE role is the specific combination of technical "
    "consulting + product feedback loop that you're looking for. Don't just say "
    "'I like the product' \u2014 say what you know: Datadog is the platform that "
    "won the observability consolidation war, the integration ecosystem is how "
    "they extend it, and you want to be at that growth edge.",
)
add_bullet_plain(
    doc,
    "Real partner/developer relationships end-to-end: Your best story is the GDOT "
    "platform adoption pivot (Card 2) \u2014 you owned the relationship with local "
    "data center teams and the central platform team, diagnosed why users weren't "
    "adopting the tool, and fixed it. Map it explicitly: 'the local operators were "
    "my external developers, the GDOT team was my internal product partner, and "
    "the problem was a friction point in the platform experience.' That's the "
    "exact TSE motion.",
)
add_bullet_plain(
    doc,
    "Observability stack knowledge: Know these cold \u2014 Metrics (counters, "
    "gauges, histograms), Logs (structured vs unstructured, pipelines), Traces "
    "(distributed tracing, spans, APM), OpenTelemetry (OTEL \u2014 the open "
    "standard for instrumentation that Datadog supports natively). Agent-based vs "
    "API-based integrations: agent runs on the host and pulls data directly; "
    "API-based polls a remote endpoint. OAuth flows: how partner integrations "
    "authenticate. You don't need to be a software engineer \u2014 you need to "
    "know how the pieces connect so you can advise on architecture decisions.",
)
add_bullet_plain(
    doc,
    "Code quality / integration design advising: The role does code reviews "
    "against Datadog's Quality Rubric. You won't be asked to write code in the "
    "interview, but Pulkit may ask how you'd approach a situation where a "
    "partner's integration works but doesn't meet platform standards. Answer: you "
    "start with curiosity (what was their constraint?), then explain the standard "
    "and why it matters for the ecosystem, then collaborate on a path to "
    "compliance rather than just blocking them. Your Card 1 story (disagreeing "
    "with engineering leads, data-first approach) maps well here.",
)
add_bullet_plain(
    doc,
    "Pushing back / influencing product: Your strongest story here is Card 2 "
    "(influencing without authority) \u2014 you had no direct authority over the "
    "GDOT product team or local operators, but you packaged external feedback into "
    "a product brief, got the feature shipped, and then drove adoption from the "
    "ground up. That is textbook TSE-to-internal-product influence. Also: the AI "
    "drill agent story shows you identified a manual bottleneck, built the case, "
    "and shipped automation \u2014 that's the 'identify friction, feed it back to "
    "eng' motion too.",
)

# --- Datadog observability cheat sheet (REFORMATTED: bold term + spoken one-liner) ---
add_heading(doc, "Datadog observability cheat sheet", level=2)
add_italic(
    doc,
    "Drill these the morning of \u2014 say each one out loud. Bold term, then a "
    "plain line you could actually say in the room.",
)
add_term_then_plain(
    doc,
    "Metrics:",
    "Numbers tracked over time \u2014 things like CPU percent, request count, or "
    "p99 latency. Counters only go up, gauges go up and down, and histograms show "
    "the distribution.",
)
add_term_then_plain(
    doc,
    "Logs:",
    "Timestamped records of individual events. Datadog ingests them, parses them, "
    "and indexes them so you can search \u2014 and structured JSON logs are far "
    "easier to query than raw unstructured text.",
)
add_term_then_plain(
    doc,
    "Traces / APM:",
    "Distributed tracing across microservices \u2014 a trace is one request's "
    "whole journey, and spans are the individual operations inside it. That's how "
    "you pinpoint exactly where the latency is coming from.",
)
add_term_then_plain(
    doc,
    "OpenTelemetry (OTEL):",
    "The vendor-neutral open standard for instrumenting code. Datadog supports it "
    "natively, so a partner can instrument once with OTEL and ship that data "
    "straight into Datadog \u2014 it's the modern default for new integrations.",
)
add_term_then_plain(
    doc,
    "Datadog Agent (agent-based vs API-based):",
    "The Agent is a lightweight process that runs ON the host, collects "
    "metrics/logs/traces locally, and ships them up \u2014 that's an agent-based "
    "integration, and it's closer to real-time. The contrast is an API-based "
    "integration, which runs no agent and instead polls a remote API on an "
    "interval \u2014 easier to set up, but less real-time.",
)
add_term_then_plain(
    doc,
    "Integration Developer Platform (IDP) / Marketplace:",
    "The IDP is where third-party developers build and publish integrations. "
    "Community ones live in integrations-extras; commercial ones go on the Datadog "
    "Marketplace \u2014 and the Partner TSE owns the quality gate between a "
    "partner's code and publication.",
)
add_term_then_plain(
    doc,
    "OAuth:",
    "The standard way partner integrations authenticate \u2014 instead of handing "
    "over a raw API key, the integration gets a scoped, revocable token to act on "
    "a customer's behalf, which is what keeps the connection secure.",
)

# --- Questions to ask Pulkit (kept) ---
add_heading(doc, "Questions to ask Pulkit", level=2)
add_bullet_plain(
    doc,
    "'What does the integration lifecycle look like end-to-end \u2014 from first "
    "contact with a partner to publication? Where does the TSE have the most "
    "leverage?'",
)
add_bullet_plain(
    doc,
    "'What's the biggest friction point partners hit on the IDP today, and how is "
    "the team thinking about fixing it?'",
)
add_bullet_plain(
    doc,
    "'How do TSEs surface product signal back to the engineering team \u2014 is "
    "there a formal process or is it more ad hoc?'",
)

# --- Mindset (kept) ---
add_heading(doc, "Mindset", level=2)
add_plain(
    doc,
    "Taylor said 'be concrete and honest rather than trying to give perfect "
    "answers.' Pulkit wants to see how you think through real situations, not "
    "whether you can recite Datadog's docs. Lead with your actual experiences, map "
    "them explicitly to what the TSE role does, and show genuine curiosity about "
    "the observability space. You have the cross-functional execution, the "
    "external-relationship stories, and the product feedback loop \u2014 make them "
    "visible.",
)

doc.save(GUIDE_OUT)
print("[4] Appended Round 1 Pulkit section + reformatted cheat sheet")

# ---------------------------------------------------------------------------
# Step 5 — rebuild the zip via python zipfile
# ---------------------------------------------------------------------------
bundle_files = [
    (GUIDE_OUT, GUIDE_NAME),
    (JD_MD, "JD.md"),
    (RESUME_PDF, "Cyrus_Shekari_Resume_datadog_7961297_v2.pdf"),
]

with zipfile.ZipFile(ZIP_OUT, "w", zipfile.ZIP_DEFLATED) as zf:
    for src_path, arc_name in bundle_files:
        zf.write(src_path, arc_name)
        print("[5] Added to zip:", arc_name, "(", src_path.stat().st_size, "bytes )")

print("[5] Zip written ->", ZIP_OUT.name, "(", ZIP_OUT.stat().st_size, "bytes )")
print("DONE")
