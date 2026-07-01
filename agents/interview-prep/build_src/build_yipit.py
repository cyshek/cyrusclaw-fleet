#!/usr/bin/env python3
"""
Build the YipitData Associate Product Manager (APM) interview-prep guide on the
NEW corrected canonical template:
  - Copy templates/Master_Interview_Prep_Guide.docx -> YipitData guide (edit in place)
  - Fill Company block (YipitData) and Role block (APM Metrics & Feeds) spoken answers
  - Append "Recruiter Screen Game Plan — Lorena Gallo (Thu Jul 2, 9:00 AM PDT, Zoom)"
    section in bold-topic + exact-spoken-line format (LOCKED style), preserving the
    real substance from build_src/yipit_section6.txt reformatted to the locked style.
  - Rebuild guides/yipitdata-apm.zip (guide docx + JD.md + resume.pdf) via zipfile.
Preserves template styling (headings black; we edit the real template copy in place).
"""
import shutil
import zipfile
import datetime
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATE = WS / "templates/Master_Interview_Prep_Guide.docx"
GDIR = WS / "guides/yipitdata-apm"
GUIDE = GDIR / "YipitData_APM_Interview_Prep_Guide.docx"
JD = GDIR / "YipitData_APM_JD.md"
RESUME = GDIR / "Cyrus_Shekari_Resume_yipitdata_7892101_v2.pdf"
ZIP = WS / "guides/yipitdata-apm.zip"

# backup current guide to /tmp
bk = Path("/tmp") / ("yipit_guide_prebuild_%s.docx" % datetime.datetime.now().strftime("%Y%m%d_%H%M%S"))
if GUIDE.exists():
    shutil.copy2(GUIDE, bk)
    print("backup ->", bk)

shutil.copy2(TEMPLATE, GUIDE)
doc = Document(str(GUIDE))


def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for r in p._p.findall(qn('w:r')):
        p._p.remove(r)


def fill_placeholder(marker, label, rest):
    for p in doc.paragraphs:
        if marker in p.text:
            clear_runs(p)
            r1 = p.add_run(label)
            r1.bold = True
            p.add_run(rest)
            return True
    return False


COMPANY_REST = (
    "YipitData is the leading alternative-data market-research and analytics firm for the "
    "disruptive economy — software, AI, cloud, e-commerce, ridesharing, payments — and recently "
    "raised $475M from The Carlyle Group at a $1B-plus valuation. Top investment funds and Fortune "
    "500 companies make high-stakes decisions on their data, so data quality isn't a feature — the "
    "data IS the product. Why I want in: that precision bar is exactly where I do my best work, and "
    "the push into AI-powered, programmatic delivery makes this the right moment to join."
)
ROLE_REST = (
    "The Associate Product Manager on Metrics & Feeds (Public Investor team) owns metrics-publishing "
    "quality and standards, partners cross-functionally with Data, Engineering, and Product, and "
    "scales the data-platform workflows from ingestion through transformation to delivery while "
    "leveraging AI tools. Why it fits: I scaled a 0-to-1 data-validation platform, defined metrics "
    "standards, drove cross-functional adoption, and have lived AI-tool automation — which is this "
    "role almost verbatim."
)

assert fill_placeholder("[COMPANY", "YipitData — what + why: ", COMPANY_REST), "company ph missing"
assert fill_placeholder("[ROLE", "Associate Product Manager, Metrics & Feeds — what + why: ", ROLE_REST), "role ph missing"

# -- Build the Recruiter Screen section at the end (before sectPr) -------------
new_body = doc.element.body
sectPr = new_body.find(qn('w:sectPr'))

# detect a List Bullet style if present
bullet_style = None
try:
    _ = doc.styles["List Bullet"]
    bullet_style = "List Bullet"
except KeyError:
    bullet_style = None


def add_par(style=None):
    p_el = OxmlElement('w:p')
    new_body.insert(list(new_body).index(sectPr), p_el)
    p = Paragraph(p_el, doc)
    if style:
        p.style = doc.styles[style]
    else:
        p.style = doc.styles["Normal"]
    return p


def heading(text, level):
    p = add_par(style=("Heading %d" % level))
    p.add_run(text)
    return p


def plain(text):
    p = add_par(style="Normal")
    p.add_run(text)
    return p


def topic_line(topic, said):
    """bold 'Topic:' run + plain spoken line"""
    p = add_par(style="Normal")
    r = p.add_run(topic + ": ")
    r.bold = True
    p.add_run(said)
    return p


def bullet(text):
    p = add_par(style=(bullet_style or "Normal"))
    p.add_run(text)
    return p


# ---- Section content ----
heading("Recruiter Screen Game Plan — Lorena Gallo (Thu Jul 2, 9:00 AM PDT, Zoom)", 2)
plain(
    "Round 1 screen for the Associate Product Manager role, Thu July 2, 2026, 9:00-9:45 AM PDT, on "
    "Zoom (https://yipitdata.zoom.us/j/88196594782), with Lorena Gallo (Recruiting Associate). "
    "Scheduled 30-45 min. YipitData sent a detailed prep brief and this section maps 1:1 to what they "
    "told you they'll cover, so nothing is a surprise. It's a fit + motivation + resume-walkthrough "
    "conversation, not a technical/case grilling, but more substantive than a pure logistics call, so "
    "come with crisp, specific stories ready. Keep answers concise, let Lorena steer, and lead with "
    "the easy wins."
)

heading("Say-it-ready answers (bold topic, then the exact line)", 3)
topic_line(
    "Walk me through your background",
    "\"I'm a Product Manager at Microsoft, where I went from supporting projects to owning a 0-to-1 "
    "data-validation platform and the AI automation on top of it. Before that I was a Product Manager "
    "Intern at Amazon Robotics, mapping dependencies across two thousand-plus robotics units, then two "
    "Product Manager internships at Microsoft driving automation frameworks. The throughline is turning "
    "messy, inconsistent data workflows into structured, scalable systems people rely on — which is "
    "exactly this APM role.\""
)
topic_line(
    "Why YipitData / SpendHound",
    "\"YipitData turns alternative data into intelligence that top funds and Fortune 500s make "
    "high-stakes decisions on — so data quality isn't a feature, the data IS the product. That precision "
    "bar is exactly where I do my best work, and the push into AI-powered, programmatic delivery makes "
    "it the right moment to join.\""
)
topic_line(
    "Why are you leaving Microsoft",
    "\"I've grown a lot and I'm proud of what I built, but I want to go deeper on data products "
    "specifically. At Microsoft the data layer is a means to an end — a way to validate resilience — "
    "rather than the core product. I want to be somewhere that treats data AS the product, where "
    "improving metrics quality or feed reliability has direct, measurable customer impact. That's "
    "YipitData.\""
)
topic_line(
    "What skillset will you bring",
    "\"Metrics and feeds quality — I defined metrics standards and enforced data quality scaling a "
    "recovery-validation platform from a two-person op into a self-service system. Cross-functional "
    "partnering across Data, Engineering, and Operations to drive adoption of scalable standards. "
    "Platform workflows end to end, ingestion through delivery. And real AI fluency — I built an LLM-"
    "powered agent that cut planning cycle time 39% and a RAG search tool that cut lookup time 83%.\""
)
topic_line(
    "SQL / Databricks / PySpark honesty",
    "\"I have solid hands-on data-pipeline and platform experience and named Databricks as an enterprise "
    "partner in my work. SQL, Databricks, and PySpark in an alt-data context are exactly the top growth "
    "area I flagged for myself, and I'm genuinely excited to deepen them here — I learn fast and I'd "
    "rather be honest than overclaim.\""
)
topic_line(
    "Work authorization",
    "\"I'm a US citizen — no sponsorship needed, now or in the future.\""
)

heading("Resume walkthrough — transition points", 3)
bullet(
    "Title framing (important): On the resume YipitData has, your roles are Product Manager / Technical "
    "Product Manager and the internships are Product Manager Intern — NOT 'TPM'. Walk through it "
    "consistently with that product/data-platform framing, leading with the product angle of each role, "
    "not the program-management label."
)
bullet(
    "Amazon Robotics (Product Manager Intern) -> Microsoft internships (x2): Internships gave me "
    "data-driven product reps at scale — mapped dependencies across 2,000+ robotics units, then drove "
    "automation frameworks at Microsoft. WHY: I wanted to build the systems, not just analyze them."
)
bullet(
    "Microsoft internships -> full-time Product Manager: I went from supporting projects to OWNING a "
    "0->1 data-validation platform and the AI automation on top of it. WHY: I add the most value in the "
    "high-ambiguity building phase, turning inconsistent workflows into structured, scalable systems."
)
bullet(
    "Have a clean 60-90 second arc ready and know the one-line WHY for each transition. Be specific — "
    "they explicitly asked you to walk through the transition points in your recent careers."
)

heading("Your skillset -> the JD", 3)
bullet(
    "Metrics + feeds quality: Scaled Azure's recovery-validation platform from a 2-person op into a "
    "self-service system sustaining 45+ annual drills — defining metrics standards and enforcing data "
    "quality across pipeline execution. That's the JD's 'owning metrics publishing quality and "
    "standards' almost verbatim."
)
bullet(
    "Cross-functional partnering: Partnered with Engineering, Data, and Operations to define "
    "requirements, prioritize improvements, and drive adoption of scalable data standards — the JD's "
    "'partnering cross-functionally with Data, Engineering, and Product.'"
)
bullet(
    "Platform workflows, ingestion->delivery: Led the 0->1 Resilience Automation Platform — product "
    "requirements plus self-service scheduling enabling repeatable, standardized data-pipeline "
    "execution — the JD's 'improving and scaling data platform workflows from ingestion through "
    "transformation to delivery.'"
)
bullet(
    "AI fluency (a real differentiator): Built an internal AI agent using LLM-powered automation + "
    "prompt engineering that cut planning cycle time by 39%, and migrated docs to a RAG-backed semantic "
    "search tool cutting lookup time 83%. The JD explicitly wants comfort with LLMs / prompt "
    "engineering / workflow automation — you've LIVED it."
)
bullet(
    "Honest growth edge: The JD wants SQL + Databricks + PySpark. Don't overclaim — you have solid "
    "hands-on data-platform experience and named Databricks as an enterprise partner, and you're "
    "actively excited to deepen SQL/Databricks/PySpark in an alt-data context. Recruiters respect "
    "candor + clear hunger to grow."
)

heading("Behavioral stories to lead with", 3)
bullet(
    "Ownership / 0->1 impact: The Service Healing / recovery-validation platform — scaled from 2 people "
    "into a self-service system, 45+ drills, $14M+ business impact. Best for 'tell me about something "
    "you owned end-to-end.'"
)
bullet(
    "Ambiguity / building structure where none exists: Pioneering Azure's first proactive resilience "
    "testing — rack-level drill program built in 4 months, 94% recovery rate, surfaced latent hardware "
    "defects. The JD literally says 'enjoy solving ambiguous problems and building structure where none "
    "exists.'"
)
bullet(
    "Cross-functional influence / high-stakes coordination: The sovereign-cloud network isolation test "
    "tied to a $1.5B+ contract (bridge lead across orgs) — use for conflict, high-stakes coordination, "
    "exec visibility."
)
bullet(
    "Data-driven decision: The Power BI toil dashboard across 140+ teams that informed the quarterly "
    "roadmap, or the 28%-faster region launches via a unified prioritization framework. Keep each to "
    "~90 seconds with STAR structure."
)

heading("Questions to ask Lorena", 3)
bullet(
    "What does the rest of the interview process look like after today, and what should I focus my prep "
    "on for the next round?"
)
bullet(
    "What does success look like for this APM in the first 6-12 months on the Metrics & Feeds / Public "
    "Investor team?"
)
bullet(
    "Who would I work most closely with day-to-day — how is the team split across Data, Engineering, "
    "and Product?"
)
bullet(
    "How does YipitData's 'impact over tenure' culture actually show up for someone early in their PM "
    "career?"
)
bullet(
    "How is the team using AI tools today on the data-platform side, and where do you see that going?"
)

heading("Logistics & mindset", 3)
bullet(
    "Time management (they flagged it): 30-45 min goes fast. Keep answers concise and specific — land "
    "the point, give one crisp example, stop. Don't ramble through the whole resume; let Lorena steer."
)
bullet(
    "Zoom logistics (from their email): Don't join more than 5 min early (the room may be in use); if "
    "you can't get in, wait 1-3 min before contacting them. Laptop charged/plugged in, good lighting, "
    "eye-level camera, quiet space, pen + paper + water ready."
)
bullet(
    "Work authorization: This role does NOT offer visa sponsorship (stated in the JD), and it's "
    "US-remote (HQ NYC). You're a US citizen, so confirm work authorization cleanly if asked — it's a "
    "clean, no-friction answer."
)
bullet(
    "Tone: Warm, organized, genuinely curious. Lorena is your advocate — make her job easy: be "
    "obviously qualified, clearly interested, concise, and easy to move forward. Thank her by name and "
    "confirm next steps at the end."
)

doc.save(str(GUIDE))
print("YipitData guide built. paras:", len(Document(str(GUIDE)).paragraphs))

# ---- VERIFY ----
v = Document(str(GUIDE))
fulltext = "\n".join(p.text for p in v.paragraphs)

# 0 placeholders
ph_hits = [m for m in ("[COMPANY", "[ROLE", "[Fill") if m in fulltext]
print("PLACEHOLDERS remaining:", ph_hits if ph_hits else "NONE (0)")

# every say-it topic line bold + ends ':'
say_topics = [
    "Walk me through your background",
    "Why YipitData / SpendHound",
    "Why are you leaving Microsoft",
    "What skillset will you bring",
    "SQL / Databricks / PySpark honesty",
    "Work authorization",
]
bad = []
checked = 0
for p in v.paragraphs:
    for t in say_topics:
        if p.text.startswith(t + ":"):
            checked += 1
            r0 = p.runs[0] if p.runs else None
            ok = bool(r0 and r0.bold and r0.text.rstrip().endswith(":"))
            if not ok:
                bad.append((t, (r0.text if r0 else None), (r0.bold if r0 else None)))
print("SAY-IT topic lines checked:", checked, "/", len(say_topics), "| bold+colon failures:", bad if bad else "NONE")

# key facts
facts = {
    "Lorena": "Lorena" in fulltext,
    "$475M": "$475M" in fulltext,
    "data IS the product / data as the product": ("data IS the product" in fulltext or "data as the product" in fulltext),
    "US citizen": "US citizen" in fulltext,
    "Databricks": "Databricks" in fulltext,
}
print("FACTS:", facts)

# headings black (no explicit non-black color override on heading runs)
hdr_colors = []
for p in v.paragraphs:
    if p.style.name.startswith("Heading") or p.style.name == "Title":
        for r in p.runs:
            c = r.font.color
            if c is not None and c.rgb is not None:
                hdr_colors.append((p.style.name, str(c.rgb), p.text[:30]))
print("HEADING explicit run colors (should be none, or 000000):", hdr_colors if hdr_colors else "none/inherit -> BLACK")

# ---- Rebuild zip ----
with zipfile.ZipFile(str(ZIP), "w", zipfile.ZIP_DEFLATED) as z:
    z.write(GUIDE, GUIDE.name)
    if JD.exists():
        z.write(JD, JD.name)
    if RESUME.exists():
        z.write(RESUME, RESUME.name)
with zipfile.ZipFile(str(ZIP)) as z:
    print("ZIP:", z.namelist())
