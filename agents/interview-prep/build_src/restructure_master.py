#!/usr/bin/env python3
"""
In-place restructure of Cyrus's ORIGINAL master guide.
Start from the .bak (his real doc, his real styles) and surgically:
  1. Consolidate Q2/Q3 into Company block + Role block (spoken-ready placeholders).
  2. Replace the 3 prose STAR stories in Section 2 with bulleted study cards.
  3. Leave Q1, Q4, Sections 3/4/5 and ALL styling untouched.
Writes -> templates/Master_Interview_Prep_Guide.docx
"""
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
import copy, glob, os

WS = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
BAK = sorted(glob.glob(str(WS / "templates/Master_Interview_Prep_Guide.docx.bak.20260629_*")))[-1]
OUT = WS / "templates/Master_Interview_Prep_Guide.docx"

shutil.copy2(BAK, OUT)
doc = Document(str(OUT))
body = doc.paragraphs

def find_idx(pred, start=0):
    for i in range(start, len(doc.paragraphs)):
        if pred(doc.paragraphs[i]):
            return i
    return -1

def clear_runs(p):
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    for r in p._p.findall(qn('w:r')):
        p._p.remove(r)

def set_two_runs(p, bold_label, plain_rest):
    """First run bold (label), second run plain."""
    clear_runs(p)
    r1 = p.add_run(bold_label); r1.bold = True
    if plain_rest:
        r2 = p.add_run(plain_rest); r2.bold = False

def set_plain(p, text):
    clear_runs(p)
    p.add_run(text)

def delete_para(p):
    p._p.getparent().remove(p._p)

def insert_par_after(ref_p, text="", style=None, bold_label=None, plain_rest=None):
    """Insert a new paragraph after ref_p, return it."""
    new_p = copy.deepcopy(ref_p._p)
    # strip its runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref_p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, ref_p._parent)
    if style is not None:
        np.style = doc.styles[style]
    if bold_label is not None:
        r1 = np.add_run(bold_label); r1.bold = True
        if plain_rest:
            np.add_run(plain_rest)
    elif text:
        np.add_run(text)
    return np

# ── 1. Q2/Q3 CONSOLIDATION ───────────────────────────────────────────────────
# Find Q2 heading, its "Script:" + fill line; Q3 heading + its lines.
i_q2 = find_idx(lambda p: p.style.name=="Heading 3" and p.text.strip().startswith("Q2:"))
i_q3 = find_idx(lambda p: p.style.name=="Heading 3" and p.text.strip().startswith("Q3:"), i_q2+1)
i_q4 = find_idx(lambda p: p.style.name=="Heading 3" and p.text.strip().startswith("Q4:"), i_q3+1)

# Rewrite Q2 heading -> Company; Q3 heading -> Role
set_plain(doc.paragraphs[i_q2], "Q2: The Company — what it is and why you want to be there")
doc.paragraphs[i_q2].style = doc.styles["Heading 3"]
set_plain(doc.paragraphs[i_q3], "Q3: The Role — what it is and why you want it")
doc.paragraphs[i_q3].style = doc.styles["Heading 3"]

# Between Q2 heading (i_q2) and Q3 heading (i_q3): there are filler paras (Script:, [Fill...], blank).
# Replace that whole block with: helper(italic) + one placeholder paragraph.
# Collect indices strictly between.
def rebuild_block(i_head, i_next_head, helper_text, label, rest):
    # delete everything between head and next head
    victims = [doc.paragraphs[k] for k in range(i_head+1, i_next_head)]
    head_p = doc.paragraphs[i_head]
    for v in victims:
        delete_para(v)
    # re-fetch head, insert helper then placeholder
    h = insert_par_after(head_p, text=helper_text, style="normal")
    for r in h.runs: r.italic = True
    insert_par_after(h, bold_label=label, plain_rest=rest, style="normal")

rebuild_block(i_q2, i_q3,
    "Say it as ONE spoken answer — what the company is, then why you want in. Filled per role.",
    "[COMPANY — what + why]: ",
    "[Fill per role: one-breath spoken answer — what the company is/does, then the specific reason you want to be there. Show basic homework; keep it natural, not a deep-dive recitation.]")

# indices shifted; recompute Q3 / Q4
i_q3 = find_idx(lambda p: p.style.name=="Heading 3" and p.text.strip().startswith("Q3:"))
i_q4 = find_idx(lambda p: p.style.name=="Heading 3" and p.text.strip().startswith("Q4:"), i_q3+1)
rebuild_block(i_q3, i_q4,
    "Say it as ONE spoken answer — what the role is, then why it's the right next step. Filled per role.",
    "[ROLE — what + why]: ",
    "[Fill per role: one-breath spoken answer — what the role actually does day-to-day, then why it's the intersection of technical depth and impact where you do your best work.]")

doc.save(str(OUT))
print("STEP 1 done (Q2/Q3 consolidated). Paras now:", len(doc.paragraphs))
