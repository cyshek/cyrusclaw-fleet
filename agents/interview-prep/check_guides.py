from docx import Document

guides = [
    ("Lance PM", "guides/lance-pm/Lance_PM_Interview_Prep_Guide.docx"),
    ("Lance PM Short", "guides/lance-pm/Lance_PM_Shorter_Interview_Prep_Guide.docx"),
    ("MediaAlpha TPM", "guides/mediaalpha-tpm/MediaAlpha_TPM_Interview_Prep_Guide.docx"),
    ("MediaAlpha TPM Short", "guides/mediaalpha-tpm/MediaAlpha_TPM_Shorter_Interview_Prep_Guide.docx"),
    ("New Relic PM", "guides/new-relic-product-manager-log-management/New_Relic_PM_Interview_Prep_Guide.docx"),
    ("New Relic PM Short", "guides/new-relic-product-manager-log-management/New_Relic_PM_Shorter_Interview_Prep_Guide.docx"),
    ("Scale AI PM-TE", "guides/scale-ai-pm-te/ScaleAI_Interview_Prep_Guide.docx"),
    ("Scale AI PM-TE Short", "guides/scale-ai-pm-te/ScaleAI_Interview_Prep_Guide_Short.docx"),
    ("Scale AI TPM-CV", "guides/scale-ai-tpm-cv/Scale_AI_Interview_Prep_Guide.docx"),
    ("Scale AI TPM-CV Short", "guides/scale-ai-tpm-cv/Scale_AI_Interview_Prep_Guide_Short.docx"),
    ("Podium PM", "guides/podium-pm/Podium_Interview_Prep_Guide.docx"),
    ("Podium PM Short", "guides/podium-pm/Podium_Interview_Prep_Guide_Short.docx"),
    ("Mintlify SE", "guides/mintlify-solutions-engineer-post-sales/Mintlify_SE_PostSales_Interview_Prep_Guide.docx"),
    ("Datadog PTSE", "guides/datadog-partner-technology-solutions-engineer/Datadog_PTSE_Interview_Prep_Guide.docx"),
]

for name, path in guides:
    try:
        doc = Document(path)
        headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith('Heading 2')]
        has_s4 = any('Section 4' in s or 'Product Thinking' in s for s in headings)
        has_s5 = any('Section 5' in s or 'Questions for the Interviewer' in s for s in headings)
        placeholders = [p.text.strip() for p in doc.paragraphs if 'Fill in here' in p.text]
        print(f"{name}: S4={'yes' if has_s4 else 'NO'} S5={'yes' if has_s5 else 'NO'} placeholders={len(placeholders)}")
    except Exception as err:
        print(f"{name}: ERROR - {err}")
