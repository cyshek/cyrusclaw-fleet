#!/usr/bin/env python3
import zipfile, shutil, re, sys
from pathlib import Path

WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATES = WORKSPACE / "templates"
GUIDES = WORKSPACE / "guides"
SUBMITTED = Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted")
MASTER_FULL = TEMPLATES / "Master_Interview_Prep_Guide.docx"

LANGCHAIN_Q2 = [
    ("Why do I want to work at LangChain:", "LangChain is right at the center of the shift everyone is building toward — making AI agents actually work in production, not just in demos. Their tools went from beloved open source to a real enterprise platform, and with the Series B momentum and Fortune 500 adoption, the work has genuine reach. I want to be at a company where the technology I'm helping customers adopt is the thing reshaping how software gets built."),
    ("Why do I want this role:", "A Sales Engineer here gets to live at the intersection of deep technical work and real customer impact — exactly where I do my best work. I've built AI agents with Copilot, driven technical adoption across cross-functional teams, and I like being the person who can both get into the weeds with engineers and clearly explain the value to a stakeholder. This role is about turning prospects into successful production users, and that translation work is what I enjoy most."),
]

LANGCHAIN_Q3 = [
    ("What does LangChain do:", "LangChain builds the foundation for agent engineering. They offer the widely-adopted open-source frameworks (LangChain and LangGraph) plus LangSmith, a platform for building, evaluating, deploying, and operating LLM agents at scale. With 100M+ monthly downloads and 35% of the Fortune 500 using LangSmith, they've become core infrastructure for teams shipping real AI products."),
    ("What is this role:", "The Sales Engineer role is a technical, customer-facing seat across the sales cycle — partnering with prospects and new enterprise customers to onboard them, run technical workshops, build enablement materials, and feed real-world friction back to product. It needs hands-on fluency with Python and the modern LLM/agent stack plus strong communication to explain complex concepts to everyone from individual developers to enterprise stakeholders."),
]

MEDIAALPHA_Q2 = [
    ("Why do I want to work at MediaAlpha:", "MediaAlpha is intentionally a small, high-leverage organization — the kind of place where an individual can have outsized impact, which is exactly the environment I want. It's a tech- and data-science-driven company, and the TPM role sits right at the nexus of engineering, business, and external partners. That cross-functional center-of-gravity is where I've consistently done my strongest work."),
    ("Why do I want this role:", "This role is about owning an end-to-end program across cross-functional teams — defining success metrics, managing the backlog, and driving execution — which mirrors how I've run the resilience program at Microsoft. I like that they explicitly want someone comfortable reviewing code with engineers and talking strategy with executive leadership; being fluent on both sides of that line is one of my strengths, and this role rewards it."),
]

MEDIAALPHA_Q3 = [
    ("What does MediaAlpha do:", "MediaAlpha is a customer-acquisition solutions provider powered by technology and data science. They build the platforms that connect consumers with providers in high-consideration insurance categories — property & casualty, health, and life insurance — running systems that handle hundreds of thousands of high-value transactions per day."),
    ("What is this role:", "The Technical Program Manager owns a functional area end-to-end: defining metrics to measure success, monitoring issues, and surfacing opportunities. Day to day that means managing backlogs and Kanban workflows, coordinating across engineering and business teams, and working with external partners to define, build, test, and release API integrations. It calls for SQL/JSON/Web API familiarity and solid Agile/SDLC experience."),
]

def escape_xml(text):
    return (text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                .replace('"',"&quot;").replace("'","&apos;"))

def extract_pPr(para_xml):
    m = re.search(r"<w:pPr>.*?</w:pPr>", para_xml, re.DOTALL)
    return m.group(0) if m else "<w:pPr/>"

def make_bullet_xml(pPr_xml, bold_label, plain_text):
    r1 = '<w:r><w:rPr><w:b w:val="1"/><w:bCs w:val="1"/><w:rtl w:val="0"/></w:rPr><w:t xml:space="preserve">' + escape_xml(bold_label) + ' </w:t></w:r>'
    r2 = '<w:r><w:rPr><w:rtl w:val="0"/></w:rPr><w:t xml:space="preserve">' + escape_xml(plain_text) + '</w:t></w:r>'
    return "<w:p>" + pPr_xml + r1 + r2 + "</w:p>"

def patch_docx(src, dst, q2_bullets, q3_bullets):
    shutil.copy2(src, dst)
    with zipfile.ZipFile(dst, "r") as zin:
        content = zin.read("word/document.xml").decode("utf-8")
        all_files = {name: zin.read(name) for name in zin.namelist()}
    placeholders = []
    idx = 0
    while True:
        pos = content.find("Fill in here", idx)
        if pos == -1:
            break
        p_start = content.rfind("<w:p ", 0, pos)
        p_end = content.find("</w:p>", pos) + len("</w:p>")
        placeholders.append((p_start, p_end, content[p_start:p_end]))
        idx = pos + 1
    if len(placeholders) != 2:
        raise ValueError("Expected 2 placeholders, found " + str(len(placeholders)) + " in " + str(dst))
    bullets_map = {0: q2_bullets, 1: q3_bullets}
    replacements = []
    for i, (p_start, p_end, para_xml) in enumerate(placeholders):
        pPr = extract_pPr(para_xml)
        new_paras = "".join(make_bullet_xml(pPr, label, text) for label, text in bullets_map[i])
        replacements.append((p_start, p_end, new_paras))
    for p_start, p_end, new_paras in reversed(replacements):
        content = content[:p_start] + new_paras + content[p_end:]
    all_files["word/document.xml"] = content.encode("utf-8")
    dst.unlink()
    with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in all_files.items():
            zout.writestr(name, data)
    print("  Patched: " + dst.name)

def verify_docx(path):
    try:
        import docx
        doc = docx.Document(str(path))
    except Exception as e:
        print("  ERROR: python-docx failed on " + path.name + ": " + str(e))
        return False
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "Fill in here" in full_text:
        print("  ERROR: [Fill in here] still present in " + path.name)
        return False
    bold_bullets = []
    for p in doc.paragraphs:
        if p.runs and p.runs[0].bold and ":" in p.runs[0].text:
            bold_bullets.append(p.runs[0].text)
    if len(bold_bullets) < 4:
        print("  ERROR: Expected 4+ bold-label bullets, found " + str(len(bold_bullets)) + " in " + path.name)
        return False
    if "Q1" not in full_text:
        print("  ERROR: Original content Q1 missing from " + path.name)
        return False
    print("  PASS: " + path.name + " - no placeholders, " + str(len(bold_bullets)) + " bold bullets, body intact")
    for b in bold_bullets:
        print("    label: " + b)
    return True

def build_zip(zip_path, files):
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files:
            f = Path(f)
            if f.exists():
                z.write(f, f.name)
                print("  Added to zip: " + f.name)
            else:
                print("  WARNING: File not found: " + str(f))
    print("  Zip created: " + str(zip_path))

def build_bundle(name, folder, zip_path, full_name, q2_bullets, q3_bullets, jd_path, resume_path):
    print("\n" + "="*60)
    print("Building bundle: " + name)
    print("="*60)
    folder.mkdir(parents=True, exist_ok=True)
    full_dst = folder / full_name
    print("\n[1/3] Patching full guide...")
    patch_docx(MASTER_FULL, full_dst, q2_bullets, q3_bullets)
    print("\n[2/3] Verifying doc...")
    ok = verify_docx(full_dst)
    if not ok:
        print("  VERIFICATION FAILED for " + name)
        return False
    print("\n[3/3] Building zip...")
    build_zip(zip_path, [full_dst, jd_path, resume_path])
    print("\nBundle complete: " + str(zip_path))
    return True

if __name__ == "__main__":
    success = True
    lc_ok = build_bundle(
        name="LangChain Sales Engineer",
        folder=GUIDES / "langchain-sales-engineer",
        zip_path=GUIDES / "langchain-sales-engineer.zip",
        full_name="LangChain_SalesEngineer_Interview_Prep_Guide.docx",
        q2_bullets=LANGCHAIN_Q2,
        q3_bullets=LANGCHAIN_Q3,
        jd_path=SUBMITTED / "langchain-b47dbcea-222c-427d-bf4d-9e30a6b60d4a/JD.md",
        resume_path=SUBMITTED / "langchain-b47dbcea-222c-427d-bf4d-9e30a6b60d4a/Cyrus_Shekari_Resume_ashby-langchain_b47dbcea_v2.pdf",
    )
    success = success and lc_ok
    ma_ok = build_bundle(
        name="MediaAlpha TPM",
        folder=GUIDES / "mediaalpha-tpm",
        zip_path=GUIDES / "mediaalpha-tpm.zip",
        full_name="MediaAlpha_TPM_Interview_Prep_Guide.docx",
        q2_bullets=MEDIAALPHA_Q2,
        q3_bullets=MEDIAALPHA_Q3,
        jd_path=SUBMITTED / "mediaalpha-8525774002/JD.md",
        resume_path=SUBMITTED / "mediaalpha-8525774002/Cyrus_Shekari_Resume_mediaalpha_8525774002_v2.pdf",
    )
    success = success and ma_ok
    print("\n" + "="*60)
    print("FINAL STATUS: " + ("ALL BUNDLES OK" if success else "SOME BUNDLES FAILED"))
    print("="*60)
    sys.exit(0 if success else 1)
