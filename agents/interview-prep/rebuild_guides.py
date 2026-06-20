"""
Rebuild all per-company guides from the new master template,
injecting back the existing Q2/Q3 fills.
"""
import copy
import shutil
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

WORKSPACE = "/home/azureuser/.openclaw/agents/interview-prep/workspace"
MASTER = f"{WORKSPACE}/templates/Master_Interview_Prep_Guide.docx"

# Q2/Q3 fills per company: list of (label_bold, rest_of_text) per bullet
# Order: [q2_bullet1, q2_bullet2, q3_bullet1, q3_bullet2]
FILLS = {
    "lance-pm": {
        "full": "guides/lance-pm/Lance_PM_Interview_Prep_Guide.docx",
        "title": "Lance PM Interview Prep Guide",
        "q2": [
            ("Why do I want to work at Lance:", " Lance is building autonomous AI agents that handle real hotel operations end-to-end — no scripts, no templates, just autonomous revenue generation. I find that mission genuinely compelling; it's infrastructure for the next wave of labor."),
            ("Why I want this role:", " This PM role maps directly to what I have been doing at Microsoft — owning a roadmap, coordinating across engineering and ops, and driving adoption of a new platform capability from scratch."),
        ],
        "q3": [
            ("What Lance does:", " Lance builds autonomous AI agents for the hospitality industry that handle calls, close sales, and manage reservations end-to-end, replacing rule-based chatbots with agents that actually complete transactions."),
            ("What this role is:", " The Product Manager owns the roadmap and lifecycle for assigned products, working across engineering, design, and go-to-market to define what gets built, prioritize the backlog, and drive outcomes for hotel operators."),
        ],
    },
    "mediaalpha-tpm": {
        "full": "guides/mediaalpha-tpm/MediaAlpha_TPM_Interview_Prep_Guide.docx",
        "title": "MediaAlpha TPM Interview Prep Guide",
        "q2": [
            ("Why do I want to work at MediaAlpha: ", "MediaAlpha sits at an intersection I find genuinely interesting — using technology and data science to solve real customer acquisition problems in insurance, a high-stakes, complex vertical where precision matters."),
            ("Why do I want this role: ", "This TPM role maps almost directly to what I've been doing. At Microsoft I owned an end-to-end automation platform initiative, coordinating across engineering, product, and operations to deliver a scalable, self-service system."),
        ],
        "q3": [
            ("What does MediaAlpha do: ", "MediaAlpha is a customer acquisition solutions provider that combines technology and data science to help insurance carriers and distributors efficiently connect with high-intent consumers at the right moment in their buying journey."),
            ("What is this role: ", "The Technical Program Manager role is the connective tissue between MediaAlpha's engineering teams, internal stakeholders, and business outcomes — owning program planning, cross-functional coordination, and delivery execution."),
        ],
    },
    "new-relic-product-manager-log-management": {
        "full": "guides/new-relic-product-manager-log-management/New_Relic_PM_Interview_Prep_Guide.docx",
        "title": "New Relic PM Interview Prep Guide",
        "q2": [
            ("Why do I want to work at New Relic: ", "New Relic is betting the whole platform on intelligent, AI-driven observability with a usage-based model that makes the economics cleaner for customers — that strategic shift is interesting, and it puts Log Management in a pivotal position."),
            ("Why do I want this role: ", "This is a full-business-owner PM role on a mature, business-critical product that still needs bold strategy — exactly the kind of challenge I'm looking for after leading large-scale platform initiatives at Microsoft."),
        ],
        "q3": [
            ("What does New Relic do: ", "New Relic is one of the leading observability platforms, and they've repositioned around an AI-driven, usage-based model — unified telemetry (logs, metrics, traces, events) in a single platform with an in-product AI assistant."),
            ("What is this role: ", "The Product Manager for Log Management is the full business owner for the logging side of the APM offering — strategy, roadmap, pricing, and cross-functional execution for one of New Relic's most-used and most-monetized capabilities."),
        ],
    },
    "scale-ai-pm-te": {
        "full": "guides/scale-ai-pm-te/ScaleAI_Interview_Prep_Guide.docx",
        "title": "Scale AI PM (T&E) Interview Prep Guide",
        "q2": [
            ("Why do I want to work at Scale AI: ", "Scale AI is doing foundational work — not building another AI application, but building the data and evaluation infrastructure that every serious AI program depends on. That's a durable position, and the public sector angle makes it even more mission-critical."),
            ("Why do I want this role: ", "The Product Manager, Public Sector GenAI T&E role combines two things I find compelling: deeply technical product ownership and high-stakes stakeholder environments where the cost of getting it wrong is real."),
        ],
        "q3": [
            ("What does Scale AI do: ", "Scale AI provides the data infrastructure and tooling that AI companies and enterprises need to build, fine-tune, and evaluate large language models — data labeling, RLHF pipelines, and now evaluation frameworks for government and enterprise."),
            ("What is this role: ", "The PM for Public Sector GenAI T&E owns the roadmap for Scale's evaluation capabilities in government and defense contexts — working at the intersection of AI safety, compliance, and mission-critical deployment."),
        ],
    },
    "scale-ai-tpm-cv": {
        "full": "guides/scale-ai-tpm-cv/Scale_AI_Interview_Prep_Guide.docx",
        "title": "Scale AI TPM (CV) Interview Prep Guide",
        "q2": [
            ("Why do I want to work at Scale AI:", " Scale AI is doing foundational work in the AI space - not just building models but making them actually reliable through data quality, evaluation, and RLHF infrastructure. That's the layer I find most compelling."),
            ("Why do I want this role:", " The TPM role in Computer Vision for Public Sector is a strong match for my background. I've spent years at Microsoft coordinating across engineering, operations, and senior stakeholders on high-stakes, zero-downtime programs."),
        ],
        "q3": [
            ("What does Scale AI do:", " Scale AI builds the data infrastructure that powers AI development - data labeling, RLHF pipelines, and evaluation frameworks that help enterprises and government agencies build reliable, production-grade AI systems."),
            ("What is this role:", " This is a TPM position on Scale's public sector delivery team, focused on computer vision. The role owns program planning, cross-functional execution, and delivery coordination for CV-based AI projects in government and defense contexts."),
        ],
    },
    "podium-pm": {
        "full": "guides/podium-pm/Podium_Interview_Prep_Guide.docx",
        "title": "Podium PM Interview Prep Guide",
        "q2": [
            ("Why do I want to work at Podium: ", "Podium is doing something genuinely exciting — bringing AI-powered communication and payments to local businesses that have historically been underserved by enterprise software. The mission is concrete and the market is huge."),
            ("Why do I want this role: ", "This PM role sits at the intersection of two things I care about: owning real product strategy end-to-end and working in a fast-moving environment where the roadmap actually changes customer outcomes."),
        ],
        "q3": [
            ("What does Podium do: ", "Podium is an AI-powered platform for local and SMB businesses — it brings together customer messaging, reviews, payments, and marketing automation in a single platform, helping businesses communicate and transact with customers more efficiently."),
            ("What is this role: ", "The Product Manager role at Podium owns a core area of the platform end-to-end — strategy, roadmap, and execution — working across engineering, design, and GTM to ship features that drive real business outcomes for local businesses."),
        ],
    },
    "mintlify-solutions-engineer-post-sales": {
        "full": "guides/mintlify-solutions-engineer-post-sales/Mintlify_SE_PostSales_Interview_Prep_Guide.docx",
        "title": "Mintlify SE Post-Sales Interview Prep Guide",
        "q2": [
            ("Why do I want to work at Mintlify: ", "You're the documentation platform for 100M+ developers and 20,000+ companies — Anthropic, Cursor, Perplexity, OpenAI. That's the infrastructure layer of the AI developer ecosystem, and it's a space I find genuinely compelling."),
            ("Why do I want this role: ", "Post-Sales SE is the technical quarterback role I'm genuinely best at — owning high-volume migrations, embedding deeply with engineering teams, and turning complex technical requirements into shipped outcomes."),
        ],
        "q3": [
            ("What does Mintlify do: ", "Mintlify is the modern documentation platform that lets companies build and ship beautiful, high-quality developer docs — with AI-powered search, analytics, and customization tools that make docs a product, not an afterthought."),
            ("What is this role: ", "The Post-Sales Solutions Engineer is the technical quarterback for scaled customers — owning 6+ migrations simultaneously, serving as the primary technical contact, and building the systems (templates, tools, playbooks) that make Mintlify's onboarding scale."),
        ],
    },
    "datadog-partner-technology-solutions-engineer": {
        "full": "guides/datadog-partner-technology-solutions-engineer/Datadog_PTSE_Interview_Prep_Guide.docx",
        "title": "Datadog Partner TSE Interview Prep Guide",
        "q2": [
            ("Why do I want to work at Datadog: ", "You've become the observability standard for the AI era, and that's a category I already live in. Monitoring distributed systems at Microsoft-scale is a significant part of my day job, and Datadog is the platform I'd point to as best-in-class."),
            ("Why do I want this role: ", "The Partner TSE sits at the exact intersection of what I'm best at — deep distributed-systems knowledge, cross-functional coordination, and technical enablement at scale — applied to the partner ecosystem rather than internal stakeholders."),
        ],
        "q3": [
            ("What does Datadog do: ", "Datadog is the leading observability and security platform for the AI era — unified, end-to-end visibility across infrastructure, applications, logs, and security for cloud-native companies and enterprises running at scale."),
            ("What is this role: ", "The Partner TSE is the technical bridge between Datadog and your third-party developer community — enabling partners to build reliable integrations, providing deep technical guidance, and driving the quality bar for the partner ecosystem."),
        ],
    },
}

def find_q_paragraphs(doc):
    """Return indices of Q2 and Q3 'Script:' paragraphs and their fill paragraphs."""
    paras = doc.paragraphs
    q2_script_idx = None
    q3_script_idx = None
    in_q2 = False
    in_q3 = False
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if 'Q2:' in txt and 'Why do you want' in txt:
            in_q2 = True
            in_q3 = False
        elif 'Q3:' in txt and ('Do you know' in txt or 'what this role' in txt.lower()):
            in_q2 = False
            in_q3 = True
        elif p.style.name.startswith('Heading') and ('Q4' in txt or 'Q5' in txt or 'Section' in txt):
            in_q2 = False
            in_q3 = False

        if txt == 'Script:':
            if in_q2 and q2_script_idx is None:
                q2_script_idx = i
            elif in_q3 and q3_script_idx is None:
                q3_script_idx = i

    return q2_script_idx, q3_script_idx


def set_paragraph_fill(para, label_text, rest_text):
    """Set a paragraph's text with bold label run and plain rest run."""
    # Clear existing runs
    for r in para.runs:
        r.text = ''
    # Remove all existing runs from XML
    p_elem = para._p
    for r_elem in p_elem.findall(qn('w:r')):
        p_elem.remove(r_elem)

    # Add bold label run
    run1 = para.add_run(label_text)
    run1.bold = True

    # Add plain rest run
    run2 = para.add_run(rest_text)
    run2.bold = False


def rebuild_guide(company_key, data):
    path = f"{WORKSPACE}/{data['full']}"
    doc = Document(MASTER)

    q2_script_idx, q3_script_idx = find_q_paragraphs(doc)
    paras = doc.paragraphs

    if q2_script_idx is None or q3_script_idx is None:
        print(f"  ERROR: Could not find Q2/Q3 script paragraphs!")
        return False

    # The fill paragraphs are immediately after each Script: paragraph
    q2_fill_idx = q2_script_idx + 1
    q3_fill_idx = q3_script_idx + 1

    # Q2: fill bullet 1 (in place) and insert bullet 2 after
    q2_para1 = paras[q2_fill_idx]
    label1, rest1 = data['q2'][0]
    set_paragraph_fill(q2_para1, label1, rest1)

    # Insert Q2 bullet 2 after bullet 1
    label2, rest2 = data['q2'][1]
    new_q2_p2 = copy.deepcopy(q2_para1._p)
    q2_para1._p.addnext(new_q2_p2)
    # Now set it (it's inserted after, find it)
    inserted_q2 = q2_para1._p.getnext()
    # Clear and set runs
    for r_elem in inserted_q2.findall(qn('w:r')):
        inserted_q2.remove(r_elem)
    # We need a paragraph object to use add_run - work directly on XML
    from docx.oxml import OxmlElement
    def add_run_to_p(p_xml, text, bold=False):
        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rpr.append(b)
            b2 = OxmlElement('w:bCs')
            rpr.append(b2)
        r.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p_xml.append(r)

    add_run_to_p(inserted_q2, label2, bold=True)
    add_run_to_p(inserted_q2, rest2, bold=False)

    # Re-fetch paragraphs (indices shifted after insert)
    paras2 = doc.paragraphs
    # Find q3 script again
    _, q3_script_idx2 = find_q_paragraphs(doc)
    q3_fill_idx2 = q3_script_idx2 + 1

    q3_para1 = paras2[q3_fill_idx2]
    label3, rest3 = data['q3'][0]
    set_paragraph_fill(q3_para1, label3, rest3)

    # Insert Q3 bullet 2
    label4, rest4 = data['q3'][1]
    new_q3_p2 = copy.deepcopy(q3_para1._p)
    q3_para1._p.addnext(new_q3_p2)
    inserted_q3 = q3_para1._p.getnext()
    for r_elem in inserted_q3.findall(qn('w:r')):
        inserted_q3.remove(r_elem)
    add_run_to_p(inserted_q3, label4, bold=True)
    add_run_to_p(inserted_q3, rest4, bold=False)

    doc.save(path)
    return True


def verify_guide(path):
    doc = Document(path)
    placeholders = [p.text for p in doc.paragraphs if 'Fill in here' in p.text]
    # Check bold labels
    bold_ok = True
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt or 'Fill in here' in txt:
            continue
        # Q2/Q3 fill bullets have bold first run
        if len(p.runs) >= 2 and p.runs[0].bold and ':' in p.runs[0].text:
            pass  # good
    return len(placeholders) == 0


for company_key, data in FILLS.items():
    path = f"{WORKSPACE}/{data['full']}"
    print(f"\nRebuilding: {data['title']}")
    try:
        ok = rebuild_guide(company_key, data)
        if ok:
            good = verify_guide(path)
            print(f"  {'OK - no placeholders' if good else 'WARNING - placeholders remain'}")
    except Exception as err:
        import traceback
        print(f"  ERROR: {err}")
        traceback.print_exc()

print("\nDone.")
