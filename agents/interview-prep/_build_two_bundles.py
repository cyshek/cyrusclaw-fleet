#!/usr/bin/env python3
"""
Build TWO interview bundles (lean scope, post-2026-06-09):
  - Podium  — Product Manager
  - Mintlify — Solutions Engineer, Post-Sales

Method = the New Relic lean pattern:
  - LITERAL copy of templates/Master_Interview_Prep_Guide.docx (keeps Arial/Georgia,
    Heading1/2/3, 'normal', numbering defs).
  - Replace ONLY the Q2 and Q3 `[Fill in here]` bullets, cloning the master's own
    placeholder bullet paragraph so the new bullets inherit identical font + numbering.
  - NO Section 5. Nothing else in the doc is touched (all 7 STAR stories, Q1, Q4 intact).
Only writes the per-company copy. Master is never modified.
"""
import copy, shutil, os
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

MASTER = "templates/Master_Interview_Prep_Guide.docx"

PODIUM_OUT = "guides/podium-product-manager/Podium_PM_Interview_Prep_Guide.docx"
MINT_OUT   = "guides/mintlify-solutions-engineer-post-sales/Mintlify_SE_PostSales_Interview_Prep_Guide.docx"

# ---------------- Podium PM ----------------
PODIUM_Q2 = [
    "Why do I want to work at Podium: You crossed $100M in AI-agent ARR in under 24 months and deployed 10,000 AI Employees that do real work for 60,000+ local businesses — texting customers, qualifying leads, booking appointments, closing deals. That's not a chatbot demo; that's human-in-the-loop AI agents operating in production at scale, which is exactly the problem space I've spent the last two years in. I've been building and operating AI agents on the infrastructure side at Microsoft, and Podium is one of the very few places turning that into real revenue for real customers right now. I want to be where AI agents are shipping outcomes, not slideware.",
    "Why do I want this role: The thing that genuinely sold me is that at Podium, PMs are builders — you use Claude, Cursor, and LLMs to prototype UIs, write code, and ship alongside engineers. That's how I already work. I built AI agents with GitHub Copilot to automate an entire drill lifecycle and prototyped a production UI with AI tooling instead of waiting on a design queue. Most PM roles still stop at the spec; this one rewards exactly the 0-to-1, get-into-the-weeds, ship-and-learn loop I'm best at, on a product I'd actually love using.",
]
PODIUM_Q3 = [
    "What does Podium do: Podium brings AI Employees to local businesses — Auto, Home Services, Aesthetics — that turn every customer conversation into revenue. The platform captures and converts leads 24/7 over text and messaging: AI agents that integrate into a business's systems, make decisions, and operate alongside the human staff to drive both new and repeat business. It started as the easiest way for a local shop to get reviews over text and has grown into a multi-product, consumer-grade platform for how local commerce actually gets done.",
    "What is this role: It's a builder-PM role that owns both the product customers use today and the platform Podium is building for tomorrow — defining strategy end-to-end, then actually prototyping and shipping it with AI tools alongside engineers. It's heavy on living in customer workflows to find what's missing, shipping fast and iterating on real usage data, and making company-level architecture and integration decisions plus designing systems and patterns other teams build on. In short: high agency, 0-to-1, customer-obsessed, and technical enough to build — which is the exact intersection I want to be operating in.",
]

# ---------------- Mintlify SE Post-Sales ----------------
MINT_Q2 = [
    "Why do I want to work at Mintlify: You're the documentation platform for 100M+ developers and 20,000+ companies — Anthropic, Microsoft, PayPal, Spotify, Coinbase — built by a ~50-person team that just raised a $45M Series B. I love that this role is explicitly about turning one-off custom work into repeatable, productized capability, and that your culture prizes \"slope over y-intercept\": learning velocity and grit over pedigree. That's exactly how I operate — I take an ambiguous, manual problem and platformize it. Being one of a small number of people who directly shapes the trajectory of a company with that kind of developer reach is the environment I want.",
    "Why do I want this role: Post-Sales SE is the technical quarterback role I'm genuinely best at — owning high-volume migrations end-to-end, coordinating an offshore vendor team against real SLAs, and building the customer-health and operational infrastructure from 0 to 1. At Microsoft I was the technical quarterback on 14+ cross-functional executions, I drove vendor and partner-team accountability, and I built self-service operational systems and dashboards from scratch to kill manual toil. This role lets me do all of that hands-on — debugging, writing scripts, shipping the process — instead of just managing it from a distance, and that's the work that energizes me.",
]
MINT_Q3 = [
    "What does Mintlify do: Mintlify is the modern documentation platform that lets companies build and ship beautiful, high-performing developer docs straight from their codebase — powering documentation for 20,000+ companies and reaching over 100 million developers a year. Docs are written as code (Markdown/MDX, Git-based), so they stay in sync with the product, and the platform handles the heavy lifting of migrations, hosting, and a great developer-and-reader experience. It's the docs layer behind a huge share of the AI and developer-tools world, including names like Anthropic and over 20% of the last YC batch.",
    "What is this role: The Post-Sales Solutions Engineer is the technical quarterback for scaled customers — owning 6+ migrations a week through 10-12 day go-live cycles, coordinating ~60% of the time with offshore vendor partners on assignment and QA, and dedicating ~40% to high-potential accounts via dedicated Slack channels for troubleshooting, onboarding, and feature triage. It's genuinely hands-on (debugging migrations, writing custom scripts, building solutions when complexity demands it) and it's also 0-to-1 systems-building: standing up CRM tracking, customer-health-score models, and operational dashboards, plus a repeatable weekly feature-communication system. It blends high-volume execution, vendor management, data-driven insight, and process-building — which maps almost one-to-one to what I've done.",
]

W = qn  # alias not needed; keep import tidy

def fill(out_path, q2_bullets, q3_bullets):
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    tmp = "/tmp/_mcopy_build.docx"
    shutil.copyfile(MASTER, tmp)
    doc = Document(tmp)

    def find_idx(contains_low, start=0):
        for i in range(start, len(doc.paragraphs)):
            if contains_low in doc.paragraphs[i].text.lower():
                return i
        return None

    def set_text(p, text):
        # Clear existing runs, then re-add text. If the bullet leads with an
        # inline label ending in a colon (e.g. "Why do I want this role: ..."),
        # split that label into its OWN BOLD run so Q2/Q3 visually separate the
        # two bundled questions. This matches the New Relic / Datadog format.
        # (Regression guard: a prior version emitted a single flat run, which
        # dropped the bold label on the Mintlify + Podium 06-10 builds.)
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        label_prefixes = (
            "Why do I want to work at ",
            "Why do I want this role:",
            "What does ",
            "What is this role:",
        )
        is_labeled = any(text.lstrip().startswith(pfx) for pfx in label_prefixes)
        if is_labeled and ":" in text:
            idx = text.index(":")
            label = text[: idx + 1]
            rest = text[idx + 1 :]
            if rest.startswith(" "):
                label = label + " "
                rest = rest[1:]
            run_label = p.add_run(label)
            run_label.bold = True
            if rest:
                run_rest = p.add_run(rest)
                run_rest.bold = False
        else:
            p.add_run(text)

    def replace_block(head_low, stop_low, bullets):
        qi = find_idx(head_low)
        assert qi is not None, f"head not found: {head_low}"
        ph = None
        for j in range(qi+1, len(doc.paragraphs)):
            if "fill in here" in doc.paragraphs[j].text.lower():
                ph = j; break
            if stop_low in doc.paragraphs[j].text.lower():
                break
        assert ph is not None, f"placeholder not found after {head_low}"
        ph_para = doc.paragraphs[ph]
        set_text(ph_para, bullets[0])
        anchor = ph_para._p
        for btxt in bullets[1:]:
            newp = copy.deepcopy(ph_para._p)
            anchor.addnext(newp)
            anchor = newp
            cp = Paragraph(newp, ph_para._parent)
            set_text(cp, btxt)
        print(f"  filled [{head_low[:30]}]: {len(bullets)} bullets")

    replace_block("why do you want to work at our company", "do you know what this role is", q2_bullets)
    replace_block("do you know what this role is", "why are you looking to leave", q3_bullets)

    doc.save(out_path)
    print("SAVED", out_path)

if __name__ == "__main__":
    fill(PODIUM_OUT, PODIUM_Q2, PODIUM_Q3)
    fill(MINT_OUT,   MINT_Q2,   MINT_Q3)
