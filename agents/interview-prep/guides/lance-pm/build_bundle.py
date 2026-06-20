#!/usr/bin/env python3
"""Build Lance PM interview bundle."""

import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

# Paths
BUNDLE_DIR = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/lance-pm")
BUNDLE_DIR.mkdir(parents=True, exist_ok=True)

SOURCES = {
    "master": Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Master_Interview_Prep_Guide.docx"),
    "shorter": Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Shorter_Master_Interview_Prep_Guide.docx"),
    "jd": Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/lance-294e2602-725c-4d84-a7a0-a1fc602acec8/JD.md"),
    "resume": Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/lance-294e2602-725c-4d84-a7a0-a1fc602acec8/Cyrus_Shekari_Resume_ashby-lance_294e2602_v2.pdf"),
}

DEST_NAMES = {
    "master": "Lance_PM_Interview_Prep_Guide.docx",
    "shorter": "Lance_PM_Shorter_Interview_Prep_Guide.docx",
    "jd": "Lance_PM_JD.md",
    "resume": "Cyrus_Shekari_Resume_Lance_PM.pdf",
}

# Q2 fills
Q2_FILLS = [
    {
        "bold_label": "Why do I want to work at Lance:",
        "plain_text": "Lance is building autonomous AI agents that handle real hotel operations end-to-end — not just a chatbot layer, but agents that take actions inside live hotel software. The combination of Computer Use Agent technology and a focused vertical is a genuinely interesting product problem, and the traction with 50+ hotels across Marriott, Hilton, and Hyatt shows it is working.",
    },
    {
        "bold_label": "Why I want this role:",
        "plain_text": "This PM role maps directly to what I have been doing at Microsoft — owning a roadmap, coordinating cross-functional teams, and driving a product from ambiguous early-stage through scaled execution. I am drawn to the pace of a YC company where decisions move fast and the PM has real influence over what gets built.",
    },
]

# Q3 fills
Q3_FILLS = [
    {
        "bold_label": "What Lance does:",
        "plain_text": "Lance builds autonomous AI agents for the hospitality industry that handle calls, close sales, and manage hotel operations by integrating directly into existing hotel software to make real-time decisions. Backed by YC W26, they work with 50+ hotels across Marriott, Hilton, and Hyatt.",
    },
    {
        "bold_label": "What this role is:",
        "plain_text": "The Product Manager owns the roadmap and lifecycle for assigned products, working across engineering, design, and business teams to define requirements, prioritize features, and ship. It is an in-person role in San Francisco focused on AI and SaaS product delivery.",
    },
]


def fill_placeholders(doc_path: Path) -> None:
    """Fill Q2 and Q3 placeholders in the given docx file."""
    doc = Document(doc_path)
    
    in_q2 = False
    in_q3 = False
    q2_count = 0
    q3_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        
        # Detect section headers
        if "Q2:" in text and ("Why" in text or "here" in text or "company" in text):
            in_q2 = True
            in_q3 = False
            continue
        
        if "Q3:" in text and ("do we do" in text or "What" in text or "know" in text):
            in_q3 = True
            in_q2 = False
            continue
        
        if "Q4:" in text:
            in_q2 = False
            in_q3 = False
            continue
        
        # Fill placeholders
        if "[Fill in here]" in text:
            if in_q2 and q2_count < len(Q2_FILLS):
                fill = Q2_FILLS[q2_count]
                q2_count += 1
                _replace_paragraph(para, fill["bold_label"], fill["plain_text"])
            elif in_q3 and q3_count < len(Q3_FILLS):
                fill = Q3_FILLS[q3_count]
                q3_count += 1
                _replace_paragraph(para, fill["bold_label"], fill["plain_text"])
    
    doc.save(doc_path)
    print(f"Saved: {doc_path.name} | Q2 filled: {q2_count} | Q3 filled: {q3_count}")


def _replace_paragraph(para, bold_label: str, plain_text: str) -> None:
    """Replace a paragraph's runs with bold label + plain text, keeping paragraph formatting."""
    # Clear existing runs without touching paragraph-level XML (style, numPr, etc.)
    for run in para.runs:
        run.text = ""
    
    # If there are no runs, we need to add them; if there are, reuse first two
    runs = para.runs
    
    # Add or update run 0 - bold label
    if len(runs) >= 1:
        r0 = runs[0]
        r0.text = bold_label
        r0.bold = True
    else:
        r0 = para.add_run(bold_label)
        r0.bold = True
    
    # Add or update run 1 - space + plain text
    if len(runs) >= 2:
        r1 = runs[1]
        r1.text = " " + plain_text
        r1.bold = False
    else:
        r1 = para.add_run(" " + plain_text)
        r1.bold = False
    
    # Clear any extra runs beyond the first two
    if len(runs) > 2:
        for extra_run in runs[2:]:
            extra_run.text = ""


def verify_doc(doc_path: Path) -> bool:
    """Verify the doc has 0 '[Fill in here]' and Q2/Q3 bold labels."""
    doc = Document(doc_path)
    
    fill_remaining = 0
    bold_label_issues = []
    
    in_q2 = False
    in_q3 = False
    
    for para in doc.paragraphs:
        text = para.text
        
        if "Q2:" in text and ("Why" in text or "here" in text or "company" in text):
            in_q2 = True
            in_q3 = False
            continue
        
        if "Q3:" in text and ("do we do" in text or "What" in text or "know" in text):
            in_q3 = True
            in_q2 = False
            continue
        
        if "Q4:" in text:
            in_q2 = False
            in_q3 = False
            continue
        
        if "[Fill in here]" in text:
            fill_remaining += 1
        
        # Check bold labels in Q2/Q3 paragraphs
        if (in_q2 or in_q3) and para.runs:
            run0 = para.runs[0]
            if run0.text and run0.text.strip():
                if not (run0.bold and run0.text.rstrip().endswith(":")):
                    bold_label_issues.append(f"  Para: '{para.text[:60]}...' | run[0].bold={run0.bold} | run[0].text='{run0.text[:30]}'")
    
    ok = fill_remaining == 0 and len(bold_label_issues) == 0
    print(f"\nVerification: {doc_path.name}")
    print(f"  [Fill in here] remaining: {fill_remaining}")
    if bold_label_issues:
        print(f"  Bold label issues ({len(bold_label_issues)}):")
        for issue in bold_label_issues:
            print(issue)
    else:
        print(f"  Bold labels: OK")
    print(f"  Result: {'PASS' if ok else 'FAIL'}")
    return ok


def main():
    print("=== STEP 1: Copy files ===")
    for key, src in SOURCES.items():
        dest = BUNDLE_DIR / DEST_NAMES[key]
        shutil.copy2(src, dest)
        print(f"Copied: {src.name} -> {dest.name}")
    
    print("\n=== STEP 2: Fill Q2 and Q3 ===")
    master_path = BUNDLE_DIR / DEST_NAMES["master"]
    shorter_path = BUNDLE_DIR / DEST_NAMES["shorter"]
    
    fill_placeholders(master_path)
    fill_placeholders(shorter_path)
    
    print("\n=== STEP 3: Verify ===")
    ok_master = verify_doc(master_path)
    ok_shorter = verify_doc(shorter_path)
    
    print("\n=== STEP 4: Zip bundle ===")
    zip_path = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/Lance_PM_Bundle.zip")
    files_to_zip = [
        BUNDLE_DIR / DEST_NAMES["master"],
        BUNDLE_DIR / DEST_NAMES["shorter"],
        BUNDLE_DIR / DEST_NAMES["jd"],
        BUNDLE_DIR / DEST_NAMES["resume"],
    ]
    
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files_to_zip:
            zf.write(f, f.name)
            print(f"Zipped: {f.name}")
    
    print(f"\nBundle zip: {zip_path}")
    print(f"Zip size: {zip_path.stat().st_size / 1024:.1f} KB")
    
    print("\n=== DONE ===")
    print(f"Master verify: {'PASS' if ok_master else 'FAIL'}")
    print(f"Shorter verify: {'PASS' if ok_shorter else 'FAIL'}")
    print("DONE")


if __name__ == "__main__":
    main()
