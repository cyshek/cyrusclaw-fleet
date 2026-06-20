#!/usr/bin/env python3
"""Build Lance PM interview bundle - v2 (handles single placeholder + insert for bullet 2)."""

import shutil
import zipfile
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Paths
BUNDLE_DIR = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/lance-pm")
BUNDLE_DIR.mkdir(parents=True, exist_ok=True)

SOURCES = {
    "master": Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Master_Interview_Prep_Guide.docx"),
    "shorter": Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Shorter_Master_Interview_Prep_Guide.docx"),
    "jd": Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/lance-294e2602-725c-4d84-a7a0-a1fc602acec8/JD.md"),
    "resume": Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/lance-294e2602-725c-4d84-a7a0-a1fc602acec8/Cyrus_Shekari_Resume_ashby-lance_294e2602_v2.pdf"),
}

DEST_NAMES = {
    "master": "Lance_PM_Interview_Prep_Guide.docx",
    "shorter": "Lance_PM_Shorter_Interview_Prep_Guide.docx",
    "jd": "Lance_PM_JD.md",
    "resume": "Cyrus_Shekari_Resume_Lance_PM.pdf",
}

# Q2 fills
Q2_FILLS = [
    {
        "bold_label": "Why do I want to work at Lance:",
        "plain_text": "Lance is building autonomous AI agents that handle real hotel operations end-to-end — not just a chatbot layer, but agents that take actions inside live hotel software. The combination of Computer Use Agent technology and a focused vertical is a genuinely interesting product problem, and the traction with 50+ hotels across Marriott, Hilton, and Hyatt shows it is working.",
    },
    {
        "bold_label": "Why I want this role:",
        "plain_text": "This PM role maps directly to what I have been doing at Microsoft — owning a roadmap, coordinating cross-functional teams, and driving a product from ambiguous early-stage through scaled execution. I am drawn to the pace of a YC company where decisions move fast and the PM has real influence over what gets built.",
    },
]

# Q3 fills
Q3_FILLS = [
    {
        "bold_label": "What Lance does:",
        "plain_text": "Lance builds autonomous AI agents for the hospitality industry that handle calls, close sales, and manage hotel operations by integrating directly into existing hotel software to make real-time decisions. Backed by YC W26, they work with 50+ hotels across Marriott, Hilton, and Hyatt.",
    },
    {
        "bold_label": "What this role is:",
        "plain_text": "The Product Manager owns the roadmap and lifecycle for assigned products, working across engineering, design, and business teams to define requirements, prioritize features, and ship. It is an in-person role in San Francisco focused on AI and SaaS product delivery.",
    },
]


def set_paragraph_content(para, bold_label: str, plain_text: str):
    """Replace paragraph runs with bold label + plain text, preserving paragraph formatting."""
    # Remove all existing runs from the paragraph XML
    p_elem = para._p
    # Remove all 'w:r' (run) elements
    for r_elem in p_elem.findall(qn('w:r')):
        p_elem.remove(r_elem)
    
    # Also remove any hyperlink elements that may contain runs
    for hl in p_elem.findall(qn('w:hyperlink')):
        p_elem.remove(hl)
    
    # Add bold label run
    r_bold = para.add_run(bold_label)
    r_bold.bold = True
    
    # Add plain text run
    r_plain = para.add_run(" " + plain_text)
    r_plain.bold = False


def insert_paragraph_after(para, bold_label: str, plain_text: str):
    """Insert a new paragraph after 'para' with the given content, copying para's style."""
    # Clone the paragraph's XML to get the same paragraph properties (indent, spacing, etc.)
    new_p = deepcopy(para._p)
    
    # Remove all run elements from the clone
    for r_elem in new_p.findall(qn('w:r')):
        new_p.remove(r_elem)
    for hl in new_p.findall(qn('w:hyperlink')):
        new_p.remove(hl)
    
    # Insert the new paragraph XML after the current one
    para._p.addnext(new_p)
    
    # Now find the new paragraph object and add runs
    # The new paragraph is now the next sibling
    # Get the document's paragraph objects - we need to find the new Paragraph wrapper
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, para._parent)
    
    r_bold = new_para.add_run(bold_label)
    r_bold.bold = True
    
    r_plain = new_para.add_run(" " + plain_text)
    r_plain.bold = False
    
    return new_para


def fill_placeholders(doc_path: Path) -> None:
    """Fill Q2 and Q3 placeholders in the given docx file."""
    doc = Document(doc_path)
    
    in_q2 = False
    in_q3 = False
    q2_count = 0
    q3_count = 0
    
    # We need to iterate carefully since we may insert paragraphs
    # Collect the paragraphs first, then process
    paragraphs = list(doc.paragraphs)
    
    for para in paragraphs:
        text = para.text
        
        # Detect section headers
        if "Q2:" in text and ("Why" in text or "here" in text or "company" in text):
            in_q2 = True
            in_q3 = False
            continue
        
        if "Q3:" in text and ("do we do" in text or "What" in text or "know" in text):
            in_q3 = True
            in_q2 = False
            continue
        
        if "Q4:" in text:
            in_q2 = False
            in_q3 = False
            continue
        
        # Fill placeholders
        if "[Fill in here" in text:
            if in_q2 and q2_count < len(Q2_FILLS):
                fill = Q2_FILLS[q2_count]
                print(f"  Filling Q2 bullet 1: {fill['bold_label']}")
                set_paragraph_content(para, fill["bold_label"], fill["plain_text"])
                q2_count += 1
                
                # Insert bullet 2 if we need it (Q2 needs 2 bullets, only 1 placeholder)
                if q2_count < len(Q2_FILLS):
                    fill2 = Q2_FILLS[q2_count]
                    print(f"  Inserting Q2 bullet 2: {fill2['bold_label']}")
                    insert_paragraph_after(para, fill2["bold_label"], fill2["plain_text"])
                    q2_count += 1
                    
            elif in_q3 and q3_count < len(Q3_FILLS):
                fill = Q3_FILLS[q3_count]
                print(f"  Filling Q3 bullet 1: {fill['bold_label']}")
                set_paragraph_content(para, fill["bold_label"], fill["plain_text"])
                q3_count += 1
                
                # Insert bullet 2 if we need it (Q3 needs 2 bullets, only 1 placeholder)
                if q3_count < len(Q3_FILLS):
                    fill2 = Q3_FILLS[q3_count]
                    print(f"  Inserting Q3 bullet 2: {fill2['bold_label']}")
                    insert_paragraph_after(para, fill2["bold_label"], fill2["plain_text"])
                    q3_count += 1
    
    doc.save(doc_path)
    print(f"Saved: {doc_path.name} | Q2 filled: {q2_count} | Q3 filled: {q3_count}")


def verify_doc(doc_path: Path) -> bool:
    """Verify the doc: 0 '[Fill in here]' remaining, Q2/Q3 bold labels correct."""
    doc = Document(doc_path)
    
    fill_remaining = 0
    q2_checks = []
    q3_checks = []
    
    in_q2 = False
    in_q3 = False
    q2_done = False
    q3_done = False
    q2_bullet_count = 0
    q3_bullet_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        
        if not q2_done and "Q2:" in text and ("Why" in text or "here" in text or "company" in text):
            in_q2 = True
            in_q3 = False
            q2_bullet_count = 0
            continue
        
        if not q3_done and "Q3:" in text and ("do we do" in text or "What" in text or "know" in text):
            # Only track the FIRST Q3 that follows Q2 (within Section 1)
            if in_q2 or q2_done:
                in_q3 = True
                in_q2 = False
                q2_done = True
                q3_bullet_count = 0
            continue
        
        if "Q4:" in text:
            if in_q2:
                in_q2 = False
                q2_done = True
            if in_q3:
                in_q3 = False
                q3_done = True
            continue
        
        if "[Fill in here" in text:
            fill_remaining += 1
        
        # Check bold labels on our expected fill paragraphs (only first 2 bullets per section)
        if in_q2 and para.runs and para.text.strip() and "Script:" not in para.text:
            r0 = para.runs[0]
            if r0.text and r0.text.strip():
                q2_bullet_count += 1
                if q2_bullet_count <= 2:
                    ok = r0.bold == True and r0.text.rstrip().endswith(":")
                    q2_checks.append((ok, r0.text[:50], r0.bold))
        
        if in_q3 and para.runs and para.text.strip() and "Script:" not in para.text:
            r0 = para.runs[0]
            if r0.text and r0.text.strip():
                q3_bullet_count += 1
                if q3_bullet_count <= 2:
                    ok = r0.bold == True and r0.text.rstrip().endswith(":")
                    q3_checks.append((ok, r0.text[:50], r0.bold))
    
    print(f"\nVerification: {doc_path.name}")
    print(f"  [Fill in here] remaining: {fill_remaining}")
    print(f"  Q2 bullets found: {len(q2_checks)}")
    for i, (ok, txt, bold) in enumerate(q2_checks):
        print(f"    Q2[{i+1}]: bold={bold} text='{txt}' → {'OK' if ok else 'FAIL'}")
    print(f"  Q3 bullets found: {len(q3_checks)}")
    for i, (ok, txt, bold) in enumerate(q3_checks):
        print(f"    Q3[{i+1}]: bold={bold} text='{txt}' → {'OK' if ok else 'FAIL'}")
    
    all_bold_ok = all(ok for ok, _, _ in q2_checks + q3_checks)
    got_expected_bullets = len(q2_checks) >= 2 and len(q3_checks) >= 2
    result = fill_remaining == 0 and all_bold_ok and got_expected_bullets
    print(f"  Result: {'PASS' if result else 'FAIL'}")
    return result


def main():
    print("=== STEP 1: Copy files ===")
    for key, src in SOURCES.items():
        dest = BUNDLE_DIR / DEST_NAMES[key]
        shutil.copy2(src, dest)
        print(f"Copied: {src.name} -> {dest.name}")
    
    print("\n=== STEP 2: Fill Q2 and Q3 ===")
    master_path = BUNDLE_DIR / DEST_NAMES["master"]
    shorter_path = BUNDLE_DIR / DEST_NAMES["shorter"]
    
    print(f"\nFilling master guide...")
    fill_placeholders(master_path)
    
    print(f"\nFilling shorter guide...")
    fill_placeholders(shorter_path)
    
    print("\n=== STEP 3: Verify ===")
    ok_master = verify_doc(master_path)
    ok_shorter = verify_doc(shorter_path)
    
    print("\n=== STEP 4: Zip bundle ===")
    zip_path = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/Lance_PM_Bundle.zip")
    files_to_zip = [
        BUNDLE_DIR / DEST_NAMES["master"],
        BUNDLE_DIR / DEST_NAMES["shorter"],
        BUNDLE_DIR / DEST_NAMES["jd"],
        BUNDLE_DIR / DEST_NAMES["resume"],
    ]
    
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files_to_zip:
            zf.write(f, f.name)
            print(f"Zipped: {f.name}")
    
    print(f"\nBundle zip: {zip_path}")
    print(f"Zip size: {zip_path.stat().st_size / 1024:.1f} KB")
    
    print("\n=== SUMMARY ===")
    print(f"Master verify: {'PASS' if ok_master else 'FAIL'}")
    print(f"Shorter verify: {'PASS' if ok_shorter else 'FAIL'}")
    print("DONE")


if __name__ == "__main__":
    main()
