#!/usr/bin/env python3
"""Build MediaAlpha TPM interview prep guides (full + shorter)."""

import sys
sys.path.insert(0, '/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/role-discovery/.venv/lib/python3.12/site-packages')

import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# --- Content ---

Q2_BULLETS = [
    (
        "Why do I want to work at MediaAlpha: ",
        "MediaAlpha sits at an intersection I find genuinely interesting — using technology and data science to solve customer acquisition in high-consideration insurance categories. What stands out to me is that you're an intentionally small team with outsized scope: the systems you run handle hundreds of thousands of high-value transactions a day, and the TPM role sits right at the nexus of engineering, business, and external customers. That combination of real scale, lean team, and a product that's deeply technical is exactly the environment where I do my best work."
    ),
    (
        "Why do I want this role: ",
        "This TPM role maps almost directly to what I've been doing. At Microsoft I owned an end-to-end reliability program — defined success metrics, coordinated across engineering and business stakeholders, managed backlogs and release workflows, and worked closely with external partners on integration quality. The MediaAlpha role asks for the same muscle: managing a functional area end-to-end, consulting on API and data-passing best practices, and keeping cross-functional teams aligned. I'm also comfortable in the API and SQL layer the JD calls out — that's been core to how I track program health and debug integration issues."
    ),
]

Q3_BULLETS = [
    (
        "What does MediaAlpha do: ",
        "MediaAlpha is a customer acquisition solutions provider that combines technology and data science to connect consumers with providers in high-consideration insurance categories — property and casualty, health, and life insurance. The business is essentially a performance marketplace: insurers and agents use MediaAlpha's platform to reach consumers who are actively shopping, and the platform needs to handle the volume and complexity that comes with that — hundreds of thousands of high-value transactions per day."
    ),
    (
        "What is this role: ",
        "The Technical Program Manager role is the connective tissue between MediaAlpha's engineering teams, internal account managers, and external partners. Day-to-day that means owning a functional area end-to-end: defining and measuring success, managing Backlogs and Kanban workflows, coordinating API integrations and releases, and consulting with partners on data-passing and API design best practices. It's a role that requires being equally comfortable reviewing code with engineers and talking strategy with business stakeholders — which is exactly how I've operated at Microsoft."
    ),
]

FULL_MASTER = '/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Master_Interview_Prep_Guide.docx'
SHORT_MASTER = '/home/azureuser/.openclaw/agents/interview-prep/workspace/templates/Shorter_Master_Interview_Prep_Guide.docx'
OUT_DIR = '/home/azureuser/.openclaw/agents/interview-prep/workspace/guides/mediaalpha-tpm/'

FULL_OUT = OUT_DIR + 'MediaAlpha_TPM_Interview_Prep_Guide.docx'
SHORT_OUT = OUT_DIR + 'MediaAlpha_TPM_Shorter_Interview_Prep_Guide.docx'

FULL_PLACEHOLDER_Q2 = '[Fill in here]'
FULL_PLACEHOLDER_Q3 = '[Fill in here. Job description usually explains both of these]'
SHORT_PLACEHOLDER_Q2 = '[Fill in here]'
SHORT_PLACEHOLDER_Q3 = '[Fill in here. Job description usually explains both of these]'


def add_bullet_paragraph(doc, parent_element, para_idx, label, body, style_name='normal'):
    """Insert a new paragraph with bold label run + plain body run after para_idx."""
    # Create a new paragraph element
    new_para = OxmlElement('w:p')

    # Copy paragraph properties from existing paragraph if available
    try:
        existing_para = doc.paragraphs[para_idx]
        pPr = existing_para._p.find(qn('w:pPr'))
        if pPr is not None:
            new_para.append(copy.deepcopy(pPr))
    except Exception:
        pass

    # Bold run (label)
    r1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    bold1 = OxmlElement('w:b')
    rPr1.append(bold1)
    r1.append(rPr1)
    t1 = OxmlElement('w:t')
    t1.text = label
    t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r1.append(t1)
    new_para.append(r1)

    # Plain run (body)
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = body
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r2.append(t2)
    new_para.append(r2)

    return new_para


def fill_placeholders(src_path, dst_path, q2_placeholder, q3_placeholder):
    """Copy src → dst then fill Q2 and Q3 placeholders."""
    shutil.copy2(src_path, dst_path)
    doc = Document(dst_path)

    body = doc.element.body

    # Find paragraph indices for Q2 and Q3 placeholders
    q2_idx = None
    q3_idx = None
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt == q2_placeholder and q2_idx is None:
            q2_idx = i
        elif txt == q3_placeholder and q3_idx is None:
            q3_idx = i

    print(f"  Q2 placeholder at para {q2_idx}: {repr(doc.paragraphs[q2_idx].text[:60])}")
    print(f"  Q3 placeholder at para {q3_idx}: {repr(doc.paragraphs[q3_idx].text[:60])}")

    # We process Q3 first (higher index) so Q2 index isn't shifted
    # But actually after removing Q2 placeholder, Q3 index shifts by (len(bullets)-1)
    # Strategy: collect the actual XML elements, then manipulate

    q2_para_elem = doc.paragraphs[q2_idx]._p
    q3_para_elem = doc.paragraphs[q3_idx]._p

    # Build replacement paragraphs for Q2
    q2_new_paras = []
    for label, body_text in Q2_BULLETS:
        new_p = OxmlElement('w:p')
        # Bold run
        r1 = OxmlElement('w:r')
        rPr1 = OxmlElement('w:rPr')
        bold1 = OxmlElement('w:b')
        rPr1.append(bold1)
        r1.append(rPr1)
        t1 = OxmlElement('w:t')
        t1.text = label
        t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r1.append(t1)
        new_p.append(r1)
        # Plain run
        r2 = OxmlElement('w:r')
        t2 = OxmlElement('w:t')
        t2.text = body_text
        t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r2.append(t2)
        new_p.append(r2)
        q2_new_paras.append(new_p)

    # Build replacement paragraphs for Q3
    q3_new_paras = []
    for label, body_text in Q3_BULLETS:
        new_p = OxmlElement('w:p')
        # Bold run
        r1 = OxmlElement('w:r')
        rPr1 = OxmlElement('w:rPr')
        bold1 = OxmlElement('w:b')
        rPr1.append(bold1)
        r1.append(rPr1)
        t1 = OxmlElement('w:t')
        t1.text = label
        t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r1.append(t1)
        new_p.append(r1)
        # Plain run
        r2 = OxmlElement('w:r')
        t2 = OxmlElement('w:t')
        t2.text = body_text
        t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r2.append(t2)
        new_p.append(r2)
        q3_new_paras.append(new_p)

    # Replace Q3 placeholder (do Q3 first since higher index)
    q3_parent = q3_para_elem.getparent()
    q3_pos = list(q3_parent).index(q3_para_elem)
    q3_parent.remove(q3_para_elem)
    for offset, new_p in enumerate(q3_new_paras):
        q3_parent.insert(q3_pos + offset, new_p)

    # Replace Q2 placeholder
    q2_parent = q2_para_elem.getparent()
    q2_pos = list(q2_parent).index(q2_para_elem)
    q2_parent.remove(q2_para_elem)
    for offset, new_p in enumerate(q2_new_paras):
        q2_parent.insert(q2_pos + offset, new_p)

    doc.save(dst_path)
    print(f"  Saved: {dst_path}")
    return dst_path


def verify_doc(path, placeholders_to_check, expected_bold_labels):
    """Verify no placeholders remain and all bold labels are present."""
    doc = Document(path)
    errors = []

    # Check no placeholders remain
    for ph in placeholders_to_check:
        for p in doc.paragraphs:
            if ph in p.text:
                errors.append(f"PLACEHOLDER STILL PRESENT: {repr(ph)}")

    # Check bold labels present
    found_labels = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.bold and ':' in r.text:
                found_labels.append(r.text)

    for expected in expected_bold_labels:
        matched = any(expected in lbl for lbl in found_labels)
        if not matched:
            errors.append(f"BOLD LABEL NOT FOUND: {repr(expected)}")

    return errors, found_labels


# --- Main ---
PLACEHOLDERS = [FULL_PLACEHOLDER_Q2, FULL_PLACEHOLDER_Q3]
EXPECTED_LABELS = [
    "Why do I want to work at MediaAlpha:",
    "Why do I want this role:",
    "What does MediaAlpha do:",
    "What is this role:",
]

print("=== Building Full Guide ===")
fill_placeholders(FULL_MASTER, FULL_OUT, FULL_PLACEHOLDER_Q2, FULL_PLACEHOLDER_Q3)
errors, labels = verify_doc(FULL_OUT, PLACEHOLDERS, EXPECTED_LABELS)
print(f"  Verification errors: {errors}")
print(f"  Bold labels found: {labels}")

print("\n=== Building Shorter Guide ===")
fill_placeholders(SHORT_MASTER, SHORT_OUT, SHORT_PLACEHOLDER_Q2, SHORT_PLACEHOLDER_Q3)
errors2, labels2 = verify_doc(SHORT_OUT, PLACEHOLDERS, EXPECTED_LABELS)
print(f"  Verification errors: {errors2}")
print(f"  Bold labels found: {labels2}")

print("\nDone.")
