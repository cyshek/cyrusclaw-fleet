import zipfile, io, tempfile, os
import docx

def check(zip_path, company, q2_labels, q3_labels):
    print("=== " + zip_path + " ===")
    with zipfile.ZipFile(zip_path) as z:
        names = z.namelist()
        print("  files:", names)
        docx_name = [n for n in names if n.endswith(".docx")][0]
        data = z.read(docx_name)
    tf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tf.write(data); tf.close()
    doc = docx.Document(tf.name)
    txt = chr(10).join(p.text for p in doc.paragraphs)
    assert "Fill in here" not in txt, "PLACEHOLDER LEFT"
    for lbl in q2_labels + q3_labels:
        assert lbl in txt, "MISSING LABEL: " + lbl
    # bold check on the 4 target bullets
    bold_ok = 0
    for p in doc.paragraphs:
        if p.runs and p.runs[0].bold and any(p.runs[0].text.strip().startswith(lbl[:12]) for lbl in q2_labels+q3_labels):
            bold_ok += 1
    has_company = company in txt
    has_sections = all(s in txt for s in ["Section 1", "Section 4", "Section 5"])
    print("  no placeholders: OK")
    print("  company name present (" + company + "):", has_company)
    print("  Q2/Q3 target bullets bold:", bold_ok, "/4")
    print("  Sections 1/4/5 intact:", has_sections)
    os.unlink(tf.name)
    print()

check("guides/langchain-sales-engineer.zip", "LangChain",
      ["Why do I want to work at LangChain:", "Why do I want this role:"],
      ["What does LangChain do:", "What is this role:"])
check("guides/mediaalpha-tpm.zip", "MediaAlpha",
      ["Why do I want to work at MediaAlpha:", "Why do I want this role:"],
      ["What does MediaAlpha do:", "What is this role:"])
print("ALL ZIP CHECKS PASSED")
