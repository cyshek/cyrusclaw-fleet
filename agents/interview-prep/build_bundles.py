#!/usr/bin/env python3
import zipfile, shutil, re, sys
from pathlib import Path

WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATES = WORKSPACE / "templates"
GUIDES = WORKSPACE / "guides"
JOB_SEARCH = Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted")
MASTER_FULL = TEMPLATES / "Master_Interview_Prep_Guide.docx"
MASTER_SHORT = TEMPLATES / "Shorter_Master_Interview_Prep_Guide.docx"

PODIUM_Q2 = [
    ("Why do I want to work at Podium:", "Podium is doing something genuinely exciting — bringing AI-powered communication and payments to local businesses that don't have the resources of a big enterprise. The scale of impact (60,000+ businesses, $100M ARR in under two years) and the speed at which they operate is the kind of environment I want to be building in right now. I'm drawn to companies where product work moves fast and actually reaches real customers quickly."),
    ("Why do I want this role:", "This PM role sits at the intersection of two things I care about: owning real product strategy end-to-end and working in a space where AI is the core of the product, not just a feature. The expectation that PMs here are builders — prototyping with AI tools, getting into the weeds with engineers — matches exactly how I like to work. I want to be in a role where I'm close to the customer and close to the code, and this one checks both boxes."),
]

PODIUM_Q3 = [
    ("What does Podium do:", "Podium is an AI-powered platform for local and SMB businesses — it brings together customer messaging, lead capture, payments, and reviews in one place. Their AI agents handle real customer interactions (texting, qualifying leads, booking appointments) so small businesses can operate like they have a full-time staff without the headcount. They're particularly strong in verticals like auto, home services, and aesthetics."),
    ("What is this role:", "The Product Manager role at Podium owns a core area of the platform end-to-end — strategy, roadmap, and shipping. It's a builder role: PMs use AI tools to prototype and contribute alongside engineers, not just write specs. The job is to balance near-term feature delivery with longer-term platform thinking, all while staying close to real customer workflows."),
]

SCALEAI_Q2 = [
    ("Why do I want to work at Scale AI:", "Scale AI is doing foundational work — not building another AI application, but building the infrastructure layer that makes reliable AI possible. The focus on high-quality data, rigorous evaluation, and now the public sector mission (where the stakes are genuinely high) is a rare combination. I want to be at a company where rigor is a competitive advantage, and Scale has built that reputation."),
    ("Why do I want this role:", "The Product Manager, Public Sector GenAI T&E role combines two things I find compelling: deeply technical product work (owning evaluation infrastructure, not just feature roadmaps) and a domain where the outcomes actually matter — government and defense use cases where you can't cut corners. My background in technical program management across complex cross-functional programs maps well to what this role demands: translating ambiguous problems into structured plans and driving execution across multiple engineering orgs."),
]

SCALEAI_Q3 = [
    ("What does Scale AI do:", "Scale AI provides the data infrastructure and tooling that AI companies and enterprises need to build, train, and evaluate their models. They're best known for data labeling and annotation at scale, but they've expanded significantly into RLHF, AI evaluation platforms, and now a major public sector business helping government agencies deploy and assess AI systems responsibly."),
    ("What is this role:", "The PM for Public Sector GenAI T&E owns the roadmap for Scale's evaluation capabilities in government contexts — the infrastructure used to continuously measure, improve, and validate the performance of AI and agentic systems for public sector clients. It's a highly technical PM role sitting at the intersection of ML systems, government requirements, and cross-org execution across Scale's engineering teams."),
]

def escape_xml(text):
    return (text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
                .replace('"',"&quot;").replace("'","&apos;"))

def extract_pPr(para_xml):
    m = re.search(r"<w:pPr>.*?</w:pPr>", para_xml, re.DOTALL)
    return m.group(0) if m else "<w:pPr/>"

def make_bullet_xml(pPr_xml, bold_label, plain_text):
    r1 = '<w:r><w:rPr><w:b w:val="1"/><w:bCs w:val="1"/><w:rtl w:val="0"/></w:rPr><w:t xml:space="preserve">' + escape_xml(bold_label) + ' </w:t></w:r>'
    r2 = '<w:r><w:rPr><w:rtl w:val="0"/></w:rPr><w:t xml:space="preserve">' + escape_xml(plain_text) + '</w:t></w:r>'
    return "<w:p>" + pPr_xml + r1 + r2 + "</w:p>"

def patch_docx(src, dst, q2_bullets, q3_bullets):
    shutil.copy2(src, dst)
    with zipfile.ZipFile(dst, "r") as zin:
        content = zin.read("word/document.xml").decode("utf-8")
        all_files = {name: zin.read(name) for name in zin.namelist()}
    placeholders = []
    idx = 0
    while True:
        pos = content.find("Fill in here", idx)
        if pos == -1:
            break
        p_start = content.rfind("<w:p ", 0, pos)
        p_end = content.find("</w:p>", pos) + len("</w:p>")
        placeholders.append((p_start, p_end, content[p_start:p_end]))
        idx = pos + 1
    if len(placeholders) != 2:
        raise ValueError(f"Expected 2 placeholders, found {len(placeholders)} in {dst}")
    bullets_map = {0: q2_bullets, 1: q3_bullets}
    replacements = []
    for i, (p_start, p_end, para_xml) in enumerate(placeholders):
        pPr = extract_pPr(para_xml)
        new_paras = "".join(make_bullet_xml(pPr, label, text) for label, text in bullets_map[i])
        replacements.append((p_start, p_end, new_paras))
    for p_start, p_end, new_paras in reversed(replacements):
        content = content[:p_start] + new_paras + content[p_end:]
    all_files["word/document.xml"] = content.encode("utf-8")
    dst.unlink()
    with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in all_files.items():
            zout.writestr(name, data)
    print(f"  Patched: {dst.name}")

def verify_docx(path):
    try:
        import docx
        doc = docx.Document(str(path))
    except Exception as e:
        print(f"  ERROR: python-docx failed on {path.name}: {e}")
        return False
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "Fill in here" in full_text:
        print(f"  ERROR: [Fill in here] still present in {path.name}")
        return False
    bold_bullets = []
    for p in doc.paragraphs:
        if p.runs and p.runs[0].bold and ":" in p.runs[0].text:
            bold_bullets.append(p.runs[0].text)
    if len(bold_bullets) < 4:
        print(f"  ERROR: Expected 4+ bold-label bullets, found {len(bold_bullets)} in {path.name}")
        for p in doc.paragraphs:
            if p.runs:
                print(f"    para={p.text[:60]!r} run0.bold={p.runs[0].bold} run0.text={p.runs[0].text[:40]!r}")
        return False
    if "Q1" not in full_text:
        print(f"  ERROR: Original content Q1 missing from {path.name}")
        return False
    print(f"  PASS: {path.name} - no placeholders, {len(bold_bullets)} bold bullets, Q1 present")
    print(f"  Bold labels: {bold_bullets}")
    return True

def build_zip(zip_path, files):
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files:
            f = Path(f)
            if f.exists():
                z.write(f, f.name)
                print(f"  Added to zip: {f.name}")
            else:
                print(f"  WARNING: File not found: {f}")
    print(f"  Zip created: {zip_path}")

def build_bundle(name, folder, zip_path, full_name, short_name, q2_bullets, q3_bullets, jd_path, resume_path):
    print("\n" + "="*60)
    print(f"Building bundle: {name}")
    print("="*60)
    folder.mkdir(parents=True, exist_ok=True)
    full_dst = folder / full_name
    short_dst = folder / short_name
    print("\n[1/4] Patching full guide...")
    patch_docx(MASTER_FULL, full_dst, q2_bullets, q3_bullets)
    print("\n[2/4] Patching shorter guide...")
    patch_docx(MASTER_SHORT, short_dst, q2_bullets, q3_bullets)
    print("\n[3/4] Verifying docs...")
    ok_full = verify_docx(full_dst)
    ok_short = verify_docx(short_dst)
    if not (ok_full and ok_short):
        print(f"  VERIFICATION FAILED for {name}")
        return False
    print("\n[4/4] Building zip...")
    build_zip(zip_path, [full_dst, short_dst, jd_path, resume_path])
    print(f"\nBundle complete: {zip_path}")
    return True

if __name__ == "__main__":
    success = True
    podium_ok = build_bundle(
        name="Podium PM",
        folder=GUIDES / "podium-pm",
        zip_path=GUIDES / "podium-pm.zip",
        full_name="Podium_Interview_Prep_Guide.docx",
        short_name="Podium_Interview_Prep_Guide_Short.docx",
        q2_bullets=PODIUM_Q2,
        q3_bullets=PODIUM_Q3,
        jd_path=JOB_SEARCH / "podium-7825829/JD.md",
        resume_path=JOB_SEARCH / "podium-7825829/Cyrus_Shekari_Resume_podium81_7825829_v2.pdf",
    )
    success = success and podium_ok
    scaleai_ok = build_bundle(
        name="Scale AI PM T&E",
        folder=GUIDES / "scale-ai-pm-te",
        zip_path=GUIDES / "scale-ai-pm-te.zip",
        full_name="ScaleAI_Interview_Prep_Guide.docx",
        short_name="ScaleAI_Interview_Prep_Guide_Short.docx",
        q2_bullets=SCALEAI_Q2,
        q3_bullets=SCALEAI_Q3,
        jd_path=JOB_SEARCH / "scale-ai-4687591005/JD.md",
        resume_path=JOB_SEARCH / "scale-ai-4687591005/Cyrus_Shekari_Resume_scaleai_4687591005_v2.pdf",
    )
    success = success and scaleai_ok
    print("\n" + "="*60)
    print("FINAL STATUS: " + ("ALL BUNDLES OK" if success else "SOME BUNDLES FAILED"))
    print("="*60)
    sys.exit(0 if success else 1)
