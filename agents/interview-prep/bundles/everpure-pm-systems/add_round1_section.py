#!/usr/bin/env python3
"""
Appends a Round 1 coaching section to the Everpure PM Systems Interview Prep Guide.
ADDENDUM ONLY — does not touch any existing content.

Uses numId=1, ilvl=0 (matching the existing simple bullet style in the doc).
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import copy

DOC_PATH = "/home/azureuser/.openclaw/agents/interview-prep/workspace/bundles/everpure-pm-systems/Everpure_PM_Systems_Interview_Prep_Guide.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_plain(doc, text):
    p = doc.add_paragraph(style="normal")
    p.add_run(text)
    return p


def add_bullet(doc, text, num_id="1", ilvl="0"):
    """Add a bullet paragraph using existing numId from the document."""
    p = doc.add_paragraph(style="normal")
    p.add_run(text)

    # Build pPr with numPr
    pPr = p._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)
    return p


def main():
    if not os.path.exists(DOC_PATH):
        print(f"ERROR: File not found: {DOC_PATH}")
        return False

    doc = Document(DOC_PATH)

    existing_count = len(doc.paragraphs)
    print(f"Existing paragraph count: {existing_count}")

    # ── Heading 1 ──────────────────────────────────────────────────────────────
    add_heading(doc, "Round 1 Prep — Mayank Bhatnagar (Thursday)", 1)

    # ── Section 1: What Mayank is evaluating ──────────────────────────────────
    add_heading(doc, "What Mayank is evaluating (60 min)", 2)
    add_plain(doc,
        "Three things per the interview packet: overall fit, market understanding, and user experience focus. "
        "This is a hiring manager screen — he wants to know who you are, whether you think like a PM for a storage product, "
        "and whether you understand users. He is not trying to trick you."
    )

    # ── Section 2: Story-to-criteria mapping ──────────────────────────────────
    add_heading(doc, "Story-to-criteria mapping", 2)
    add_plain(doc, "Use these story anchors when he probes each area:")
    add_bullet(doc,
        "Fit → Lead with the 'tell me about yourself' script. Hit: technical program ownership at scale, data-driven, "
        "bridge between eng and outcomes. Transition line: 'I'm ready for my next step where I can bring this intersection "
        "of technical depth and product strategy to a PM role.'"
    )
    add_bullet(doc,
        "Market understanding → Q2/Q3 answers above cover the baseline. If he goes deeper: FlashArray = block storage, "
        "latency-sensitive workloads (databases, VMs, mission-critical apps). FlashBlade = unstructured/file/object data at scale "
        "(AI/ML pipelines, backup, analytics). Everpure differentiates on simplicity + performance + Purity OS (software layer). "
        "You don't need to be a storage expert — you need to show genuine curiosity and that you did homework."
    )
    add_bullet(doc,
        "User experience focus → This is NOT a design/UX interview. 'User experience' here means: do you think about the people "
        "who size, deploy, and operate these systems? Internal users = Solutions Engineers, field teams, ops. External users = enterprise IT teams. "
        "Best story: the platform adoption pivot (Section 2 Q3 / Q4) — you discovered real users weren't using the tool, diagnosed root friction, "
        "fixed it. That is exactly what Mayank means by UX focus."
    )

    # ── Section 3: FlashArray & FlashBlade cheat sheet ────────────────────────
    add_heading(doc, "FlashArray & FlashBlade — 60-second cheat sheet", 2)
    add_plain(doc, "Know these cold so you can drop them naturally:")
    add_bullet(doc,
        "FlashArray: all-flash block storage. Designed for latency-sensitive workloads — databases, VMs, Tier-1 apps. "
        "Customers care about consistent sub-millisecond latency and simple management."
    )
    add_bullet(doc,
        "FlashBlade: all-flash file + object storage. Designed for unstructured data at scale — AI/ML training data, analytics, "
        "backup/recovery. Customers care about throughput and consolidation."
    )
    add_bullet(doc,
        "Purity OS: the software layer that runs on both platforms. Handles data services (deduplication, compression, snapshots, replication). "
        "This is where a lot of PM work lives — features ship through Purity updates."
    )
    add_bullet(doc,
        "Everpure's pitch: simpler to manage than competitors (EMC/NetApp/HPE), better economics through deduplication, and strong "
        "telemetry/analytics (Pure1 cloud management). The telemetry angle is directly relevant to this PM role — "
        "'Drive Data-Informed Product Evolution' is job #1."
    )
    add_bullet(doc,
        "If asked about competition: Dell EMC PowerStore/PowerScale, NetApp, HPE Alletra. Everpure wins on simplicity and software-driven management."
    )

    # ── Section 4: Questions to ask Mayank ────────────────────────────────────
    add_heading(doc, "Questions to ask Mayank", 2)
    add_plain(doc, "Use 1–2 of these (they're already in Section 5, but these are tailored for Round 1 with the HM):")
    add_bullet(doc,
        "'What does a great PM in this role look like 6 months in — what are they shipping or changing?' "
        "(lets him paint his ideal hire; you can reflect it back)"
    )
    add_bullet(doc,
        "'What's the biggest gap between where the FlashArray/FlashBlade roadmap is today and where you want it to be?' "
        "(shows you think in roadmap terms, not just execution)"
    )
    add_bullet(doc,
        "'How does the team decide which telemetry signals to act on versus which to monitor?' "
        "(signals data-informed product thinking, relevant to job #1)"
    )

    # ── Section 5: Mindset for Thursday ───────────────────────────────────────
    add_heading(doc, "Mindset for Thursday", 2)
    add_plain(doc,
        "Brian's email said 'no weird gotcha questions — just real conversations.' Lean into that. Mayank is a hiring manager "
        "who wants to feel confident you can do the job. Your job is to make him see the signal clearly: technical depth, data-first thinking, "
        "cross-functional execution, and genuine curiosity about the product space. You have all of that — just make it visible."
    )

    doc.save(DOC_PATH)
    print(f"Saved: {DOC_PATH}")

    # ── Verification ──────────────────────────────────────────────────────────
    doc2 = Document(DOC_PATH)
    paragraphs = doc2.paragraphs

    heading1_found = False
    h2_sections = []
    bullet_counts = {}
    current_h2 = None
    in_round1 = False

    for p in paragraphs:
        text = p.text.strip()
        style = p.style.name
        numPr = p._element.find(f'.//{{{W}}}numPr')

        if style == "Heading 1" and "Round 1 Prep" in text:
            heading1_found = True
            in_round1 = True

        if in_round1 and style == "Heading 2":
            current_h2 = text
            h2_sections.append(text)
            bullet_counts[current_h2] = 0

        if in_round1 and current_h2 and numPr is not None:
            bullet_counts[current_h2] += 1

    print("\n=== VERIFICATION ===")
    print(f"Heading 1 'Round 1 Prep' found: {heading1_found}")
    print(f"H2 subsection count: {len(h2_sections)}")
    for s in h2_sections:
        print(f"  H2: '{s}' — {bullet_counts.get(s, 0)} bullets")

    expected_h2 = [
        "What Mayank is evaluating (60 min)",
        "Story-to-criteria mapping",
        "FlashArray & FlashBlade — 60-second cheat sheet",
        "Questions to ask Mayank",
        "Mindset for Thursday",
    ]
    expected_bullets = {
        "Story-to-criteria mapping": 3,
        "FlashArray & FlashBlade — 60-second cheat sheet": 5,
        "Questions to ask Mayank": 3,
    }

    all_ok = True
    for h2 in expected_h2:
        if h2 not in h2_sections:
            print(f"  MISSING H2: {h2}")
            all_ok = False

    for h2, expected in expected_bullets.items():
        actual = bullet_counts.get(h2, 0)
        if actual != expected:
            print(f"  BULLET MISMATCH in '{h2}': expected {expected}, got {actual}")
            all_ok = False

    if all_ok:
        print("\n✅ All checks passed.")
    else:
        print("\n❌ Some checks failed — review above.")

    return all_ok


if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)
