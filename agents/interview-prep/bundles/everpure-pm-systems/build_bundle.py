"""
Build script for Everpure PM Systems interview prep bundle.
Fills Q2 and Q3 placeholders in both full and shorter guide copies,
then zips everything into a bundle.
"""

import shutil
import zipfile
from pathlib import Path
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Paths ───────────────────────────────────────────────────────────────────
WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
TEMPLATES  = WORKSPACE / "templates"
OUTPUT_DIR = WORKSPACE / "bundles" / "everpure-pm-systems"

FULL_TEMPLATE    = TEMPLATES / "Master_Interview_Prep_Guide.docx"
SHORTER_TEMPLATE = TEMPLATES / "Shorter_Master_Interview_Prep_Guide.docx"

FULL_OUT    = OUTPUT_DIR / "Everpure_PM_Systems_Interview_Prep_Guide.docx"
SHORTER_OUT = OUTPUT_DIR / "Everpure_PM_Systems_Shorter_Interview_Prep_Guide.docx"

JD_SRC      = Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/pure-storage-7671846/JD.md")
RESUME_PDF  = Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/pure-storage-7671846/Cyrus_Shekari_Resume_purestorage_7671846_v2.pdf")

ZIP_OUT = OUTPUT_DIR / "Everpure_PM_Systems_PrepBundle.zip"

# ─── Q2 + Q3 content ─────────────────────────────────────────────────────────
# Each Q = list of (bold_label, plain_rest) tuples — one per bullet
Q2_BULLETS = [
    ("What Everpure does:", " Everpure builds high-performance all-flash storage platforms (FlashArray, FlashBlade) used in thousands of global deployments — they're a leader in enterprise storage making data faster, simpler, and more sustainable."),
    ("Why I want this role:", " I want to bring my experience driving data-informed technical programs at Microsoft to a product role where I can bridge engineering depth and customer outcomes — this role's focus on telemetry-driven roadmap decisions, cross-functional execution, and shipping hardware+software improvements is exactly the kind of work I'm built for."),
]

Q3_BULLETS = [
    ("What is this role:", " PM for Systems on the FlashArray & FlashBlade platforms — owning the end-to-end roadmap for defined functional areas, from defining requirements with Solutions Engineers and customers to driving launch readiness, while aligning Engineering, Purity software, and QA teams."),
    ("What Everpure is about:", " Everpure is reshaping enterprise data storage — FlashArray targets block storage for latency-sensitive workloads, FlashBlade is purpose-built for unstructured data at scale. The company values technical curiosity, data-driven decisions, and fast iteration."),
]


def add_bold_bullet(doc, paragraph, bold_label, plain_text):
    """
    Replace paragraph content with two runs:
      run[0] = bold_label (bold=True)
      run[1] = plain_text (bold=False)
    Preserves the paragraph's existing style (list-bullet or normal).
    """
    # Clear existing runs
    p_elem = paragraph._p
    # Remove all <w:r> and <w:hyperlink> children
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink', 'bookmarkStart', 'bookmarkEnd'):
            p_elem.remove(child)

    # Add bold run for label
    run_bold = paragraph.add_run(bold_label)
    run_bold.bold = True

    # Add plain run for rest
    run_plain = paragraph.add_run(plain_text)
    run_plain.bold = False

    return paragraph


def fill_placeholders(doc, q2_bullets, q3_bullets):
    """
    Find the two placeholder paragraphs and replace them with bullet lists.
    Q2 placeholder: '[Fill in here]'
    Q3 placeholder: '[Fill in here. Job description usually explains both of these]'
    Returns the indices of modified paragraphs for verification.
    """
    modified = []
    
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        
        if txt == '[Fill in here]':
            # Q2 — replace this paragraph with first bullet, insert second after
            bullets = q2_bullets
            add_bold_bullet(doc, para, bullets[0][0], bullets[0][1])
            modified.append(i)
            
            # Insert second bullet paragraph after this one
            # We need to insert a new paragraph after para in the XML
            new_para = OxmlElement('w:p')
            # Copy pPr (paragraph properties) from existing para if present
            pPr_src = para._p.find(qn('w:pPr'))
            if pPr_src is not None:
                new_para.insert(0, copy.deepcopy(pPr_src))
            
            # Insert after current paragraph element
            para._p.addnext(new_para)
            
            # Now get a Paragraph wrapper around it and add runs
            # Find the new paragraph in doc.paragraphs
            # We need to work with the XML element directly
            new_para_obj = docx.text.paragraph.Paragraph(new_para, doc)
            add_bold_bullet(doc, new_para_obj, bullets[1][0], bullets[1][1])
            modified.append(i + 0.5)
        
        elif txt == '[Fill in here. Job description usually explains both of these]':
            # Q3 — same pattern
            bullets = q3_bullets
            add_bold_bullet(doc, para, bullets[0][0], bullets[0][1])
            modified.append(i)
            
            new_para = OxmlElement('w:p')
            pPr_src = para._p.find(qn('w:pPr'))
            if pPr_src is not None:
                new_para.insert(0, copy.deepcopy(pPr_src))
            
            para._p.addnext(new_para)
            
            new_para_obj = docx.text.paragraph.Paragraph(new_para, doc)
            add_bold_bullet(doc, new_para_obj, bullets[1][0], bullets[1][1])
            modified.append(i + 0.5)
    
    return modified


def verify_bold_bullets(doc, labels):
    """
    For each label string, find the paragraph containing it and verify run[0].bold == True.
    Returns list of (label, found, bold_ok) tuples.
    """
    results = []
    for label in labels:
        found = False
        bold_ok = False
        for para in doc.paragraphs:
            if label in para.text:
                found = True
                if para.runs and para.runs[0].bold:
                    bold_ok = True
                break
        results.append((label, found, bold_ok))
    return results


def count_placeholders(doc):
    count = 0
    for para in doc.paragraphs:
        if '[Fill in here' in para.text:
            count += 1
    return count


def build_doc(template_path, output_path, q2_bullets, q3_bullets, doc_label):
    print(f"\n{'='*60}")
    print(f"Building: {doc_label}")
    print(f"  Template: {template_path}")
    print(f"  Output:   {output_path}")
    
    shutil.copy2(template_path, output_path)
    doc = docx.Document(output_path)
    
    # Verify placeholders exist before filling
    placeholders_before = count_placeholders(doc)
    print(f"  Placeholders before fill: {placeholders_before}")
    assert placeholders_before == 2, f"Expected 2 placeholders, found {placeholders_before}"
    
    modified = fill_placeholders(doc, q2_bullets, q3_bullets)
    print(f"  Modified paragraph indices: {modified}")
    
    doc.save(output_path)
    
    # Re-open to verify
    doc2 = docx.Document(output_path)
    
    placeholders_after = count_placeholders(doc2)
    print(f"  Placeholders after fill: {placeholders_after}")
    
    labels = [
        "What Everpure does:",
        "Why I want this role:",
        "What is this role:",
        "What Everpure is about:",
    ]
    verify_results = verify_bold_bullets(doc2, labels)
    print(f"  Bold label verification:")
    all_ok = True
    for label, found, bold_ok in verify_results:
        status = "✅" if (found and bold_ok) else "❌"
        print(f"    {status} '{label}' — found={found}, bold={bold_ok}")
        if not (found and bold_ok):
            all_ok = False
    
    assert placeholders_after == 0, f"FAIL: {placeholders_after} placeholders remain!"
    assert all_ok, "FAIL: Some bold labels missing or not bold!"
    
    print(f"  ✅ {doc_label} PASSED all checks")
    return verify_results


def build_zip():
    print(f"\n{'='*60}")
    print("Building ZIP bundle...")
    
    files_to_zip = [
        (FULL_OUT,    FULL_OUT.name),
        (SHORTER_OUT, SHORTER_OUT.name),
        (JD_SRC,      "JD.md"),
        (RESUME_PDF,  RESUME_PDF.name),
    ]
    
    for src, arcname in files_to_zip:
        assert src.exists(), f"Missing file: {src}"
    
    with zipfile.ZipFile(ZIP_OUT, 'w', zipfile.ZIP_DEFLATED) as zf:
        for src, arcname in files_to_zip:
            zf.write(src, arcname)
            print(f"  Added: {arcname}")
    
    # Verify zip contents
    with zipfile.ZipFile(ZIP_OUT, 'r') as zf:
        names = zf.namelist()
        print(f"  ZIP contents: {names}")
        assert len(names) == 4, f"Expected 4 files in zip, got {len(names)}"
    
    print(f"  ✅ ZIP saved: {ZIP_OUT}")
    print(f"  Size: {ZIP_OUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    # Build full guide
    full_results = build_doc(FULL_TEMPLATE, FULL_OUT, Q2_BULLETS, Q3_BULLETS, "Full Guide")
    
    # Build shorter guide with identical content
    shorter_results = build_doc(SHORTER_TEMPLATE, SHORTER_OUT, Q2_BULLETS, Q3_BULLETS, "Shorter Guide")
    
    # Build zip
    build_zip()
    
    print(f"\n{'='*60}")
    print("✅ BUILD COMPLETE")
    print(f"  Full guide:    {FULL_OUT}")
    print(f"  Shorter guide: {SHORTER_OUT}")
    print(f"  Bundle zip:    {ZIP_OUT}")
