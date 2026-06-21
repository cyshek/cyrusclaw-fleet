"""
Build script for Datadog Partner TSE Interview Prep Bundle.
Fills Q2/Q3 placeholders and appends Round 1 coaching section.
"""

import shutil
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ────────────────────────────────────────────────────────────────────
WORKSPACE = Path("/home/azureuser/.openclaw/agents/interview-prep/workspace")
JOB_SEARCH = Path("/home/azureuser/.openclaw/agents/job-search/workspace/projects/job-search/applications/submitted/datadog-7961297")
OUT_DIR = WORKSPACE / "bundles" / "datadog-partner-tse"
TEMPLATE = WORKSPACE / "templates" / "Master_Interview_Prep_Guide.docx"
GUIDE_NAME = "Datadog_Partner_TSE_Interview_Prep_Guide.docx"
GUIDE_OUT = OUT_DIR / GUIDE_NAME
RESUME_PDF = JOB_SEARCH / "Cyrus_Shekari_Resume_datadog_7961297_v2.pdf"
JD_MD = JOB_SEARCH / "JD.md"
ZIP_OUT = OUT_DIR / "Datadog_Partner_TSE_PrepBundle.zip"

OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Step 1: Fresh copy of template ───────────────────────────────────────────
shutil.copy2(TEMPLATE, GUIDE_OUT)
print(f"[1] Copied template → {GUIDE_OUT.name}")

# ── Helpers ───────────────────────────────────────────────────────────────────

def clear_paragraph(para):
    """Remove all runs from a paragraph."""
    for run in para.runs:
        run._r.getparent().remove(run._r)
    # Also clear any leftover XML run elements
    p_elem = para._p
    for r_elem in p_elem.findall(qn('w:r')):
        p_elem.remove(r_elem)


def add_bold_label_bullet(para, label: str, rest: str):
    """
    Populate para with two runs:
      run[0]: bold=True  → label (e.g. "What Datadog does:")
      run[1]: bold=False → rest of text
    """
    clear_paragraph(para)
    r0 = para.add_run(label)
    r0.bold = True
    r1 = para.add_run(" " + rest)
    r1.bold = False


def insert_paragraph_after(para, text="", style=None):
    """Insert a new paragraph immediately after `para`."""
    new_p = OxmlElement('w:p')
    para._p.addnext(new_p)
    new_para = para.__class__(new_p, para._p.getparent())
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


# ── Step 2: Open and patch Q2 / Q3 ───────────────────────────────────────────
doc = Document(GUIDE_OUT)
paragraphs = doc.paragraphs

# Locate the two placeholder paragraphs by their placeholder text
q2_para = None
q3_para = None
for i, p in enumerate(paragraphs):
    if p.text.strip() == "[Fill in here]":
        q2_para = p
    elif "[Fill in here. Job description" in p.text:
        q3_para = p

assert q2_para is not None, "Could not find Q2 placeholder"
assert q3_para is not None, "Could not find Q3 placeholder"
print(f"[2] Located Q2 placeholder (para index search complete)")
print(f"[2] Located Q3 placeholder (para index search complete)")

# ── Q2: Two bullets ───────────────────────────────────────────────────────────
# Bullet 1 — repurpose the existing placeholder para
add_bold_label_bullet(
    q2_para,
    "What Datadog does:",
    "Datadog is the leading cloud observability and monitoring platform — metrics, logs, traces, "
    "and security in one place, used by engineering and DevOps teams worldwide to understand what "
    "is happening in their infrastructure and applications."
)
q2_para.style = doc.styles["normal"]

# Bullet 2 — insert a NEW paragraph after the first
# We need to insert it right after q2_para in the XML
q2_bullet2_p = OxmlElement('w:p')
q2_para._p.addnext(q2_bullet2_p)
# Build the new paragraph object by finding it in the refreshed list
# (Insert after, so it appears as next sibling)
# We'll add runs directly via OxmlElement for reliability

def make_two_run_para(doc_obj, label: str, rest: str, style_name="normal"):
    """Return a new detached paragraph element with bold label + plain rest."""
    # Create paragraph XML
    new_p = OxmlElement('w:p')

    # Optional: preserve paragraph properties matching the style
    # run 0 — bold label
    r0 = OxmlElement('w:r')
    rPr0 = OxmlElement('w:rPr')
    b0 = OxmlElement('w:b')
    rPr0.append(b0)
    r0.append(rPr0)
    t0 = OxmlElement('w:t')
    t0.text = label
    r0.append(t0)
    new_p.append(r0)

    # run 1 — plain rest
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = " " + rest
    # xml:space preserve so leading space kept
    t1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r1.append(t1)
    new_p.append(r1)

    return new_p


# Replace the empty element we inserted with a proper one
q2_para._p.addnext(
    make_two_run_para(
        doc,
        "Why I want this role:",
        "I want to move into a role where I'm working at the intersection of technical depth and external impact — "
        "advising partners on integration design, diagnosing real problems, and feeding that signal back into the product. "
        "That combination of consultant, engineer, and product voice is exactly where I do my best work."
    )
)
# Remove the blank element we inserted a moment ago
q2_para._p.getnext().getnext()  # sanity — don't remove, the addnext already placed correctly

# Wait — we inserted a blank element AND then addnext again. Let's verify the structure:
# Actually the second addnext goes AFTER q2_para, pushing the blank one further.
# Let me remove the blank one.
# The blank q2_bullet2_p is now AFTER the new make_two_run_para element... let's check:
# q2_para._p  → q2 content
# q2_para._p.getnext() → make_two_run_para result (most recently added "after")
# q2_para._p.getnext().getnext() → q2_bullet2_p (the blank one)
blank_elem = q2_para._p.getnext().getnext()
if blank_elem is not None and len(blank_elem) == 0:
    blank_elem.getparent().remove(blank_elem)
    print("[2] Removed stray blank paragraph element")

print("[2] Q2 — 2 bullets inserted")

# ── Q3: Two bullets ───────────────────────────────────────────────────────────
add_bold_label_bullet(
    q3_para,
    "What is this role:",
    "Partner TSE at Datadog — you're the primary technical contact for third-party developers building integrations "
    "on the Integration Developer Platform (IDP). You guide them from architectural design through publication on the "
    "Datadog Marketplace, do code reviews against the Quality Rubric, troubleshoot complex issues, and surface friction "
    "points back to the internal product and engineering teams."
)
q3_para.style = doc.styles["normal"]

q3_para._p.addnext(
    make_two_run_para(
        doc,
        "What Datadog is about:",
        "Datadog started as infrastructure monitoring and has expanded into a full observability platform: metrics, "
        "APM (application performance monitoring), log management, security, and now AI observability. $2.7B in revenue, "
        "26% YoY growth. The integration ecosystem is how Datadog extends its platform — the Partner TSE role is central "
        "to that growth engine."
    )
)
print("[2] Q3 — 2 bullets inserted")

# ── Step 3: FORMATTING GATE — verify bold labels ──────────────────────────────
doc.save(GUIDE_OUT)

# Re-open and verify
doc2 = Document(GUIDE_OUT)
bold_check_passes = 0
bold_check_total = 4

for p in doc2.paragraphs:
    text = p.text.strip()
    for label in ["What Datadog does:", "Why I want this role:", "What is this role:", "What Datadog is about:"]:
        if text.startswith(label):
            run0 = p.runs[0] if p.runs else None
            is_bold = run0 is not None and run0.bold is True
            print(f"  GATE CHECK: '{label}' → runs[0].bold = {run0.bold if run0 else 'NO RUNS'} → {'PASS' if is_bold else 'FAIL'}")
            if is_bold:
                bold_check_passes += 1

print(f"\n[3] Formatting gate: {bold_check_passes}/{bold_check_total} bold labels verified")
if bold_check_passes < bold_check_total:
    raise AssertionError(f"FORMATTING GATE FAILED: only {bold_check_passes}/{bold_check_total} bold labels confirmed")

# ── Step 4: Append Round 1 coaching section ───────────────────────────────────
doc3 = Document(GUIDE_OUT)

def add_heading(doc_obj, text, level):
    p = doc_obj.add_heading(text, level=level)
    return p

def add_para(doc_obj, text, style="Normal"):
    p = doc_obj.add_paragraph(style=style)
    p.add_run(text)
    return p

def add_bullet(doc_obj, text):
    """Add a bullet paragraph using em-dash prefix (no List Bullet style in template)."""
    p = doc_obj.add_paragraph(style="normal")
    p.add_run("\u2022 " + text)
    return p

# Heading 1
add_heading(doc3, "Round 1 Prep — Pulkit Chandra (Hiring Manager)", level=1)

# Heading 2: What Pulkit is evaluating
add_heading(doc3, "What Pulkit is evaluating", level=2)
add_para(
    doc3,
    "Per Taylor's email, this is a conversational interview — not a case study. Pulkit wants to understand how you "
    "think and how you've worked. Five themes from the prep note: (1) why you're here / why Datadog specifically, "
    "(2) real partner or developer relationships you've owned end-to-end, (3) observability stack knowledge and "
    "architecture decisions, (4) code quality and integration design advising, (5) times you pushed back or "
    "influenced product/eng based on external signal."
)

# Heading 2: Theme-by-theme prep
add_heading(doc3, "Theme-by-theme prep", level=2)

add_bullet(
    doc3,
    "Why you're here / why Datadog: Lead with the transition story — you've been driving large-scale technical "
    "programs at Microsoft, you're most energized at the intersection of technical depth and external impact, and "
    "Datadog's Partner TSE role is the specific combination of technical consulting + product feedback loop that "
    "you're looking for. Don't just say 'I like the product' — say what you know: Datadog is the platform that won "
    "the observability consolidation war, the integration ecosystem is how they extend it, and you want to be at "
    "that growth edge."
)
add_bullet(
    doc3,
    "Real partner/developer relationships end-to-end: Your best story is the GDOT platform adoption pivot "
    "(Section 2 Q3/Q6) — you owned the relationship with local data center teams and the central platform team, "
    "diagnosed why users weren't adopting the tool, and fixed it. Map it explicitly: 'the local operators were my "
    "external developers, the GDOT team was my internal product partner, and the problem was a friction point in "
    "the platform experience.' That's the exact TSE motion."
)
add_bullet(
    doc3,
    "Observability stack knowledge: Know these cold — Metrics (counters, gauges, histograms), Logs (structured vs "
    "unstructured, pipelines), Traces (distributed tracing, spans, APM), OpenTelemetry (OTEL — the open standard "
    "for instrumentation that Datadog supports natively). Agent-based vs API-based integrations: agent runs on the "
    "host and pulls data directly; API-based polls a remote endpoint. OAuth flows: how partner integrations "
    "authenticate. You don't need to be a software engineer — you need to know how the pieces connect so you can "
    "advise on architecture decisions."
)
add_bullet(
    doc3,
    "Code quality / integration design advising: The role does code reviews against Datadog's Quality Rubric. You "
    "won't be asked to write code in the interview, but Pulkit may ask how you'd approach a situation where a "
    "partner's integration works but doesn't meet platform standards. Answer: you start with curiosity (what was "
    "their constraint?), then explain the standard and why it matters for the ecosystem, then collaborate on a path "
    "to compliance rather than just blocking them. Your Section 2 Q2 story (disagreeing with engineering leads, "
    "data-first approach) maps well here."
)
add_bullet(
    doc3,
    "Pushing back / influencing product: Your strongest story here is Q6 (influencing without authority) — you had "
    "no direct authority over the GDOT product team or local operators, but you packaged external feedback into a "
    "product brief, got the feature shipped, and then drove adoption from the ground up. That is textbook "
    "TSE-to-internal-product influence. Also: the AI drill agent story shows you identified a manual bottleneck, "
    "built the case, and shipped automation — that's the 'identify friction, feed it back to eng' motion too."
)

# Heading 2: Datadog observability cheat sheet
add_heading(doc3, "Datadog observability cheat sheet", level=2)
add_para(doc3, "Know these cold:")

add_bullet(
    doc3,
    "Metrics: numerical measurements over time (CPU %, request count, latency p99). Counters go up; gauges go up "
    "and down; histograms show distributions."
)
add_bullet(
    doc3,
    "Logs: text records of events. Datadog ingests, parses, and indexes them. Key concept: structured logs (JSON) "
    "are easier to query than unstructured."
)
add_bullet(
    doc3,
    "Traces / APM: distributed tracing across microservices. A trace = one request's journey; spans = individual "
    "operations within it. Lets you find where latency lives."
)
add_bullet(
    doc3,
    "OpenTelemetry (OTEL): vendor-neutral open standard for instrumentation. Datadog supports it natively — "
    "partners can instrument with OTEL and ship to Datadog. This is the modern default for new integrations."
)
add_bullet(
    doc3,
    "Integration Developer Platform (IDP): Datadog's platform for third-party developers to build and publish "
    "integrations. Integrations live in integrations-extras (community) or the Datadog Marketplace (commercial). "
    "The TSE owns the quality gate."
)
add_bullet(
    doc3,
    "Datadog Agent: lightweight process that runs on the host, collects metrics/logs/traces, and ships them to "
    "Datadog. Many integrations are agent-based (agent pulls from a local service). Contrast: API-based "
    "integrations poll a remote API — no agent needed but less real-time."
)

# Heading 2: Questions to ask Pulkit
add_heading(doc3, "Questions to ask Pulkit", level=2)
add_bullet(
    doc3,
    "'What does the integration lifecycle look like end-to-end — from first contact with a partner to publication? "
    "Where does the TSE have the most leverage?'"
)
add_bullet(
    doc3,
    "'What's the biggest friction point partners hit on the IDP today, and how is the team thinking about fixing it?'"
)
add_bullet(
    doc3,
    "'How do TSEs surface product signal back to the engineering team — is there a formal process or is it more ad hoc?'"
)

# Heading 2: Mindset
add_heading(doc3, "Mindset", level=2)
add_para(
    doc3,
    "Taylor said 'be concrete and honest rather than trying to give perfect answers.' Pulkit wants to see how you "
    "think through real situations, not whether you can recite Datadog's docs. Lead with your actual experiences, "
    "map them explicitly to what the TSE role does, and show genuine curiosity about the observability space. You "
    "have the cross-functional execution, the external-relationship stories, and the product feedback loop — make "
    "them visible."
)

doc3.save(GUIDE_OUT)
print("[4] Round 1 coaching section appended and saved")

# ── Step 5: Placeholder check ─────────────────────────────────────────────────
doc_final = Document(GUIDE_OUT)
placeholder_count = 0
for p in doc_final.paragraphs:
    if "[Fill in here" in p.text or "[fill in" in p.text.lower() or "[placeholder" in p.text.lower():
        placeholder_count += 1
        print(f"  PLACEHOLDER FOUND: {p.text!r}")
print(f"[5] Placeholder check: {placeholder_count} placeholders remaining")

# ── Step 6: Zip the bundle ────────────────────────────────────────────────────
bundle_files = [
    (GUIDE_OUT, GUIDE_NAME),
    (JD_MD, "JD.md"),
    (RESUME_PDF, "Cyrus_Shekari_Resume_datadog_7961297_v2.pdf"),
]

with zipfile.ZipFile(ZIP_OUT, "w", zipfile.ZIP_DEFLATED) as zf:
    for src_path, arc_name in bundle_files:
        zf.write(src_path, arc_name)
        print(f"[6] Added to zip: {arc_name} ({src_path.stat().st_size:,} bytes)")

print(f"\n[6] Zip written → {ZIP_OUT}")
print(f"    Zip size: {ZIP_OUT.stat().st_size:,} bytes")

# ── Summary ───────────────────────────────────────────────────────────────────
print("\n" + "="*60)
print("BUILD SUMMARY")
print("="*60)
print(f"  Guide:      {GUIDE_OUT.name}")
print(f"  Bold labels: {bold_check_passes}/4")
print(f"  Placeholders: {placeholder_count}")
print(f"  Zip:         {ZIP_OUT.name}")
print(f"  Status:      {'SUCCESS' if bold_check_passes == 4 and placeholder_count == 0 else 'ISSUES FOUND'}")
