"""Final verification of all rebuilt guides."""
from docx import Document

all_guides = [
    # Full guides
    ("Lance PM (full)", "guides/lance-pm/Lance_PM_Interview_Prep_Guide.docx"),
    ("MediaAlpha TPM (full)", "guides/mediaalpha-tpm/MediaAlpha_TPM_Interview_Prep_Guide.docx"),
    ("New Relic PM (full)", "guides/new-relic-product-manager-log-management/New_Relic_PM_Interview_Prep_Guide.docx"),
    ("Scale AI PM-TE (full)", "guides/scale-ai-pm-te/ScaleAI_Interview_Prep_Guide.docx"),
    ("Scale AI TPM-CV (full)", "guides/scale-ai-tpm-cv/Scale_AI_Interview_Prep_Guide.docx"),
    ("Podium PM (full)", "guides/podium-pm/Podium_Interview_Prep_Guide.docx"),
    ("Mintlify SE (full)", "guides/mintlify-solutions-engineer-post-sales/Mintlify_SE_PostSales_Interview_Prep_Guide.docx"),
    ("Datadog PTSE (full)", "guides/datadog-partner-technology-solutions-engineer/Datadog_PTSE_Interview_Prep_Guide.docx"),
    # Shorter guides
    ("Lance PM (short)", "guides/lance-pm/Lance_PM_Shorter_Interview_Prep_Guide.docx"),
    ("MediaAlpha TPM (short)", "guides/mediaalpha-tpm/MediaAlpha_TPM_Shorter_Interview_Prep_Guide.docx"),
    ("New Relic PM (short)", "guides/new-relic-product-manager-log-management/New_Relic_PM_Shorter_Interview_Prep_Guide.docx"),
    ("Scale AI PM-TE (short)", "guides/scale-ai-pm-te/ScaleAI_Interview_Prep_Guide_Short.docx"),
    ("Scale AI TPM-CV (short)", "guides/scale-ai-tpm-cv/Scale_AI_Interview_Prep_Guide_Short.docx"),
    ("Podium PM (short)", "guides/podium-pm/Podium_Interview_Prep_Guide_Short.docx"),
]

all_ok = True
for name, path in all_guides:
    try:
        doc = Document(path)
        headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith('Heading')]
        placeholders = [p.text.strip() for p in doc.paragraphs if 'Fill in here' in p.text]

        # Check Q2/Q3 bold labels
        bold_issues = []
        in_q2_q3 = False
        q_section = None
        for p in doc.paragraphs:
            txt = p.text.strip()
            if ('Q2:' in txt or 'Q2: ' in txt) and p.style.name.startswith('Heading'):
                q_section = 'Q2'
                in_q2_q3 = True
            elif ('Q3:' in txt or 'Q3: ' in txt) and p.style.name.startswith('Heading'):
                q_section = 'Q3'
                in_q2_q3 = True
            elif p.style.name.startswith('Heading') and in_q2_q3:
                if 'Q4' in txt or 'Q5' in txt or 'Section 2' in txt or 'Section 3' in txt or 'Behavioral' in txt or 'Leaving' in txt:
                    in_q2_q3 = False

            if in_q2_q3 and p.text.strip() and not p.style.name.startswith('Heading') and p.text.strip() != 'Script:' and p.text.strip() != 'Script:.':
                if p.runs:
                    first_run = p.runs[0]
                    if not first_run.bold or ':' not in first_run.text:
                        bold_issues.append(f"{q_section}: '{p.text[:60]}'")

        status_parts = []
        if placeholders:
            status_parts.append(f"PLACEHOLDERS={placeholders}")
            all_ok = False
        if bold_issues:
            status_parts.append(f"BOLD_MISSING={bold_issues}")
            all_ok = False
        if not status_parts:
            status_parts.append("OK")

        print(f"{name}: {' | '.join(status_parts)}")
    except Exception as err:
        print(f"{name}: ERROR - {err}")
        all_ok = False

print(f"\n{'ALL GUIDES CLEAN' if all_ok else 'SOME ISSUES FOUND'}")
