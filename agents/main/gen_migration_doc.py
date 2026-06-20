#!/usr/bin/env python3
"""Generate OpenClaw Fleet Contingency Migration Guide as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/home/azureuser/.openclaw/workspace/OpenClaw_Migration_Guide.docx"

def set_heading_color(paragraph, rgb):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb)

def add_code_block(doc, code_text):
    """Add a shaded code block paragraph."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # dark red for commands
        # shade the background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        # tight spacing
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.left_indent = Cm(0.5)

def add_note(doc, text, icon='💡'):
    """Add a note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f'{icon}  {text}')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x88)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point."""
    style_name = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style_name)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        p.add_run(text).font.size = Pt(10.5)
    else:
        p.add_run(text).font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_step(doc, number, title, detail=None):
    """Add a numbered step."""
    p = doc.add_paragraph()
    r = p.add_run(f'{number}  {title}')
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    if detail:
        pd = doc.add_paragraph()
        pd.add_run(detail).font.size = Pt(10.5)
        pd.paragraph_format.left_indent = Cm(0.8)
        pd.paragraph_format.space_before = Pt(0)
        pd.paragraph_format.space_after = Pt(2)

def add_table(doc, headers, rows):
    """Add a simple table."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()  # spacing after table

def section_heading(doc, number, title, time_estimate=None):
    """Add a section heading."""
    text = f'Section {number}: {title}'
    if time_estimate:
        text += f'  (⏱ {time_estimate})'
    h = doc.add_heading(text, level=1)
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()

def subsection_heading(doc, title):
    """Add a subsection heading."""
    h = doc.add_heading(title, level=2)
    if h.runs:
        h.runs[0].font.size = Pt(12)
        h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def body(doc, text):
    """Add normal body text."""
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    return p

# ============================================================
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_heading('OpenClaw Fleet', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(24)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_heading('Contingency Migration Guide', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
tagline = doc.add_paragraph('🔒  Break-Glass — Read But Don\'t Act Until You Actually Lose Access')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].bold = True
tagline.runs[0].font.size = Pt(11)
tagline.runs[0].font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Last updated: 2026-06-20  |  Current Azure IP: ').font.size = Pt(10)
r = meta.add_run('40.65.93.84')
r.bold = True; r.font.size = Pt(10)
meta.add_run('  |  Target: Hetzner CAX21').font.size = Pt(10)

doc.add_page_break()

# ============================================================
# INTRO
# ============================================================
doc.add_heading('Introduction', level=1).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
intro_box = doc.add_paragraph()
intro_box.add_run(
    'This is a break-glass guide. Don\'t act on anything here until you actually lose access. '
    'Estimated cost if you do need to migrate: ~$25/mo.'
)
intro_box.runs[0].font.size = Pt(11)
intro_box.runs[0].bold = True
pPr = intro_box._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'FFF2CC')
pPr.append(shd)

doc.add_paragraph()
body(doc,
    'This document covers migrating the OpenClaw fleet from a company-sponsored Azure VM + GitHub Copilot '
    'Enterprise to a fully personal setup running on Hetzner with GitHub Copilot Individual. '
    'The current setup costs ~$0 out of pocket (Azure credit + company Copilot). After migration: '
    '~$25/mo total — about $115/mo cheaper than Azure alone.')

doc.add_paragraph()

# Quick reference table
doc.add_heading('Quick Reference', level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
add_table(doc,
    ['Item', 'Current (Company)', 'Post-Migration (Personal)'],
    [
        ['LLM Provider', 'GitHub Copilot Enterprise (SSO)', 'GitHub Copilot Individual ($19/mo)'],
        ['VM', 'Azure Standard_D4as_v7 (~$140/mo)', 'Hetzner CAX21 (~$5.50/mo)'],
        ['GitHub Repo', '✅ cyshek/cyrusclaw-fleet — personal', 'No change needed'],
        ['Discord Bot', '✅ Personal bot token', 'No change needed'],
        ['Public IP', '40.65.93.84', '<new Hetzner IP>'],
        ['Monthly Cost', '~$0 (credits)', '~$25'],
    ]
)

# ============================================================
# SECTION 1: BEFORE YOU LOSE ACCESS
# ============================================================
doc.add_page_break()
section_heading(doc, 1, 'Before You Lose Access', '~10 min')

body(doc, 'Do this while still employed. Once access is cut, some of these become harder or impossible.')

subsection_heading(doc, '1.1 — Sign Up for GitHub Copilot Pro (Personal Account)')
add_bullet(doc, '🖱️  Navigate to: https://github.com/settings/copilot')
add_bullet(doc, 'Select "Pro" ($19/mo) — includes full model catalog (Claude, GPT-5)')
add_bullet(doc, 'Sign in as personal account: cyshek — NOT your corporate account')
add_note(doc, 'The "Individual" tier ($10/mo) also works but may have rate limits on heavy cron activity.', '⚠️')
doc.add_paragraph()

subsection_heading(doc, '1.2 — Verify GitHub PAT is Personal (Not Corporate SSO)')
add_bullet(doc, '🔑  Your GitHub PAT must belong to the cyshek personal account')
add_bullet(doc, 'Run this check from the Azure VM:')
add_code_block(doc, '''curl -s -H "Authorization: token $GITHUB_TOKEN" https://api.github.com/user \\
  | python3 -c "import sys,json; d=json.load(sys.stdin); print(d['login'], d['type'])"
# Expected output: cyshek  User''')
add_bullet(doc, '🔑  If it prints a corporate org or fails → create a new PAT:')
add_bullet(doc, '🖱️  https://github.com/settings/tokens → Generate new token → Update ~/.openclaw/.env', level=1)
doc.add_paragraph()

subsection_heading(doc, '1.3 — Verify Third-Party API Keys Are on Personal Accounts')
add_bullet(doc, '🖱️  RESIDENTIAL_PROXY — log in to the provider portal, confirm it\'s a personal subscription')
add_bullet(doc, '🖱️  CAPSOLVER_API_KEY — https://dashboard.capsolver.com')
add_bullet(doc, '🖱️  TWOCAPTCHA_API_KEY — https://2captcha.com/enterpage')
add_note(doc, 'These keys are in ~/.openclaw/.env on the Azure VM. If any are corporate accounts, you\'ll need personal alternatives.', '🔑')
doc.add_paragraph()

subsection_heading(doc, '1.4 — Take a Backup Snapshot (Optional but Recommended)')
add_bullet(doc, 'Run from the Azure VM (SSH: azureuser@40.65.93.84):')
add_code_block(doc, '''tar -czf ~/openclaw-backup-$(date +%Y%m%d).tar.gz \\
  --exclude='~/.openclaw/cache' --exclude='~/.openclaw/browser' \\
  --exclude='~/.openclaw/logs' --exclude='~/.openclaw/media' \\
  ~/.openclaw/ ~/.config/systemd/user/openclaw-gateway*''')
add_bullet(doc, 'This creates a timestamped tarball in your home dir (~several hundred MB)')
add_bullet(doc, 'Download it locally: scp azureuser@40.65.93.84:~/openclaw-backup-*.tar.gz ~/Downloads/')
doc.add_paragraph()

# ============================================================
# SECTION 2: DAY ZERO
# ============================================================
doc.add_page_break()
section_heading(doc, 2, 'Day Zero', '~15 min')

body(doc,
    'The moment corporate SSO is cut. Your agents are still running on the Azure VM — '
    'they just need to re-authenticate Copilot to your personal account. This is fast.')

add_note(doc, 'The Azure VM stays running regardless of your employment status (it\'s under your personal Azure subscription with credit). You just lose the company Copilot token.', '✅')
doc.add_paragraph()

add_step(doc, '2.1', 'SSH into the existing VM')
add_code_block(doc, 'ssh azureuser@40.65.93.84')
doc.add_paragraph()

add_step(doc, '2.2', 'Re-authenticate GitHub Copilot to your personal cyshek account')
add_code_block(doc, '''openclaw auth login --provider github-copilot
# Follow the device-code flow — it will print a URL + one-time code
# 🖱️  Open: https://github.com/login/device
# Enter the code shown in the terminal
# Sign in as cyshek (PERSONAL account — not corporate)''')
doc.add_paragraph()

add_step(doc, '2.3', 'Verify everything is healthy')
add_code_block(doc, 'openclaw doctor')
add_note(doc, 'Expect all agents to show green. If you see Copilot auth errors, re-run step 2.2.', '✅')
doc.add_paragraph()

add_step(doc, '2.4', '🖱️  Send a test message in your Discord channel')
add_bullet(doc, 'Open your personal Discord server')
add_bullet(doc, 'Send a message to main agent (e.g. "hey, you there?")')
add_bullet(doc, 'Confirm it responds within ~10 seconds')
doc.add_paragraph()

add_step(doc, '2.5', 'Check all cron jobs are still queued')
add_code_block(doc, 'openclaw cron list')
add_note(doc, 'Done. Agents are live on the Azure VM. The Hetzner migration can happen this week at your own pace — no rush.', '✅')
doc.add_paragraph()

# ============================================================
# SECTION 3: MOVE TO HETZNER
# ============================================================
doc.add_page_break()
section_heading(doc, 3, 'Move to Hetzner', '~45 min total')

body(doc,
    'Migrate the fleet from Azure ($140/mo) to Hetzner CAX21 (~$5.50/mo). '
    'Do this within the first week after day zero. The Azure VM stays up as a fallback '
    'until you confirm Hetzner is stable.')

subsection_heading(doc, '3.1 — Provision the Hetzner VM  (⏱ ~10 min)')
add_step(doc, '3.1.1', '🖱️  Create a Hetzner account at https://hetzner.com')
add_bullet(doc, 'Create a new Project named: openclaw-fleet')
doc.add_paragraph()

add_step(doc, '3.1.2', '🖱️  Create a new server with these settings:')
add_table(doc,
    ['Setting', 'Value'],
    [
        ['Server type', 'CAX21 — ARM, 4 vCPU, 8 GB RAM (~€4.51/mo)'],
        ['Image', 'Ubuntu 22.04 LTS'],
        ['Location', 'Falkenstein (EU) or Ashburn (US-East)'],
        ['SSH key', 'Paste your public key (cat ~/.ssh/id_rsa.pub  or  id_ed25519.pub)'],
        ['Firewall', 'Allow TCP 22 inbound only; block all other inbound'],
    ]
)
add_note(doc, 'CAX21 has 8 GB RAM — half of current Azure VM (16 GB). Workable, but monitor. If tight, upgrade to CAX31 (16 GB, ~€10/mo).', '⚠️')
doc.add_paragraph()

add_step(doc, '3.1.3', 'Note the new IP and set it as a variable for the rest of this section:')
add_code_block(doc, 'NEW_IP="<hetzner-ip-here>"   # replace with your actual new IP')
doc.add_paragraph()

add_step(doc, '3.1.4', 'Install OpenClaw on the new VM:')
add_code_block(doc, '''ssh root@$NEW_IP "curl -fsSL https://openclaw.dev/install.sh | sudo bash && \\
  useradd -m -s /bin/bash azureuser"''')
doc.add_paragraph()

subsection_heading(doc, '3.2 — Rsync Data from Azure → Hetzner  (⏱ ~10 min)')
body(doc, 'Run the following from the Azure VM (ssh azureuser@40.65.93.84 first):')
add_code_block(doc, '''# Core data — excludes caches, browser profile, raw logs, media blobs
rsync -avz --progress \\
  --exclude='cache/' --exclude='browser/' --exclude='logs/' \\
  --exclude='media/' --exclude='tmp/' \\
  ~/.openclaw/ azureuser@$NEW_IP:~/.openclaw/

# Custom scripts
rsync -avz ~/.openclaw/bin/ azureuser@$NEW_IP:~/.openclaw/bin/

# Systemd service files
rsync -avz ~/.config/systemd/user/openclaw-gateway* \\
  azureuser@$NEW_IP:~/.config/systemd/user/''')

doc.add_paragraph()
body(doc, 'What gets synced (critical):')
add_table(doc,
    ['Path', 'Purpose'],
    [
        ['~/.openclaw/openclaw.json', 'Main config'],
        ['~/.openclaw/.env', '🔑 API keys / tokens'],
        ['~/.openclaw/secrets.env', '🔑 Gateway token'],
        ['~/.openclaw/workspace/', 'main agent memory, logs, MEMORY.md'],
        ['~/.openclaw/agents/', 'All 8 peer agent workspaces'],
        ['~/.openclaw/shared-memory/', 'Fleet-wide shared memory'],
        ['~/.openclaw/bin/', 'Custom scripts (voice-note.sh, etc.)'],
        ['~/.openclaw/state/openclaw.sqlite', 'Sessions, crons, state DB'],
        ['~/.config/systemd/user/openclaw-gateway*', 'Service files'],
    ]
)
add_note(doc, 'Do NOT copy ~/.openclaw/credentials/github-copilot.token.json — the token is tied to auth flow and must be re-done on the new machine.', '⚠️')
doc.add_paragraph()

subsection_heading(doc, '3.3 — Start OpenClaw on Hetzner  (⏱ ~10 min)')
add_code_block(doc, '''ssh azureuser@$NEW_IP

# Fix permissions on scripts
chmod +x ~/.openclaw/bin/*.sh

# Enable systemd lingering (keeps gateway alive after SSH disconnect)
sudo loginctl enable-linger azureuser

# Load and enable service files
systemctl --user daemon-reload
systemctl --user enable openclaw-gateway.service
systemctl --user enable openclaw-gateway-watchdog.timer

# Re-auth Copilot (MUST redo — token doesn't transfer)
openclaw auth login --provider github-copilot
# 🖱️  Follow device-code flow → sign in as cyshek (personal)

# Start the gateway
systemctl --user start openclaw-gateway.service

# Verify
openclaw doctor
openclaw cron list''')
doc.add_paragraph()

subsection_heading(doc, '3.4 — Parallel Run + Cut Over  (⏱ ~15 min over 24–48h)')

add_step(doc, '3.4.1', 'Run BOTH VMs for 24–48 hours. Watch Hetzner logs for errors:')
add_code_block(doc, 'journalctl --user -u openclaw-gateway -f')
doc.add_paragraph()

add_step(doc, '3.4.2', '🖱️  Update anything that hard-coded the Azure IP (40.65.93.84):')
add_bullet(doc, 'Discord bot webhook settings (if any had the IP hard-coded)')
add_bullet(doc, 'Any cloudflare tunnels or reverse proxies')
add_bullet(doc, 'Check: cat ~/.openclaw/state/cloudflared.log | tail -5')
add_note(doc, 'The Discord bot TOKEN is portable — no change needed. Only hard-coded IP references need updating.', '✅')
doc.add_paragraph()

add_step(doc, '3.4.3', 'When Hetzner is stable, stop the Azure VM (don\'t delete yet):')
add_code_block(doc, '''# Via Azure CLI (from any machine):
az vm deallocate --resource-group personal-agents --name openclaw-vm

# OR: 🖱️  Azure Portal → Virtual Machines → openclaw-vm → Stop''')
doc.add_paragraph()

# ============================================================
# SECTION 4: WRAP UP
# ============================================================
doc.add_page_break()
section_heading(doc, 4, 'Wrap Up (After Confirming Hetzner is Stable)')

body(doc, 'Wait one week. If Hetzner is running clean, delete the Azure VM and cancel the Azure subscription.')

add_step(doc, '4.1', 'Delete the Azure VM and resource group:')
add_code_block(doc, '''# Via CLI:
az group delete --name personal-agents --yes --no-wait

# OR: 🖱️  Azure Portal → Resource Groups → personal-agents → Delete resource group
# Type "personal-agents" to confirm deletion''')
doc.add_paragraph()

add_step(doc, '4.2', '🖱️  Cancel Azure billing / remove payment method if no longer needed:')
add_bullet(doc, 'Navigate to: https://portal.azure.com → Subscriptions')
add_bullet(doc, 'Subscription ID: 4b2a3918-407a-4f0f-902d-4808c58614e3')
add_bullet(doc, 'Cancel or remove payment method if no other Azure services are in use')
doc.add_paragraph()

add_step(doc, '4.3', 'Clean up stale Azure secrets from the Hetzner VM:')
add_code_block(doc, 'rm -rf ~/.openclaw/secrets/azure-sp.env')
doc.add_paragraph()

add_step(doc, '4.4', 'Confirm monthly cost breakdown:')
add_table(doc,
    ['Item', 'Monthly Cost'],
    [
        ['GitHub Copilot Pro', '$19.00'],
        ['Hetzner CAX21 (ARM, 4 vCPU, 8 GB)', '~$5.50 (~€4.51)'],
        ['Misc (proxy, APIs)', '~$2.00'],
        ['TOTAL', '~$26.50'],
        ['vs. Azure today', '~$140–150'],
        ['Monthly savings', '~$115 🎉'],
    ]
)
doc.add_paragraph()

# ============================================================
# APPENDIX: COPILOT RATE LIMITS
# ============================================================
doc.add_page_break()
doc.add_heading('Appendix: If Copilot Rate-Limits You', level=1).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

body(doc,
    'If GitHub Copilot Individual rate-limits you (likely on heavy days with 40–60 subagent runs), '
    'switch to direct Anthropic API. This costs more (~$25–80/mo for LLM) but has no rate limits.')

subsection_heading(doc, 'A.1 — Add the Anthropic API Key')
add_code_block(doc, '''# 🖱️  Get key from: https://console.anthropic.com → API Keys → Create
# 🔑 Then add to ~/.openclaw/.env on the Hetzner VM:
echo 'ANTHROPIC_API_KEY=sk-ant-...' >> ~/.openclaw/.env
echo 'OPENAI_API_KEY=sk-...' >> ~/.openclaw/.env    # for embeddings (~$1/mo)''')
doc.add_paragraph()

subsection_heading(doc, 'A.2 — Update Model Strings in openclaw.json')
body(doc, 'Change all model references from github-copilot/ prefix to anthropic/:')
add_code_block(doc, '''# Edit ~/.openclaw/openclaw.json — change these model strings:

# BEFORE:
#   "primary": "github-copilot/claude-sonnet-4.6"
#   "fallbacks": ["github-copilot/claude-sonnet-4.6", "github-copilot/gpt-5.5"]
#   "imageModel": { "primary": "github-copilot/claude-opus-4.8" }

# AFTER:
#   "primary": "anthropic/claude-sonnet-4-5"
#   "fallbacks": ["anthropic/claude-sonnet-4-5", "anthropic/claude-haiku-3-5"]
#   "imageModel": { "primary": "anthropic/claude-opus-4-5" }

# Also update memorySearch in all agents:
#   "memorySearch": { "provider": "openai", "model": "text-embedding-3-small" }
#   (add OPENAI_API_KEY to .env — text-embedding-3-small is ~$1/mo at this scale)''')
doc.add_paragraph()

subsection_heading(doc, 'A.3 — Restart the Gateway')
add_code_block(doc, 'systemctl --user restart openclaw-gateway.service')
add_note(doc, 'The model changes take effect after restart. Run openclaw doctor to verify.', '✅')
doc.add_paragraph()

subsection_heading(doc, 'A.4 — Anthropic Cost Estimate')
add_table(doc,
    ['Model', 'Input Rate', 'Output Rate', 'Est. Monthly (~30M tokens)'],
    [
        ['Claude Sonnet 4.5', '$3/MTok', '$15/MTok', '~$72–175'],
        ['Claude Haiku 3.5 (fallback)', '$0.80/MTok', '$4/MTok', '~$20–50'],
        ['OpenAI text-embedding-3-small', '$0.02/MTok', 'n/a', '~$1'],
        ['Total + Hetzner CAX21', '', '', '~$80–185/mo'],
    ]
)
add_note(doc, 'Recommendation: Start with Copilot Individual. Only switch to Anthropic direct if you actually hit rate limits.', '💡')

doc.add_paragraph()

# ============================================================
# QUICK REFERENCE CHECKLIST (final page)
# ============================================================
doc.add_page_break()
doc.add_heading('Quick Checklist — Day-Zero + Week-One', level=1).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_heading('Before Termination (While Still Employed)', level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
for item in [
    '☐  Check GITHUB_TOKEN owner (curl -s api.github.com/user) — must print cyshek',
    '☐  Verify RESIDENTIAL_PROXY / CAPSOLVER / TWOCAPTCHA account ownership',
    '☐  🖱️  Sign up for GitHub Copilot Individual on cyshek account (https://github.com/settings/copilot)',
    '☐  (Optional) Take a backup snapshot: tar -czf ~/openclaw-backup-$(date +%Y%m%d).tar.gz ...',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
doc.add_heading('Day Zero (The Moment Access Is Cut)', level=2).runs[0].font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
for item in [
    '☐  ssh azureuser@40.65.93.84',
    '☐  openclaw auth login --provider github-copilot  (re-auth to cyshek personal account)',
    '☐  openclaw doctor  (verify all agents healthy)',
    '☐  Send a test Discord message — confirm agents respond',
    '☐  openclaw cron list  (verify jobs still queued)',
    '☐  ✅ Done — agents are live. Rest can happen this week.',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
doc.add_heading('Week One (Hetzner Migration)', level=2).runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
for item in [
    '☐  🖱️  Create Hetzner account + CAX21 server (Ubuntu 22.04)',
    '☐  Set NEW_IP="<hetzner-ip>"',
    '☐  Install OpenClaw on new VM: ssh root@$NEW_IP "curl -fsSL https://openclaw.dev/install.sh | sudo bash"',
    '☐  Rsync ~/.openclaw/ to Hetzner (see Section 3.2)',
    '☐  openclaw auth login --provider github-copilot  (re-auth on NEW machine)',
    '☐  systemctl --user start openclaw-gateway.service',
    '☐  openclaw doctor  (verify all agents healthy on Hetzner)',
    '☐  Run both VMs 24–48h in parallel — watch for errors',
    '☐  Cut over — stop Azure VM: az vm deallocate -g personal-agents -n openclaw-vm',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
doc.add_heading('Week Two (Cleanup)', level=2).runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
for item in [
    '☐  Confirm Hetzner stable for 1 full week',
    '☐  az group delete --name personal-agents --yes --no-wait',
    '☐  🖱️  Cancel Azure billing: https://portal.azure.com → Subscriptions → 4b2a3918-...',
    '☐  rm -rf ~/.openclaw/secrets/azure-sp.env  (stale Azure creds)',
    '☐  🎉  Enjoy ~$115/mo in savings',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

doc.add_paragraph()
doc.add_heading('Optional: If Copilot Rate-Limits You', level=2).runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
for item in [
    '☐  🔑  Get ANTHROPIC_API_KEY from https://console.anthropic.com',
    '☐  🔑  Get OPENAI_API_KEY for embeddings (~$1/mo)',
    '☐  Edit ~/.openclaw/openclaw.json — change model strings (see Appendix A.2)',
    '☐  systemctl --user restart openclaw-gateway.service',
]:
    p = doc.add_paragraph(item)
    p.runs[0].font.size = Pt(10.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

# Footer note
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run('Azure subscription ID: 4b2a3918-407a-4f0f-902d-4808c58614e3  |  Current IP: 40.65.93.84  |  GitHub: github.com/cyshek')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
fr.font.italic = True

# ============================================================
doc.save(OUTPUT_PATH)
print(f"✅ Saved: {OUTPUT_PATH}")

import os
size = os.path.getsize(OUTPUT_PATH)
print(f"📄 File size: {size:,} bytes ({size/1024:.1f} KB)")
