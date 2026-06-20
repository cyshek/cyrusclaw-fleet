#!/usr/bin/env python3
"""Build a .docx of the Anduril FDE interview prep with a clickable TOC.

Each question becomes a bookmarked heading; the TOC at the top has internal
hyperlinks (real bookmark jumps) so you can click straight to a question's script.
Works in Word and in Google Docs (upload .docx -> opens as native Doc, links intact).
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "output/anduril-fde-interview-prep.txt"
OUT = "output/anduril-fde-interview-prep.docx"

# ---------- low-level helpers for bookmarks + internal hyperlinks ----------
_bm_id = [0]

def add_bookmark(paragraph, name):
    bid = str(_bm_id[0]); _bm_id[0] += 1
    start = OxmlElement('w:bookmarkStart'); start.set(qn('w:id'), bid); start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), bid)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)

def add_internal_link(paragraph, anchor, text):
    h = OxmlElement('w:hyperlink'); h.set(qn('w:anchor'), anchor)
    r = OxmlElement('w:r'); rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '1155CC'); rpr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rpr.append(u)
    r.append(rpr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); h.append(r); paragraph._p.append(h)

# ---------- parse the source into blocks ----------
with open(SRC) as f:
    lines = f.read().split('\n')

# A "question" is a heading line that sits between ==== rules.
# Section banners are the ###### blocks.
blocks = []  # list of dicts: {type:'section'|'question'|'body', text/title, lines}
i = 0
n = len(lines)

def is_rule(s): return set(s.strip()) == {'='} and len(s.strip()) >= 10
def is_hash(s): return s.startswith('####')

cur_body = []
def flush_body():
    global cur_body
    if cur_body:
        blocks.append({'type': 'body', 'lines': cur_body})
        cur_body = []

while i < n:
    line = lines[i]
    # Section banner: line of #, a "#  SECTION..." line, line of #
    if is_hash(line) and i + 1 < n and lines[i+1].startswith('#') and 'SECTION' in lines[i+1]:
        flush_body()
        title = lines[i+1].lstrip('# ').strip()
        # consume the 3 banner lines
        j = i + 2
        while j < n and is_hash(lines[j]):
            j += 1
        blocks.append({'type': 'section', 'title': title})
        i = j
        continue
    # Question heading: rule / title / rule
    if is_rule(line) and i + 2 < n and is_rule(lines[i+2]):
        flush_body()
        title = lines[i+1].strip()
        blocks.append({'type': 'question', 'title': title})
        i += 3
        continue
    cur_body.append(line)
    i += 1
flush_body()

# ---------- build the doc ----------
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10.5)

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Anduril — Forward Deployed Engineer (Mission Autonomy)"); r.bold = True; r.font.size = Pt(16)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Interview Prep — Cyrus Shekari  ·  Technical Operations Engineer  ·  Costa Mesa, CA")
r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
doc.add_paragraph()

# assign anchors to questions
qcount = 0
for b in blocks:
    if b['type'] == 'question':
        qcount += 1
        b['anchor'] = f"q{qcount}"

# TOC
toc_h = doc.add_paragraph(); rr = toc_h.add_run("CONTENTS — click to jump"); rr.bold = True; rr.font.size = Pt(13)
rr.font.color.rgb = RGBColor(0x11,0x11,0x11)
doc.add_paragraph().add_run("").font.size = Pt(2)

for b in blocks:
    if b['type'] == 'section':
        p = doc.add_paragraph()
        rr = p.add_run(b['title']); rr.bold = True; rr.font.size = Pt(10.5)
        rr.font.color.rgb = RGBColor(0x33,0x33,0x33)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    elif b['type'] == 'question':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        bullet = p.add_run("•  "); bullet.font.color.rgb = RGBColor(0x99,0x99,0x99)
        add_internal_link(p, b['anchor'], b['title'])

doc.add_page_break()

# Body content
def add_body(body_lines):
    # collapse leading/trailing blank lines
    while body_lines and not body_lines[0].strip(): body_lines.pop(0)
    while body_lines and not body_lines[-1].strip(): body_lines.pop()
    for ln in body_lines:
        stripped = ln.rstrip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if not stripped:
            p.add_run("").font.size = Pt(4)
            continue
        if stripped.startswith('BEATS'):
            r = p.add_run(stripped); r.bold = True; r.font.color.rgb = RGBColor(0xB0,0x57,0x00)
            r.font.size = Pt(9.5)
        elif stripped.startswith('SCRIPT'):
            r = p.add_run(stripped); r.bold = True; r.font.color.rgb = RGBColor(0x0B,0x57,0x94)
        elif stripped.startswith('"'):
            r = p.add_run(stripped); r.font.size = Pt(10.5); r.italic = True
        elif stripped.startswith('#') or set(stripped) == {'='}:
            continue  # drop stray rules/hashes inside body
        else:
            r = p.add_run(stripped); r.font.size = Pt(10)

for b in blocks:
    if b['type'] == 'section':
        # Real Heading 1 -> Google Docs renders a native heading: one-click anchor + outline entry
        p = doc.add_heading(b['title'], level=1)
        add_bookmark(p, b.get('anchor', ''))
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x0B,0x3D,0x6B)
    elif b['type'] == 'question':
        # Real Heading 2 -> one-click jump target in Google Docs (kills the extra 'Bookmark' click)
        p = doc.add_heading(b['title'], level=2)
        add_bookmark(p, b['anchor'])
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x11,0x11,0x11)
    elif b['type'] == 'body':
        add_body(list(b['lines']))

doc.save(OUT)
print("wrote", OUT, "| questions:", qcount)
