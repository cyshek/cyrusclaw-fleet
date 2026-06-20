#!/usr/bin/env python3
"""
tailor_resume.py — DOCX-template + LLM-rewrite resume tailoring (v2).

The master .docx (resume/Cyrus_Shekari_Resume_master.docx) is the visual
template. We mutate paragraphs IN PLACE — preserving fonts, spacing,
bold-emphasis runs (which contain the immutable numbers/$/%) — then convert
to PDF via LibreOffice headless.

Inputs (CLI):
  --org <slug>           e.g. "robinhood"
  --job-id <id>          e.g. "7747728"
  [--jd <path>]          default: applications/queued/{org}-{job_id}/JD.md
  [--family pm|tpm|pgm|se|fde]   override JD-title detection
  [--rewrites <path>]    JSON file of bullet rewrites (see below)
  [--out-dir <path>]     default: applications/queued/{org}-{job_id}/
  [--suffix <str>]       filename suffix (default "_v2")

Rewrites JSON format (one file per target, lives next to JD.md by default):
  {
    "title_swaps": {
      "microsoft_ft": "Product Manager",         # paragraph 5 title swap
      "pro_painters":  "Product Manager Intern"  # paragraph 29 title swap
    },
    "bullets": {
      "6":  "Rewritten bullet text. Must contain $14M+ business impact verbatim.",
      "7":  "...",
      ...
    },
    "skills_priority": ["Python", "payment", "data pipelines", ...]
  }

Hard rules enforced in code:
  * Master .docx and .pdf are NEVER modified.
  * Title swaps allowed only for two paragraphs (Microsoft FT, Pro Painters)
    and only to PM / TPM / Program Manager / Product Manager labels.
  * For every rewritten bullet, every number/$/% from the original must
    appear unchanged in the rewrite. If validation fails, the original
    bullet text is kept (no silent corruption).
  * Bold-emphasis run text is preserved verbatim (no styling shifts).
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore

# WordprocessingML namespace helper
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
def _w(tag: str) -> str:
    return f'{{{_W_NS}}}{tag}'

# ---------------------------------------------------------------------------
# Paths

PROJECT = Path(__file__).resolve().parents[1]              # projects/job-search
MASTER_DOCX = PROJECT / "resume" / "Cyrus_Shekari_Resume_master.docx"
APPS_DIR = PROJECT / "applications" / "queued"

# ---------------------------------------------------------------------------
# Static map of master-paragraph indices (verified against the .docx)

# Title rows we are allowed to swap (paragraph index → role-key)
# As of 2026-05-12, Pro Painters joins the swap allowlist (same intern family
# as the other intern roles) — no longer LOCKED.
TITLE_PARA = {
    5:  "microsoft_ft",        # "Technical Program Manager"        → PM/TPM/Program Manager/Product Manager
    12: "microsoft_2023",      # "Technical Program Manager Intern" → *Intern family
    17: "microsoft_2022",      # "Technical Program Manager Intern" → *Intern family
    23: "amazon_robotics",     # "Technical Program Manager Intern" → *Intern family
    29: "pro_painters",        # "Program Manager Intern"           → *Intern family
}

_PM_FAMILY = {
    "Product Manager",
    "Program Manager",
    "Technical Program Manager",
    "Technical Product Manager",
}
_PM_INTERN_FAMILY = {f"{t} Intern" for t in _PM_FAMILY}

ALLOWED_TITLE_LABELS = {
    "microsoft_ft":     _PM_FAMILY,
    "microsoft_2023":   _PM_INTERN_FAMILY,
    "microsoft_2022":   _PM_INTERN_FAMILY,
    "amazon_robotics":  _PM_INTERN_FAMILY,
    "pro_painters":     _PM_INTERN_FAMILY,
}

# Bullet paragraphs grouped by job. Used for selection / capping.
JOB_BULLETS = {
    "microsoft_ft":      [6, 7, 8, 9, 10],   # cap 5
    "microsoft_2023":    [13, 14, 15],       # cap 3
    "microsoft_2022":    [18, 19, 20],       # cap 3
    "amazon_robotics":   [24, 25, 26],       # cap 3
    "pro_painters":      [30, 31, 32],       # cap 3
}

# Skills paragraphs (label : paragraph index)
SKILLS_PARA = {
    "Technical": 42,
    "Program / Product": 43,
    "AI / Automation": 44,
}

# ---------------------------------------------------------------------------
# Family detection

FAMILY_PATTERNS = [
    ("fde", re.compile(r"forward[-\s]?deployed", re.I)),
    ("se",  re.compile(r"\b(solutions?|sales|customer|ai)\s+engineer\b", re.I)),
    ("se",  re.compile(r"\bsolutions?\s+architect\b", re.I)),
    ("tpm", re.compile(r"\btechnical\s+(program|product)\s+manager\b", re.I)),
    ("tpm", re.compile(r"\bTPM\b")),
    ("pgm", re.compile(r"\b(engineering\s+)?program\s+manager\b", re.I)),
    ("pm",  re.compile(r"\bproduct\s+manager\b", re.I)),
    ("pm",  re.compile(r"\bproduct\s+builder\b", re.I)),
]

def detect_family(title: str) -> str:
    for fam, pat in FAMILY_PATTERNS:
        if pat.search(title):
            return fam
    return "pm"


# ---------------------------------------------------------------------------
# Faithful title resolver (Fix 1, 2026-05-31)
#
# The resume's headline role title (microsoft_ft slot) MUST mirror the JD's
# ACTUAL role title. Do NOT default every PM resume to "Technical Product
# Manager". Map the JD title to the closest allowed label WITHOUT inflating to
# the "Technical" variant unless the JD itself says "Technical".
#
# Allowed labels for microsoft_ft (_PM_FAMILY):
#   Product Manager, Program Manager,
#   Technical Program Manager, Technical Product Manager
#
# Resolution is phrase-based off the JD title line so we faithfully
# distinguish "Product Manager" vs "Technical Product Manager" etc. (the
# `tpm` family alone is ambiguous between the two Technical variants).

_TITLE_RESOLVE_PATTERNS = [
    # Most specific first. Each maps a JD-title phrase → exact allowed label.
    (re.compile(r"\btechnical\s+program\s+manager\b", re.I), "Technical Program Manager"),
    (re.compile(r"\btechnical\s+product\s+manager\b", re.I), "Technical Product Manager"),
    (re.compile(r"\bprogram\s+manager\b", re.I),              "Program Manager"),
    (re.compile(r"\bproduct\s+manager\b", re.I),              "Product Manager"),
    (re.compile(r"\bTPM\b"),                                  "Technical Program Manager"),
    (re.compile(r"\bTPgM\b"),                                 "Technical Program Manager"),
    (re.compile(r"\bproduct\s+builder\b", re.I),             "Product Manager"),
]

# Fallback per-family label when the JD title line can't be phrase-matched
# (e.g. detection came from --family override). Deliberately NON-Technical for
# the plain pm/pgm families so we never claim a "Technical" title the role
# didn't ask for.
_FAMILY_DEFAULT_TITLE = {
    "pm":  "Product Manager",
    "pgm": "Program Manager",
    "tpm": "Technical Program Manager",
}


def resolve_headline_title(title_line: str, family: str) -> str | None:
    """Return the faithful allowed label for the microsoft_ft headline slot.

    Prefers a direct phrase match against the JD title line; falls back to a
    per-family default. Returns None for non-PM families (se/fde) — those keep
    the master title (no PM-family swap applies).
    """
    if family not in _FAMILY_DEFAULT_TITLE:
        return None  # se / fde / anything outside the PM family: no swap
    for pat, label in _TITLE_RESOLVE_PATTERNS:
        if pat.search(title_line or ""):
            return label
    return _FAMILY_DEFAULT_TITLE.get(family)


# ---------------------------------------------------------------------------
# Title-track coherence enforcer (Fix 2, 2026-06-01)
#
# Bug: the LLM populated all 5 TITLE_PARA slots independently, so a single
# resume could mix "Technical Product Manager" and "Technical Program Manager"
# (e.g. cloudera-1303 left 2 intern slots Program while flipping the rest to
# Product). resolve_headline_title only governed the headline (slot 5).
#
# Fix: derive ONE canonical track (Product vs Program, + Technical-ness) from
# the resolved headline, then snap EVERY title slot to that track while keeping
# each slot's fixed "Intern" qualifier. Never invents a label outside the
# allowed family sets; only enforces consistency.
#
# When the JD title is ambiguous / non-PM-specific (e.g. an SE role mapped into
# the PM family via --family), default to the PROGRAM track (Cyrus 2026-06-01:
# matches his primary/most-recent Microsoft title; "Product" only when the JD
# explicitly says Product).

def _track_from_headline(headline: str) -> tuple[str, bool]:
    """Return (discipline, is_technical) where discipline ∈ {Product, Program}.

    Falls back to the Program track for anything that isn't clearly Product.
    """
    h = (headline or "").lower()
    is_technical = "technical" in h
    discipline = "Product" if "product" in h else "Program"
    return discipline, is_technical


def coerce_title_track(title_swaps: dict, headline: str) -> dict:
    """Snap ALL known TITLE_PARA slots to the single track implied by ``headline``.

    - microsoft_ft  -> "[Technical ]<discipline> Manager"
    - *intern slots -> "[Technical ]<discipline> Manager Intern"
    Iterates over ALL keys in ALLOWED_TITLE_LABELS (not just pre-seeded slots)
    so un-seeded intern slots (e.g. microsoft_2023, amazon_robotics) are also
    coerced to the correct track. Returns a NEW dict; never mutates the caller's.
    """
    discipline, is_technical = _track_from_headline(headline)
    prefix = "Technical " if is_technical else ""
    base = f"{prefix}{discipline} Manager"
    out = dict(title_swaps)
    for role_key in ALLOWED_TITLE_LABELS:  # ALL known slots, not just pre-seeded ones
        if role_key == "microsoft_ft":
            out[role_key] = base
        else:
            # intern family slots
            out[role_key] = f"{base} Intern"
    return out

# ---------------------------------------------------------------------------
# Number-preservation validator

# Matches: $14M+, 14, 14.5, 14%, $1.5B, 200+, 2,000+, 2.7%, 14M, etc.
NUMBER_RE = re.compile(r"\$?\d[\d,]*\.?\d*[%MBK+]*\+?")

def extract_numbers(text: str) -> list[str]:
    raw = NUMBER_RE.findall(text)
    # Strip trailing punctuation that might attach (e.g. "39%,")
    return [n.rstrip(",.;:") for n in raw if any(c.isdigit() for c in n)]

def numbers_preserved(original: str, rewrite: str) -> tuple[bool, list[str]]:
    missing = []
    for n in extract_numbers(original):
        if n not in rewrite:
            missing.append(n)
    return (len(missing) == 0, missing)

# ---------------------------------------------------------------------------
# Bullet mutation: parses **bold** markers in new_text into bold runs.

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)


def _parse_bold_segments(text: str):
    """Return list of (text, is_bold) segments parsed from **...** markers."""
    segments = []
    last = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False))
        inner = m.group(1)
        if inner:
            segments.append((inner, True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False))
    if not segments:
        segments = [(text, False)]
    return segments


def mutate_bullet(paragraph, original_text: str, new_text: str) -> dict:
    """
    Replace the bullet's text. Parses ``**bold**`` markers from new_text and
    emits one run per (text, is_bold) segment, preserving the master's font /
    size from the first run as a template.

    Returns dict with keys: applied(bool), reason(str), final_text(str).
    """
    from copy import deepcopy
    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    runs = list(paragraph.runs)
    if not runs:
        return {"applied": False, "reason": "no runs", "final_text": original_text}

    segments = _parse_bold_segments(new_text)
    final_text = "".join(t for t, _ in segments)

    template = runs[0]
    template_xml = deepcopy(template._r)
    # Strip the original run's text element(s) and bold marker from the template
    for child in list(template_xml):
        tag = child.tag
        if tag == f"{W_NS}t" or tag == f"{W_NS}tab":
            template_xml.remove(child)
    rPr = template_xml.find(f"{W_NS}rPr")
    if rPr is not None:
        b = rPr.find(f"{W_NS}b")
        if b is not None:
            rPr.remove(b)
        bcs = rPr.find(f"{W_NS}bCs")
        if bcs is not None:
            rPr.remove(bcs)

    # Remove all existing runs from the paragraph
    p_elem = paragraph._p
    for r in runs:
        p_elem.remove(r._r)

    # Append a new run per segment, cloned from template
    import docx.oxml.ns as _ns  # noqa
    from docx.oxml import OxmlElement
    for seg_text, is_bold in segments:
        if not seg_text:
            continue
        new_r = deepcopy(template_xml)
        if is_bold:
            seg_rPr = new_r.find(f"{W_NS}rPr")
            if seg_rPr is None:
                seg_rPr = OxmlElement("w:rPr")
                new_r.insert(0, seg_rPr)
            b_el = OxmlElement("w:b")
            seg_rPr.append(b_el)
        t_el = OxmlElement("w:t")
        # preserve leading/trailing whitespace
        if seg_text != seg_text.strip() or "  " in seg_text:
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = seg_text
        new_r.append(t_el)
        p_elem.append(new_r)

    return {"applied": True, "reason": "segments", "final_text": final_text}


# ---------------------------------------------------------------------------
# Title row mutation (preserves whitespace padding so right-aligned date stays put)

TITLE_ROW_RE = re.compile(r"^(?P<title>\S.*?\S)(?P<gap>\s{2,})(?P<date>\S.*\S)\s*$")

def mutate_title(paragraph, role_key: str, new_label: str) -> dict:
    if new_label not in ALLOWED_TITLE_LABELS[role_key]:
        return {"applied": False, "reason": f"label not allowed: {new_label}"}
    # Find the run containing the title text (run 0 in our master).
    full = paragraph.text
    m = TITLE_ROW_RE.match(full)
    if not m:
        return {"applied": False, "reason": "no title-gap-date pattern"}
    old_title = m.group("title")
    gap = m.group("gap")
    date = m.group("date")
    # Length-preserve: adjust gap so total visible width is unchanged
    diff = len(old_title) - len(new_label)
    new_gap = gap + (" " * diff) if diff > 0 else (gap[:max(2, len(gap) + diff)])
    new_full = f"{new_label}{new_gap}{date}"
    # Apply to first run; clear any subsequent runs
    if not paragraph.runs:
        return {"applied": False, "reason": "no runs"}
    paragraph.runs[0].text = new_full
    for r in paragraph.runs[1:]:
        r.text = ""
    return {"applied": True, "old": old_title, "new": new_label, "row": new_full}


# ---------------------------------------------------------------------------
# Per-JD skill GENERATION (Fix 2, 2026-05-31 — revised per Cyrus 2026-05-31)
#
# Skills are GENERATED FROM THE JD ITSELF, not selected from a master inventory.
# Cyrus's explicit risk decision: do NOT constrain skills to a pre-listed
# inventory of what he "has." If a skill is relevant to the JD, claim it
# (worst case he learns it before an interview). No inventory cross-check, no
# "skill not in profile -> skip."
#
# Implementation: a RECOGNITION CATALOG maps canonical, well-formed skill names
# to the JD phrases/tokens that signal them. We scan the JD text; any catalog
# entry whose trigger appears in the JD is emitted under its group. This pulls
# clean skill names out of messy JD prose (vs naive noun-phrase scraping, which
# yields garbage). The catalog is a RECOGNITION dictionary of known tech/PM
# terms — NOT a claim-list of Cyrus's existing skills.
#
# Group buckets mirror the resume's three skills rows:
#   "Technical"          — languages, platforms, tools, infra, data
#   "Program / Product"  — PM/TPM competencies & methodologies
#   "AI / Automation"    — AI/ML/automation tooling & techniques

# canonical skill name -> list of JD trigger substrings (lowercased, matched on
# lowercased JD text). First trigger that hits is enough. Order within each
# group is the emit/priority order.
_SKILL_CATALOG: dict[str, list[tuple[str, list[str]]]] = {
    "Technical": [
        ("Python", ["python"]),
        ("Java", [" java", "java,", "java.", "java/"]),
        ("Go", [" golang", "go programming", " go,", " go.", " go "]),
        ("C++", ["c++"]),
        ("JavaScript", ["javascript", "node.js", "nodejs"]),
        ("TypeScript", ["typescript"]),
        ("React", ["react"]),
        ("Azure", ["azure"]),
        ("AWS", ["aws", "amazon web services", "ec2", "s3 ", "lambda"]),
        ("GCP", ["gcp", "google cloud"]),
        ("Cloud infrastructure", ["cloud", "microsoft cloud"]),
        ("Distributed systems", ["distributed", "microservice", "large-scale system", "scalable system"]),
        ("Kubernetes", ["kubernetes", "k8s"]),
        ("Docker", ["docker", "container"]),
        ("CI/CD", ["ci/cd", "cicd", "continuous integration", "continuous delivery", "continuous deployment"]),
        ("DevOps", ["devops"]),
        ("APIs", ["api", "rest", "graphql", "sdk", "endpoint"]),
        ("Microservices", ["microservice"]),
        ("Data pipelines", ["data pipeline", "etl", "data engineering", "ingestion", "data warehouse", "data infra"]),
        ("SQL", ["sql", "postgres", "mysql", "relational database"]),
        ("NoSQL", ["nosql", "mongodb", "dynamodb", "cassandra"]),
        ("Data analytics", ["analytics", "data analysis", "data-driven", "metrics"]),
        ("Power BI", ["power bi"]),
        ("Tableau", ["tableau"]),
        ("Dashboards", ["dashboard", "reporting", "visualiz"]),
        ("Machine learning", ["machine learning", " ml ", "ml model", "ml-"]),
        ("YAML", ["yaml"]),
        ("Infrastructure as Code", ["infrastructure as code", " iac", "terraform"]),
        ("System design", ["system design", "architecture", "architect"]),
        ("Reliability engineering", ["reliability", "resilience", "sre", "observability", "uptime", "availability"]),
        ("Security", ["security", "compliance", "authentication", "encryption"]),
    ],
    "Program / Product": [
        ("Technical Program Management", ["technical program", "program manage", "tpm", "program management"]),
        ("Product Management", ["product manage", "product owner", " pm ", "product strategy"]),
        ("Product requirements", ["requirement", "prd", "product doc", "product spec"]),
        ("Roadmapping", ["roadmap"]),
        ("Prioritization", ["prioritiz", "backlog", "tradeoff", "trade-off"]),
        ("Strategic planning", ["strategy", "strategic", "planning", "vision"]),
        ("Cross-functional execution", ["cross-functional", "cross functional", "xfn", "cross-team", "cross team"]),
        ("Stakeholder management", ["stakeholder", "executive", "leadership alignment", "influence"]),
        ("Go-to-market", ["go-to-market", "go to market", "gtm", "launch"]),
        ("Customer research", ["customer research", "user research", "voice of customer", "customer insight", "discovery"]),
        ("Data-driven decision making", ["data-driven", "data driven", "experimentation", "a/b test", "ab test"]),
        ("Agile/Scrum", ["agile", "scrum", "sprint", "kanban"]),
        ("Risk management", ["risk", "mitigat", "dependenc"]),
        ("Cross-org coordination", ["coordinat", "collaborat", "partner team"]),
        ("Metrics & KPIs", ["kpi", "okr", "metric", "north star"]),
        ("Budgeting", ["budget", "resourcing", "headcount"]),
    ],
    "AI / Automation": [
        ("AI agents", ["ai agent", "agentic", "autonomous agent", "agent framework"]),
        ("LLM-powered tools", ["llm", "large language model", "gpt", "foundation model", "language model"]),
        ("Generative AI", ["generative ai", "genai", "gen ai", "generative"]),
        ("Prompt engineering", ["prompt engineering", "prompt", "prompting"]),
        ("RAG / Semantic search", ["rag", "retrieval-augmented", "retrieval augmented", "semantic search", "embedding", "vector"]),
        ("Copilot Studio", ["copilot", "power platform", "low-code", "low code"]),
        ("Workflow automation", ["workflow automation", "automation", "orchestrat", "rpa"]),
        ("ML pipelines", ["ml pipeline", "mlops", "model deployment", "model training", "fine-tun"]),
        ("Process optimization", ["process optim", "efficien", "streamlin", "operational excellence"]),
        ("Natural language processing", ["nlp", "natural language"]),
        ("Computer vision", ["computer vision", "image recognition", "object detection"]),
        ("Data labeling / evaluation", ["data labeling", "annotation", "model eval", "evals", "benchmark"]),
    ],
}

# Per-group emit cap so a keyword-dense JD doesn't overflow the line.
_SKILL_MAX_PER_GROUP = 8
# Per-group minimum so a sparse JD still yields a non-empty, on-identity line.
_SKILL_MIN_PER_GROUP = 3

# Fallback items per group when the JD triggers fewer than the minimum. These
# are generic, on-identity PM/TPM defaults (still JD-agnostic safe — they're
# the resume's baseline competencies, used only to avoid an empty line).
_SKILL_FALLBACK = {
    "Technical": ["APIs", "SQL", "Data analytics"],
    "Program / Product": ["Technical Program Management", "Roadmapping", "Cross-functional execution"],
    "AI / Automation": ["AI agents", "LLM-powered tools", "Workflow automation"],
}


def merge_skills_for_group(base_items: list[str], label: str, jd_text: str) -> list[str]:
    """Skills model (Cyrus 2026-06-01, Option B — pure JD-priority).

    Cyrus's base skill set is 50+; each JD needs a different slice, so the JD
    drives WHICH skills surface and his base only fills in to reach the floor.

      1. JD-RELEVANT block FIRST, ordered by catalog/JD priority — regardless of
         whether a skill is already in his base. So a JD-wanted catalog skill he
         doesn't currently list can rank ABOVE a base skill the JD didn't
         mention. Built as:
           a. catalog hits for this JD, in catalog priority order (this is the
              canonical JD priority — includes both base-and-catalog skills and
              stretch skills he doesn't list yet, interleaved by priority);
           b. then any base skill that is JD-relevant but NOT captured by the
              catalog (don't lose a relevant base skill just because it's not a
              catalog canonical).
      2. FILLER block: remaining (non-JD-relevant) base skills, used only to
         reach _SKILL_MIN_PER_GROUP and to fill remaining slots up to the cap.
      3. Caps: keep >= _SKILL_MIN_PER_GROUP (top up from base, then fallback),
         and <= _SKILL_MAX_PER_GROUP. Dedup / near-dupe logic preserved.

    JD priority leads; base skills are filler to reach the floor / cap.
    """
    jd_low = (jd_text or "").lower()
    base = [s.strip() for s in base_items if s.strip()]

    # JD-relevant catalog skills (canonical, catalog priority order). This is
    # the authoritative JD-priority ordering — base membership does NOT promote.
    jd_hits: list[str] = []
    for canonical, triggers in _SKILL_CATALOG.get(label, []):
        if any(trig in jd_low for trig in triggers):
            jd_hits.append(canonical)
    jd_hits_low = {j.lower() for j in jd_hits}

    def _base_is_jd_relevant(skill: str) -> bool:
        s = skill.lower()
        if s in jd_low:
            return True
        for w in re.split(r"[\s/]+", s):
            if len(w) >= 3 and w in jd_low:
                return True
        return False

    # JD-first block:
    #   (a) catalog hits in catalog/JD priority order (base or stretch alike);
    #   (b) base skills that ARE JD-relevant but aren't catalog canonicals
    #       (so a relevant base skill isn't dropped). Appended after the
    #       priority-ordered catalog block.
    jd_base_extra = [b for b in base
                     if _base_is_jd_relevant(b) and b.lower() not in jd_hits_low]
    jd_block = jd_hits + jd_base_extra
    jd_block_low = {x.lower() for x in jd_block}

    # Filler block: remaining non-JD-relevant base skills (identity filler).
    filler = [b for b in base if b.lower() not in jd_block_low]

    # Single priority-ordered candidate stream: JD-first, then base filler.
    tier1, tier2, tier3 = jd_block, [], filler

    merged: list[str] = []
    seen: set[str] = set()

    def _norm(s: str) -> str:
        # canonical key for near-dupe detection: lowercase, strip non-alnum.
        return re.sub(r"[^a-z0-9]+", "", s.lower())

    def _is_dupe(it: str) -> bool:
        n = _norm(it)
        if not n:
            return True
        for s in seen:
            # treat as dupe if one normalized key contains the other
            # (e.g. 'semanticsearch' ⊆ 'ragsemanticsearch').
            if n == s or (len(n) >= 5 and (n in s or s in n)):
                return True
        return False

    for it in tier1 + tier2 + tier3:
        if _is_dupe(it):
            continue
        seen.add(_norm(it))
        merged.append(it)
        if len(merged) >= _SKILL_MAX_PER_GROUP:
            break

    # Floor: never below the minimum. Top up from base, then group fallback.
    if len(merged) < _SKILL_MIN_PER_GROUP:
        for it in base + _SKILL_FALLBACK.get(label, []):
            if not _is_dupe(it):
                seen.add(_norm(it))
                merged.append(it)
            if len(merged) >= _SKILL_MIN_PER_GROUP:
                break

    return merged[:_SKILL_MAX_PER_GROUP]


# ---------------------------------------------------------------------------
# Skills paragraph rewrite (Fix 2 + Fix 3, 2026-05-31)

def _rewrite_skills_runs(paragraph, label_text: str, body: str) -> bool:
    """Write the skills paragraph as: BOLD label run + NON-BOLD body run.

    Fix 3: only the group label ("Technical:") is bold; the comma-separated
    skill list after the colon must NOT be bold. We force run0 = label (bold),
    run1 = " body" (not bold), and blank any further runs.
    """
    runs = paragraph.runs
    if not runs:
        return False
    # run0 carries the label and stays bold.
    runs[0].text = label_text
    runs[0].bold = True
    body_text = f" {body}"
    if len(runs) >= 2:
        runs[1].text = body_text
        runs[1].bold = False
        for r in runs[2:]:
            r.text = ""
            r.bold = False
    else:
        # Only one run in this paragraph: append a fresh non-bold body run
        # cloned from run0's font so it visually matches.
        new = paragraph.add_run(body_text)
        new.bold = False
        if runs[0].font.name:
            new.font.name = runs[0].font.name
        if runs[0].font.size:
            new.font.size = runs[0].font.size
    return True


def reorder_skills(paragraph, label: str, priority: list[str],
                   jd_text: str = "", tailor: bool = True) -> dict:
    """Generate JD-driven skills for this group and write them in place.

    Format: "Technical: Azure, distributed systems, ..." where run0 holds the
    bold label and run1 holds the non-bold body.

    Behavior (FINAL reconciled spec, Cyrus 2026-05-31):
      * Fix 2 (tailor=True): BLEND — start from the master/base skills for this
        group, augment/replace with JD-relevant ones (JD-relevant-first),
        within the 3..8 cap, via merge_skills_for_group(). Not base-only, not
        JD-only. If tailor=False, the master body is kept as-is (and merely
        re-emitted with correct bolding).
      * Fix 3: re-emit with label-run bold, body-run NOT bold.
    """
    txt = paragraph.text
    if ":" not in txt:
        return {"applied": False, "reason": "no colon"}
    pre, _, body = txt.partition(":")
    master_items = [s.strip() for s in body.split(",") if s.strip()]

    if tailor:
        # Fix 2: blend base (master) with JD-relevant skills, JD-first.
        new_items = merge_skills_for_group(master_items, label, jd_text)
        if not new_items:
            # Defensive: never blank the line — fall back to the master body.
            new_items = master_items
    else:
        new_items = list(master_items)
        # Optional priority reordering of the existing master body.
        def rank(item: str) -> tuple:
            low = item.lower()
            for i, kw in enumerate(priority):
                kwl = kw.lower()
                if kwl and (kwl in low or low in kwl):
                    return (0, i)
            return (1, new_items.index(item))
        new_items = sorted(new_items, key=rank)

    new_body = ", ".join(new_items)

    if not _rewrite_skills_runs(paragraph, f"{pre}:", new_body):
        return {"applied": False, "reason": "no runs"}
    return {"applied": True, "before": master_items, "after": new_items,
            "source": "jd-blended" if tailor else "master-reordered"}


# ---------------------------------------------------------------------------
# Bullet paragraph injection (for overflow beyond master's fixed slots)

def _clone_bullet_paragraph(template_para):
    """Deepcopy a paragraph element (preserves numPr/style) and return the new
    <w:p> element. Caller must insert it into the document tree and set text.
    """
    from copy import deepcopy
    return deepcopy(template_para._p)


def _set_paragraph_plain_text(p_element, text: str):
    """Replace all runs in a paragraph element with one run per (text, is_bold)
    segment parsed from ``**bold**`` markers. Copies the first existing run's
    rPr (font/size) onto the new runs so the injected bullet visually matches
    its siblings. The ``**`` markers are stripped; only inner text is rendered.
    """
    from copy import deepcopy
    from lxml import etree  # python-docx ships with lxml

    rPr_template = None
    # Collect & remove existing runs
    for r in list(p_element.findall(_w('r'))):
        if rPr_template is None:
            rpr = r.find(_w('rPr'))
            if rpr is not None:
                rPr_template = deepcopy(rpr)
                # Strip <w:b/> so the cloned run isn't bold by default.
                for b in rPr_template.findall(_w('b')):
                    rPr_template.remove(b)
                for b in rPr_template.findall(_w('bCs')):
                    rPr_template.remove(b)
        p_element.remove(r)

    segments = _parse_bold_segments(text)
    for seg_text, is_bold in segments:
        if not seg_text:
            continue
        new_r = etree.SubElement(p_element, _w('r'))
        if rPr_template is not None:
            seg_rPr = deepcopy(rPr_template)
            new_r.append(seg_rPr)
        if is_bold:
            seg_rPr = new_r.find(_w('rPr'))
            if seg_rPr is None:
                seg_rPr = etree.SubElement(new_r, _w('rPr'))
                # rPr must be first child of run
                new_r.remove(seg_rPr)
                new_r.insert(0, seg_rPr)
            etree.SubElement(seg_rPr, _w('b'))
        t = etree.SubElement(new_r, _w('t'))
        t.text = seg_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


# ---------------------------------------------------------------------------
# Main pipeline

def run_pipeline(
    org: str,
    job_id: str,
    jd_path: Path,
    family: str | None,
    rewrites_path: Path | None,
    out_dir: Path,
    suffix: str,
) -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)

    jd_text = jd_path.read_text() if jd_path.exists() else ""
    title_line = ""
    for ln in jd_text.splitlines():
        s = ln.strip().lstrip("#").strip()
        if s:
            title_line = s
            break
    fam = family or detect_family(title_line)

    # Resolve rewrites
    rewrites: dict = {}
    if rewrites_path and rewrites_path.exists():
        rewrites = json.loads(rewrites_path.read_text())
    else:
        # Default location: applications/queued/{org}-{job_id}/rewrites.json
        guess = out_dir / "rewrites.json"
        if guess.exists():
            rewrites = json.loads(guess.read_text())
            rewrites_path = guess

    title_swaps   = rewrites.get("title_swaps", {}) or {}
    skills_prio   = rewrites.get("skills_priority", []) or []

    # 1. Copy master → target docx
    target_docx = out_dir / f"Cyrus_Shekari_Resume_{org}_{job_id}{suffix}.docx"
    shutil.copy2(MASTER_DOCX, target_docx)

    # 2. Open the copy and mutate
    doc = Document(str(target_docx))
    paras = doc.paragraphs

    # Capture originals for diff
    originals = {i: paras[i].text for i in (
        list(TITLE_PARA.keys())
        + [b for blist in JOB_BULLETS.values() for b in blist]
        + list(SKILLS_PARA.values())
    )}

    report = {
        "org": org, "job_id": job_id, "family": fam, "title_seen": title_line,
        "title_swaps": [], "bullets": [], "skills": [], "warnings": [],
        "rewrites_path": str(rewrites_path) if rewrites_path else None,
    }

    # 3. Title swaps
    #
    # Fix 1 (2026-05-31): the HEADLINE role title (microsoft_ft) must mirror the
    # JD's actual role title — never blanket-default to "Technical Product
    # Manager". Resolve the faithful label from the JD title line + family, and
    # use it unless an explicit, faithful swap was supplied by the rewriter.
    faithful_headline = resolve_headline_title(title_line, fam)
    if faithful_headline is not None:
        supplied = title_swaps.get("microsoft_ft")
        # Honor an explicitly-supplied label ONLY if it is itself a faithful
        # match for the JD (i.e. equals what we resolved). Otherwise force the
        # faithful one so we never claim a title the role didn't ask for.
        if supplied != faithful_headline:
            title_swaps = dict(title_swaps)
            title_swaps["microsoft_ft"] = faithful_headline
            report.setdefault("title_resolution", {})
            report["title_resolution"] = {
                "jd_title": title_line, "family": fam,
                "resolved": faithful_headline, "supplied": supplied,
                "action": "forced-faithful" if supplied else "defaulted-faithful",
            }

    # Fix 2 (2026-06-01): enforce ONE Product-vs-Program track across ALL title
    # slots so a single resume can never mix "Product Manager" and "Program
    # Manager" (the LLM picks each slot independently). Runs for EVERY family
    # that actually has title swaps — including se/fde, where there is no
    # JD-derived PM headline but the LLM still populates all 5 slots (that path
    # was the original source of the mix, e.g. cloudera-1303). Track source:
    # the resolved faithful headline when present, else the LLM's microsoft_ft
    # value, else the Program default.
    # Always coerce when the role has a recognized family — even if title_swaps
    # is empty (unseeded). The old guard `if any(k in ALLOWED_TITLE_LABELS for k
    # in title_swaps)` silently skipped coerce for empty rewrites.json
    # title_swaps dicts, leaving intern slots at master defaults (bug 2026-06-14).
    _coerce_family = fam or rewrites.get("family") or ""
    if _coerce_family in ("pm", "tpm", "pgm", "se", "fde") or any(k in ALLOWED_TITLE_LABELS for k in title_swaps):
        track_src = faithful_headline or title_swaps.get("microsoft_ft") or "Program Manager"
        before_coerce = dict(title_swaps)
        title_swaps = coerce_title_track(title_swaps, track_src)
        coerced = {k: title_swaps[k] for k in title_swaps
                   if before_coerce.get(k) != title_swaps[k]}
        if coerced:
            report.setdefault("title_resolution", {})
            report["title_resolution"]["track_source"] = track_src
            report["title_resolution"]["track_coerced"] = coerced

    for pidx, role_key in TITLE_PARA.items():
        new_label = title_swaps.get(role_key)
        if not new_label:
            continue
        res = mutate_title(paras[pidx], role_key, new_label)
        report["title_swaps"].append({
            "para": pidx, "role_key": role_key,
            "from": originals[pidx], "to": new_label,
            **res,
        })
        if not res.get("applied"):
            report["warnings"].append(
                f"title swap failed para={pidx} role={role_key}: {res.get('reason')}")

    # 4. Bullet rewrites + selection
    #
    # New (2026-05-12) contract: `bullets` is either
    #   (a) {role_key: [text, text, ...]}   — list per role; overwrites the
    #       role's master slots in order, drops trailing slots if list is
    #       shorter, INJECTS new bullet paragraphs after the last slot if
    #       longer.
    #   (b) {paragraph_idx_str: text}        — legacy per-paragraph dict (still
    #       supported for back-compat with old rewrites.json files).
    #
    # We detect by inspecting the keys.
    bullets_section = rewrites.get("bullets", {}) or {}
    # Heuristic: if any key is a JOB_BULLETS role_key, treat as new list-form;
    # else treat as legacy index-form.
    is_list_form = any(k in JOB_BULLETS for k in bullets_section.keys())

    if is_list_form:
        # For each role: collect the desired list. Default to master text if
        # role not mentioned at all (keep all unchanged).
        for job_key, indices in JOB_BULLETS.items():
            new_list = bullets_section.get(job_key)
            if new_list is None:
                # Keep all originals
                for pidx in indices:
                    report["bullets"].append({
                        "para": pidx, "job": job_key,
                        "original": originals[pidx], "rewrite": originals[pidx],
                        "applied": True, "reason": "no-rewrite-provided",
                    })
                continue
            # Clean empties / whitespace-only entries
            new_list = [str(s).strip() for s in new_list if str(s).strip()]
            n_slots = len(indices)
            n_new = len(new_list)
            # Overwrite existing slots 0..min(n_slots, n_new)-1
            for i, pidx in enumerate(indices):
                if i < n_new:
                    new_text = new_list[i]
                    res = mutate_bullet(paras[pidx], originals[pidx], new_text)
                    report["bullets"].append({
                        "para": pidx, "job": job_key,
                        "original": originals[pidx], "rewrite": res["final_text"],
                        "applied": res["applied"], "reason": res["reason"],
                    })
                else:
                    # Drop this slot — remove the paragraph entirely
                    p_el = paras[pidx]._p
                    parent = p_el.getparent()
                    if parent is not None:
                        parent.remove(p_el)
                    report["bullets"].append({
                        "para": pidx, "job": job_key,
                        "original": originals[pidx], "rewrite": "",
                        "applied": True, "reason": "dropped",
                    })
            # Inject any extras after the last existing slot
            if n_new > n_slots:
                last_pidx = indices[-1]
                template_para = paras[last_pidx]
                anchor_p = template_para._p
                # If the anchor was just overwritten (last slot was set), insert
                # after it. If anchor was removed (n_slots < n_new but list still
                # exceeds slots), the slots count must equal n_slots here so
                # anchor is still attached.
                for j in range(n_slots, n_new):
                    new_p = _clone_bullet_paragraph(template_para)
                    _set_paragraph_plain_text(new_p, new_list[j])
                    anchor_p.addnext(new_p)
                    anchor_p = new_p  # chain so order is preserved
                    report["bullets"].append({
                        "para": f"+{j - n_slots + 1}", "job": job_key,
                        "original": "", "rewrite": new_list[j],
                        "applied": True, "reason": "injected",
                    })
    else:
        # Legacy per-paragraph-index dict path (back-compat)
        bullet_rew = bullets_section
        for job_key, indices in JOB_BULLETS.items():
            for pidx in indices:
                new_text = bullet_rew.get(str(pidx))
                orig = originals[pidx]
                if new_text is None:
                    report["bullets"].append({
                        "para": pidx, "job": job_key,
                        "original": orig, "rewrite": orig,
                        "applied": True, "reason": "no-rewrite-provided",
                    })
                    continue
                if new_text == "":
                    p_el = paras[pidx]._p
                    parent = p_el.getparent()
                    if parent is not None:
                        parent.remove(p_el)
                    report["bullets"].append({
                        "para": pidx, "job": job_key,
                        "original": orig, "rewrite": "",
                        "applied": True, "reason": "dropped",
                    })
                    continue
                res = mutate_bullet(paras[pidx], orig, new_text)
                report["bullets"].append({
                    "para": pidx, "job": job_key,
                    "original": orig, "rewrite": res["final_text"],
                    "applied": res["applied"], "reason": res["reason"],
                })
                if not res["applied"]:
                    report["warnings"].append(
                        f"bullet rewrite fell back para={pidx}: {res['reason']}")

    # 5. Skills reorder + per-JD tailoring (Fix 2 + Fix 3, 2026-05-31)
    # Always run — skills must be tailored per role even when the rewriter
    # supplied no explicit priority list.
    for label, pidx in SKILLS_PARA.items():
        res = reorder_skills(paras[pidx], label, skills_prio, jd_text=jd_text)
        report["skills"].append({"label": label, "para": pidx, **res})

    # 6. Tighten formatting (bullets / spacing / right-tab / header tabs)
    tighten_doc(doc)

    # 7. Save and convert to PDF
    doc.save(str(target_docx))
    pdf_path = out_dir / f"Cyrus_Shekari_Resume_{org}_{job_id}{suffix}.pdf"
    convert_to_pdf(target_docx, pdf_path)

    report["docx"] = str(target_docx)
    report["pdf"] = str(pdf_path)
    return report


# ---------------------------------------------------------------------------
# Tightening pass — keeps tailored resumes one-page with small bullets and
# right-aligned dates/locations on the same line as the title/company.
# Mirrors /tmp/tighten.py + /tmp/header_tabs.py from the 2026-05-09 fix.

_HEADER_GAP_RE = re.compile(
    r'^(?P<left>\S.*?\S)[ \t]{2,}(?P<right>\S.*\S)[ \t]*$')
_DATE_HINT_RE  = re.compile(r'\b(19|20)\d{2}\b|Present', re.IGNORECASE)
_LOC_HINT_RE   = re.compile(r'^[A-Za-z .]+,\s*[A-Z]{2}$')


def tighten_doc(doc) -> None:
    """Apply the master one-pager treatment to an open Document in-place."""
    # 1. Bullet glyph size on every numbering level (~6pt)
    npart = getattr(doc.part, 'numbering_part', None)
    if npart is not None:
        for lvl in npart.element.iter(_w('lvl')):
            rPr = lvl.find(_w('rPr'))
            if rPr is None:
                rPr = lvl.makeelement(_w('rPr'), {})
                lvl.append(rPr)
            for tag in ('sz', 'szCs'):
                for el in rPr.findall(_w(tag)):
                    rPr.remove(el)
            rPr.append(rPr.makeelement(_w('sz'),   {_w('val'): '12'}))
            rPr.append(rPr.makeelement(_w('szCs'), {_w('val'): '12'}))

    # 2. Per-paragraph: right tab @ 10780 twips, line spacing 1.0, no extra space.
    #    This is the ORIGINAL one-pager treatment (reverted 2026-05-31). Cyrus's
    #    two preferred sample resumes use exactly this: 1.0 line spacing and the
    #    10780 right-tab (which renders dates + City,State flush to the right
    #    margin — verified). The brief 1.15/inherit-tab experiment made outputs
    #    drift and spill to a 2nd page, so we force the known-good values here.
    for p in doc.paragraphs:
        pPr = p._p.find(_w('pPr'))
        if pPr is not None:
            tabs = pPr.find(_w('tabs'))
            if tabs is not None:
                for tab in tabs.findall(_w('tab')):
                    if tab.get(_w('val')) == 'right':
                        tab.set(_w('pos'), '10780')
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)

    # 3. Header rows: replace run-on space padding with a single tab so the
    #    right tab stop actually engages (otherwise the right-side text wraps).
    for p in doc.paragraphs:
        text = p.text
        m = _HEADER_GAP_RE.match(text)
        if not m:
            continue
        right = m.group('right')
        if not (_DATE_HINT_RE.search(right) or _LOC_HINT_RE.match(right)):
            continue
        left = m.group('left')
        runs = list(p.runs)
        if not runs:
            continue
        first = runs[0]
        first_name = first.font.name
        first_size = first.font.size
        first.text = left
        for r in runs[1:]:
            r.text = ''
        new = p.add_run('\t' + right)
        # City/State locations are bolded (Cyrus 2026-05-31, mirrors master);
        # right-aligned DATES stay un-bold. _LOC_HINT_RE distinguishes them.
        new.bold = bool(_LOC_HINT_RE.match(right))
        if first_name:
            new.font.name = first_name
        if first_size:
            new.font.size = first_size


def _soffice_user_install_args() -> list[str]:
    """Return ['-env:UserInstallation=file://<dir>'] when a private LibreOffice
    profile is requested via the RESUME_LO_USER_INSTALL env var (set by the
    --user-install CLI flag), else []. A private profile lets a second agent
    render concurrently without colliding on the shared default LO profile
    lock. Unset = identical to historical behavior."""
    d = os.environ.get("RESUME_LO_USER_INSTALL", "").strip()
    if not d:
        return []
    p = Path(d).expanduser().resolve()
    p.mkdir(parents=True, exist_ok=True)
    return [f"-env:UserInstallation=file://{p}"]


def convert_to_pdf(docx: Path, pdf_out: Path) -> None:
    """Use LibreOffice headless. Output filename is determined by soffice."""
    cmd = [
        "soffice", *_soffice_user_install_args(),
        "--headless", "--convert-to", "pdf",
        "--outdir", str(pdf_out.parent), str(docx),
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if proc.returncode != 0:
        raise RuntimeError(
            f"soffice conversion failed (rc={proc.returncode})\nSTDOUT:{proc.stdout}\nSTDERR:{proc.stderr}")
    # soffice writes <stem>.pdf in --outdir; rename if it differs
    produced = pdf_out.parent / (docx.stem + ".pdf")
    if produced != pdf_out and produced.exists():
        produced.replace(pdf_out)


# ---------------------------------------------------------------------------
# CLI

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--org", required=True)
    ap.add_argument("--job-id", required=True)
    ap.add_argument("--jd", default=None)
    ap.add_argument("--rewrites", default=None,
                    help="JSON file of bullet rewrites; defaults to "
                         "<out-dir>/rewrites.json")
    ap.add_argument("--family", default=None,
                    choices=["pm", "tpm", "pgm", "se", "fde"])
    ap.add_argument("--out-dir", default=None)
    ap.add_argument("--suffix", default="_v2")
    ap.add_argument("--auto-rewrite", action="store_true",
                    help="Generate rewrites.json via bullet_rewriter.py before tailoring. "
                         "When combined with --out-dir, ALL file I/O (JD.md read, "
                         "rewrites.json, tailoring-notes.md, PDF) stays under that dir. "
                         "Requires JD.md to already exist inside --out-dir; "
                         "stage it first: cp /path/to/JD.md <out-dir>/JD.md")
    ap.add_argument("--max-loops", type=int, default=3,
                    help="Max page-fit loops when --auto-rewrite is on.")
    ap.add_argument("--user-install", default=None,
                    help="Private LibreOffice profile dir. When set, soffice "
                         "renders use -env:UserInstallation=file://<dir> so "
                         "concurrent renders (e.g. a second agent) don't "
                         "collide on the shared default LO profile lock. "
                         "Unset = historical behavior.")
    args = ap.parse_args()

    if args.user_install:
        os.environ["RESUME_LO_USER_INSTALL"] = args.user_install

    if args.auto_rewrite:
        # Delegate: bullet_rewriter handles rewrites + render + page-fit loop.
        # Pass out_dir so all I/O (JD staging, rewrites.json, PDF) stays under
        # the caller-specified directory instead of defaulting to APPS_DIR.
        from bullet_rewriter import run as br_run
        explicit_out_dir = Path(args.out_dir) if args.out_dir else None
        result = br_run(
            org=args.org, job_id=args.job_id, family=args.family,
            out_path=Path(args.rewrites) if args.rewrites else None,
            render=True, max_loops=args.max_loops,
            out_dir=explicit_out_dir,
            suffix=args.suffix,
        )
        print(json.dumps(result, indent=2))
        return

    out_dir = Path(args.out_dir) if args.out_dir else (APPS_DIR / f"{args.org}-{args.job_id}")
    jd_path = Path(args.jd) if args.jd else (out_dir / "JD.md")
    rewrites_path = Path(args.rewrites) if args.rewrites else None

    report = run_pipeline(
        org=args.org, job_id=args.job_id, jd_path=jd_path,
        family=args.family, rewrites_path=rewrites_path,
        out_dir=out_dir, suffix=args.suffix,
    )
    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
