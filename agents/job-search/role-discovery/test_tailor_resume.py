"""Regression tests for tailor_resume.py standing behavior changes (2026-05-31).

Covers the three Cyrus-directed fixes:
  1. Headline title mirrors the JD's ACTUAL role title (no blanket
     "Technical Product Manager" default).
  2. Skills section is tailored per-JD (relevant subset, drop irrelevant,
     vary role-to-role) using ONLY master-inventory skills (no fabrication).
  3. Skills bolding: only the group label is bold; the body list is not.
"""
import re

import tailor_resume as t


# ---------------------------------------------------------------------------
# Fix 1 — faithful headline title

def test_plain_pm_resolves_to_product_manager_not_technical():
    fam = t.detect_family("Product Manager, Payments")
    assert fam == "pm"
    assert t.resolve_headline_title("Product Manager, Payments", fam) == "Product Manager"


def test_senior_pm_still_plain_product_manager():
    fam = t.detect_family("Senior Product Manager")
    assert t.resolve_headline_title("Senior Product Manager", fam) == "Product Manager"


def test_technical_product_manager_faithful():
    fam = t.detect_family("Technical Product Manager")
    assert t.resolve_headline_title("Technical Product Manager", fam) == "Technical Product Manager"


def test_technical_program_manager_faithful():
    assert t.resolve_headline_title(
        "Technical Program Manager", t.detect_family("Technical Program Manager")
    ) == "Technical Program Manager"


def test_program_manager_faithful():
    assert t.resolve_headline_title(
        "Program Manager, Infrastructure", t.detect_family("Program Manager, Infrastructure")
    ) == "Program Manager"


def test_tpm_acronym_maps_to_technical_program_manager():
    assert t.resolve_headline_title("TPM - Platform", t.detect_family("TPM - Platform")) \
        == "Technical Program Manager"


def test_non_pm_families_get_no_swap():
    for title in ("Forward Deployed Engineer", "Solutions Engineer", "Solutions Architect"):
        fam = t.detect_family(title)
        assert t.resolve_headline_title(title, fam) is None


def test_resolved_title_always_in_allowlist():
    for title in ("Product Manager", "Technical Product Manager",
                  "Technical Program Manager", "Program Manager"):
        lbl = t.resolve_headline_title(title, t.detect_family(title))
        assert lbl in t.ALLOWED_TITLE_LABELS["microsoft_ft"]


# ---------------------------------------------------------------------------
# Fix 2 (2026-06-01) — single-track coherence across ALL title slots

_MIXED_SWAPS = {
    "microsoft_ft":     "Technical Product Manager",
    "microsoft_2023":   "Technical Product Manager Intern",
    "microsoft_2022":   "Technical Program Manager Intern",   # off-track
    "amazon_robotics":  "Technical Program Manager Intern",   # off-track
    "pro_painters":     "Product Manager Intern",
}


def test_coerce_snaps_all_slots_to_product_track():
    out = t.coerce_title_track(_MIXED_SWAPS, "Technical Product Manager")
    assert out["microsoft_ft"] == "Technical Product Manager"
    for k in ("microsoft_2023", "microsoft_2022", "amazon_robotics", "pro_painters"):
        assert out[k] == "Technical Product Manager Intern", (k, out[k])
    # no within-resume Program/Product mixing
    assert all("Program" not in v for v in out.values())


def test_coerce_snaps_all_slots_to_program_track():
    out = t.coerce_title_track(_MIXED_SWAPS, "Technical Program Manager")
    assert out["microsoft_ft"] == "Technical Program Manager"
    for k in ("microsoft_2023", "microsoft_2022", "amazon_robotics", "pro_painters"):
        assert out[k] == "Technical Program Manager Intern", (k, out[k])
    assert all("Product" not in v for v in out.values())


def test_coerce_non_technical_track():
    out = t.coerce_title_track(_MIXED_SWAPS, "Product Manager")
    assert out["microsoft_ft"] == "Product Manager"
    assert out["microsoft_2023"] == "Product Manager Intern"
    assert all("Technical" not in v for v in out.values())


def test_coerce_ambiguous_defaults_to_program():
    # A non-PM-specific headline falls back to the Program track.
    out = t.coerce_title_track(_MIXED_SWAPS, "Manager")
    assert out["microsoft_ft"] == "Program Manager"
    assert all("Product" not in v for v in out.values())


def test_coerce_output_always_in_allowlist():
    for headline in ("Product Manager", "Technical Product Manager",
                     "Technical Program Manager", "Program Manager"):
        out = t.coerce_title_track(_MIXED_SWAPS, headline)
        assert out["microsoft_ft"] in t.ALLOWED_TITLE_LABELS["microsoft_ft"]
        for k in ("microsoft_2023", "microsoft_2022", "amazon_robotics", "pro_painters"):
            assert out[k] in t.ALLOWED_TITLE_LABELS[k], (headline, k, out[k])


def test_coerce_does_not_mutate_input():
    src = dict(_MIXED_SWAPS)
    t.coerce_title_track(src, "Technical Program Manager")
    assert src == _MIXED_SWAPS


# ---------------------------------------------------------------------------
# Fix 2 (revised 2026-05-31) — skills GENERATED from the JD, no inventory gate

# ---------------------------------------------------------------------------
# Fix 2 (FINAL reconciled 2026-05-31) — BLEND: base skills + JD-relevant augment

# A realistic base for the Technical group (mirrors the master .docx body).
_BASE_TECH = ["Azure", "distributed systems", "CI/CD", "APIs",
              "data pipelines", "Power BI", "YAML", "SQL"]
_BASE_PP = ["Technical Program Management", "Product requirements", "Roadmapping",
            "Cross-functional execution", "Stakeholder management", "Agile/Scrum"]


def test_blend_keeps_base_identity_when_jd_sparse():
    # JD mentions nothing technical -> base skills carry the line (identity kept).
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", "We value teamwork.")
    assert set(out).issubset(set(_BASE_TECH))  # nothing invented when JD adds nothing
    assert len(out) >= t._SKILL_MIN_PER_GROUP


def test_blend_augments_with_jd_skill_not_in_base():
    # Kubernetes is NOT a base skill; JD asks for it -> it gets added (augment).
    jd = "You will run Kubernetes clusters and Docker containers daily."
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", jd)
    assert "Kubernetes" in out  # JD augment, even though not in base
    # base identity not wholesale-wiped: at least one base skill survives
    assert any(b in out for b in _BASE_TECH)


def test_blend_jd_relevant_base_floats_first():
    # SQL + APIs are base AND in JD -> they should lead (tier1), ahead of
    # non-JD base skills like YAML.
    jd = "Heavy SQL and REST API work."
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", jd)
    assert out.index("SQL") < out.index("YAML")
    assert out.index("APIs") < out.index("YAML")


def test_jd_stretch_skill_outranks_nonjd_base():
    # Option B (2026-06-01): a JD-wanted catalog skill Cyrus does NOT list in
    # base must rank ABOVE a base skill the JD didn't mention. Python is a
    # high-catalog-priority stretch skill (not in _BASE_TECH); YAML is base but
    # not in this JD -> Python must lead YAML.
    jd = "Strong Python scripting required; build data pipelines."
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", jd)
    assert "Python" in out, out
    assert "YAML" not in out[:out.index("Python")] if "YAML" in out else True
    if "YAML" in out:
        assert out.index("Python") < out.index("YAML"), out
    # And a JD catalog hit leads the line (JD-priority-first, not base-first).
    assert out[0] in ("Python", "Data pipelines"), out


def test_blend_jd_skill_can_replace_weak_base_when_full():
    # Dense JD: JD-relevant skills should crowd out weaker non-JD base skills
    # rather than overflow past the cap.
    jd = ("python kubernetes docker aws gcp react typescript microservices "
          "devops machine learning")
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", jd)
    assert len(out) <= t._SKILL_MAX_PER_GROUP
    # a JD-relevant non-base skill made it in
    assert any(s in out for s in ("Kubernetes", "Docker", "AWS", "Python", "React"))
    # a weak non-JD base skill (YAML, not in this JD) got displaced
    assert "YAML" not in out


def test_blend_claims_jd_skill_even_if_not_in_profile():
    # Cyrus's accepted risk: claim a JD skill that's neither base nor 'had'.
    jd = "Deep AWS Lambda and GCP experience required."
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", jd)
    assert "AWS" in out and "GCP" in out


def test_blend_vary_between_roles():
    jd_pm = "Own product roadmaps, write PRDs, drive go-to-market and A/B testing."
    jd_tpm = "Coordinate Azure distributed systems, CI/CD, Kubernetes reliability."
    pm_tech = t.merge_skills_for_group(_BASE_TECH, "Technical", jd_pm)
    tpm_tech = t.merge_skills_for_group(_BASE_TECH, "Technical", jd_tpm)
    pm_pp = t.merge_skills_for_group(_BASE_PP, "Program / Product", jd_pm)
    assert pm_tech != tpm_tech  # per-role variation
    assert "Roadmapping" in pm_pp


def test_blend_respects_caps():
    jd = ("python java go c++ javascript typescript react azure aws gcp "
          "kubernetes docker ci/cd devops api microservice sql nosql tableau")
    out = t.merge_skills_for_group(_BASE_TECH, "Technical", jd)
    assert t._SKILL_MIN_PER_GROUP <= len(out) <= t._SKILL_MAX_PER_GROUP


def test_blend_min_floor_when_everything_sparse():
    out = t.merge_skills_for_group([], "Technical", "")
    assert len(out) >= t._SKILL_MIN_PER_GROUP  # fallback prevents empty line


def test_blend_dedups_near_duplicates():
    # base has 'semantic search'; JD triggers catalog 'RAG / Semantic search'.
    # Only ONE of the two near-dupes should survive (no redundant pair).
    base = ["semantic search", "workflow automation"]
    jd = "We use RAG retrieval and semantic search heavily."
    out = t.merge_skills_for_group(base, "AI / Automation", jd)
    sem = [s for s in out if "semantic search" in s.lower()]
    assert len(sem) == 1, f"expected one semantic-search variant, got {sem}"


# ---------------------------------------------------------------------------
# Fix 3 — only the group label is bold (run-level)

class _FakeFont:
    def __init__(self):
        self.name = "Calibri"
        self.size = None


class _FakeRun:
    def __init__(self, text, bold=None):
        self.text = text
        self.bold = bold
        self.font = _FakeFont()


class _FakeParagraph:
    def __init__(self, runs):
        self.runs = runs

    @property
    def text(self):
        return "".join(r.text for r in self.runs)

    def add_run(self, text):
        r = _FakeRun(text)
        self.runs.append(r)
        return r


def test_bolding_label_bold_body_not():
    p = _FakeParagraph([
        _FakeRun("Technical:", bold=True),
        _FakeRun(" Azure, SQL, APIs", bold=None),
    ])
    res = t.reorder_skills(p, "Technical", [], jd_text="Azure SQL APIs", tailor=True)
    assert res["applied"]
    # run0 = label, bold True; run1 = body, bold False
    assert p.runs[0].text == "Technical:" and p.runs[0].bold is True
    assert p.runs[1].bold is False
    assert p.runs[1].text.lstrip().startswith(("Azure", "SQL", "APIs"))


def test_bolding_single_run_appends_nonbold_body():
    p = _FakeParagraph([_FakeRun("Technical: Azure, SQL, APIs", bold=True)])
    res = t.reorder_skills(p, "Technical", [], jd_text="Azure SQL APIs", tailor=True)
    assert res["applied"]
    assert p.runs[0].text == "Technical:" and p.runs[0].bold is True
    assert len(p.runs) >= 2 and p.runs[1].bold is False


def test_body_has_no_bold_markers():
    p = _FakeParagraph([
        _FakeRun("Program / Product:", bold=True),
        _FakeRun(" Roadmapping, Agile/Scrum", bold=None),
    ])
    t.reorder_skills(p, "Program / Product", [], jd_text="roadmap agile", tailor=True)
    body = "".join(r.text for r in p.runs[1:])
    assert "**" not in body  # no markdown bold leakage


if __name__ == "__main__":
    import sys
    import pytest
    sys.exit(pytest.main([__file__, "-q"]))


# --- Location bold in tighten_doc (Cyrus 2026-05-31: bold city/state on every resume) ---

def test_tighten_bolds_city_state_location_not_dates():
    """tighten_doc must bold the right-aligned City, State token on header rows,
    while leaving right-aligned DATE tokens un-bold."""
    from docx import Document as _Doc
    import tailor_resume as tr
    doc = _Doc()
    # Company + location row (space-padded like the master), single bold run
    p_loc = doc.add_paragraph()
    r = p_loc.add_run("Microsoft" + " " * 40 + "Seattle, WA")
    r.bold = True
    # Title + date row, single bold run
    p_date = doc.add_paragraph()
    r2 = p_date.add_run("Product Manager" + " " * 30 + "March 2024 - Present")
    r2.bold = True
    tr.tighten_doc(doc)
    # Location row: trailing run is the City, State and must be bold
    loc_runs = [(rr.text.strip(), rr.bold) for rr in p_loc.runs if rr.text.strip()]
    assert ("Seattle, WA", True) in loc_runs, loc_runs
    # Date row: trailing date run must NOT be bold
    date_runs = [(rr.text.strip(), rr.bold) for rr in p_date.runs if rr.text.strip()]
    trailing = [b for t, b in date_runs if "2024" in t]
    assert trailing and trailing[0] is False, date_runs
